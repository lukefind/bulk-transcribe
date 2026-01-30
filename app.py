#!/usr/bin/env python3
"""
Web UI for Audio Transcription Tool using OpenAI Whisper
"""

# Prevent PyTorch from spawning subprocesses (causes duplicate dock icons on macOS)
import os
import sys
from pathlib import Path

# Set up ffmpeg path if running in PyInstaller bundle
if getattr(sys, 'frozen', False):
    base_path = Path(sys._MEIPASS)
    # PyInstaller puts binaries in different locations depending on structure
    possible_bin_paths = [
        base_path / 'bin',
        base_path / 'Frameworks' / 'bin',
        Path(sys.executable).parent / 'bin',
        Path(sys.executable).parent.parent / 'Frameworks' / 'bin',
    ]
    for bin_path in possible_bin_paths:
        if bin_path.exists() and (bin_path / 'ffmpeg').exists():
            os.environ['PATH'] = str(bin_path) + os.pathsep + os.environ.get('PATH', '')
            break

os.environ['OMP_NUM_THREADS'] = '1'
os.environ['MKL_NUM_THREADS'] = '1'
os.environ['OPENBLAS_NUM_THREADS'] = '1'
os.environ['VECLIB_MAXIMUM_THREADS'] = '1'
os.environ['NUMEXPR_NUM_THREADS'] = '1'
os.environ['TOKENIZERS_PARALLELISM'] = 'false'

# Suppress noisy third-party warnings early (before any imports that trigger them)
from logger import configure_runtime_noise, log_event
configure_runtime_noise()

# Prevent multiprocessing from creating new app instances
import multiprocessing
multiprocessing.set_start_method('fork', force=True)

import subprocess
import threading
import queue
import secrets
from dataclasses import replace
import re
from flask import Flask, render_template, request, jsonify, send_file, g, Response
import zipfile
import io
import tempfile
import shutil
from werkzeug.utils import secure_filename
from werkzeug.middleware.proxy_fix import ProxyFix
import json
from datetime import datetime, timezone
from typing import Optional

from transcribe_options import TranscribeOptions, get_preset, postprocess_segments
import session_store
import compute_backend

app = Flask(__name__)

# Apply ProxyFix for HTTPS detection behind reverse proxy (Caddy, nginx)
# This ensures request.is_secure works correctly
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1, x_host=1, x_prefix=1)

# Configure max upload size from environment
app.config['MAX_CONTENT_LENGTH'] = session_store.get_max_upload_mb() * 1024 * 1024


@app.errorhandler(413)
def request_entity_too_large(error):
    """Handle upload size limit exceeded."""
    max_mb = session_store.get_max_upload_mb()
    return jsonify({
        'error': f'File too large. Maximum upload size is {max_mb}MB.'
    }), 413

SUPPORTED_FORMATS = ['.mp3', '.wav', '.m4a', '.flac', '.ogg', '.webm', '.mp4', '.mov', '.avi']
AVAILABLE_MODELS = ['tiny', 'base', 'small', 'medium', 'large', 'large-v3']
MODEL_INFO = {
    'tiny': {'size': '~75 MB', 'ram': '~1 GB', 'speed': 'Fastest', 'accuracy': 'Basic'},
    'base': {'size': '~145 MB', 'ram': '~1 GB', 'speed': 'Fast', 'accuracy': 'Good'},
    'small': {'size': '~465 MB', 'ram': '~2 GB', 'speed': 'Medium', 'accuracy': 'Better'},
    'medium': {'size': '~1.5 GB', 'ram': '~5 GB', 'speed': 'Slow', 'accuracy': 'High'},
    'large': {'size': '~3 GB', 'ram': '~10 GB', 'speed': 'Slowest', 'accuracy': 'Best'},
    'large-v3': {'size': '~3 GB', 'ram': '~10 GB', 'speed': 'Slowest', 'accuracy': 'Best'},
}
COMMON_LANGUAGES = [
    ('', 'Auto-detect'),
    ('en', 'English'),
    ('es', 'Spanish'),
    ('fr', 'French'),
    ('de', 'German'),
    ('it', 'Italian'),
    ('pt', 'Portuguese'),
    ('ja', 'Japanese'),
    ('ko', 'Korean'),
    ('zh', 'Chinese'),
    ('ru', 'Russian'),
    ('ar', 'Arabic'),
    ('hi', 'Hindi'),
]

transcription_status = {
    'running': False,
    'cancelled': False,
    'cancel_after_file': False,  # Stop after current file(s) complete
    'current_file': '',
    'current_file_progress': '',
    'current_file_percent': 0,
    'current_file_start': None,  # Timestamp when current file started
    'current_file_elapsed': 0,  # Seconds elapsed on current file
    'last_transcribed_text': '',
    'completed': 0,
    'total': 0,
    'results': [],
    'error': None,
    'downloading_model': False,
    'active_jobs': 0,
    'active_workers': 0,
    'current_files': [],  # List of files currently being processed (for parallel mode)
    'start_time': None,  # Job start timestamp
    'elapsed_seconds': 0,  # Total elapsed time
}

# Thread-safe lock for updating transcription_status
import threading
status_lock = threading.Lock()

# Request counter for periodic cleanup
_request_counter = 0
_cleanup_interval = 50  # Run cleanup every N requests


def _cleanup_orphaned_running_jobs():
    """
    Mark any jobs with status='running' as failed on startup.
    
    This handles the case where the server was restarted while jobs were running.
    Those jobs are now orphaned (no thread processing them) and should be marked
    as failed so the UI doesn't show them as stuck.
    """
    if not session_store.is_server_mode():
        return
    
    try:
        sessions_dir = session_store.get_sessions_dir()
        if not os.path.exists(sessions_dir):
            return
        
        orphaned_count = 0
        for session_id in os.listdir(sessions_dir):
            session_path = os.path.join(sessions_dir, session_id)
            if not os.path.isdir(session_path):
                continue
            
            jobs_dir = os.path.join(session_path, 'jobs')
            if not os.path.exists(jobs_dir):
                continue
            
            for job_id in os.listdir(jobs_dir):
                job_path = os.path.join(jobs_dir, job_id)
                if not os.path.isdir(job_path):
                    continue
                
                manifest_path = os.path.join(job_path, 'manifest.json')
                if not os.path.exists(manifest_path):
                    continue
                
                try:
                    manifest = session_store.read_json(manifest_path)
                    if manifest and manifest.get('status') == 'running':
                        # This job was running when server stopped - mark as failed
                        manifest['status'] = 'failed'
                        manifest['finishedAt'] = datetime.now(timezone.utc).isoformat()
                        manifest['error'] = {
                            'code': 'SERVER_RESTART',
                            'message': 'Job was interrupted by server restart. You can re-run this job using the retry button.'
                        }
                        session_store.atomic_write_json(manifest_path, manifest)
                        orphaned_count += 1
                except Exception:
                    pass  # Skip jobs we can't read
        
        if orphaned_count > 0:
            print(f"[startup] Marked {orphaned_count} orphaned running job(s) as failed")
    except Exception as e:
        print(f"[startup] Warning: Failed to cleanup orphaned jobs: {e}")


# Run cleanup on module load (server startup)
_cleanup_orphaned_running_jobs()


# =============================================================================
# Session Middleware
# =============================================================================

@app.before_request
def before_request_session():
    """Ensure session ID exists for all requests."""
    global _request_counter
    
    # Get or generate session ID
    raw_session_id = request.cookies.get(session_store.COOKIE_NAME)
    
    # Handle shared sessions: "shared:<mode>:<real_session_id>"
    g.share_mode = None
    g.is_shared_session = False
    
    if raw_session_id and raw_session_id.startswith('shared:'):
        parts = raw_session_id.split(':', 2)
        if len(parts) == 3:
            g.share_mode = parts[1]  # 'read' or 'edit'
            session_id = parts[2]    # The actual session ID
            g.is_shared_session = True
        else:
            # Malformed, generate new
            session_id = session_store.new_id(32)
            g._new_session_id = session_id
    elif raw_session_id:
        session_id = raw_session_id
    else:
        session_id = session_store.new_id(32)
        g._new_session_id = session_id
    
    g.session_id = session_id
    
    # Periodic cleanup (exclude current session to prevent mid-request deletion)
    _request_counter += 1
    if _request_counter >= _cleanup_interval:
        _request_counter = 0
        try:
            session_store.cleanup_expired_sessions(exclude_session_id=session_id)
        except Exception:
            pass  # Don't fail requests due to cleanup errors


@app.after_request
def after_request_session(response):
    """Set session cookie if new, touch session metadata, and add security headers."""
    # Set cookie for new sessions
    if hasattr(g, '_new_session_id'):
        session_store.set_new_session_cookie(request, response, g._new_session_id)
    
    # Touch session to update lastSeenAt (skip for static files and healthz)
    if hasattr(g, 'session_id') and not request.path.startswith('/static') and request.path != '/healthz':
        try:
            session_store.touch_session(g.session_id)
        except Exception:
            pass  # Don't fail requests due to session touch errors
    
    # Security headers for API endpoints and downloads
    if request.path.startswith('/api/'):
        response.headers['Cache-Control'] = 'no-store'
        response.headers['Pragma'] = 'no-cache'
        response.headers['X-Content-Type-Options'] = 'nosniff'
    
    return response


# CSRF-protected endpoints in server mode
_CSRF_PROTECTED_ENDPOINTS = {
    'api_upload_files',
    'api_create_job',
    'api_cancel_job',
}


def _require_csrf():
    """Check CSRF token for protected endpoints in server mode."""
    if not session_store.is_server_mode():
        return None  # No CSRF in local mode
    
    if request.endpoint not in _CSRF_PROTECTED_ENDPOINTS:
        return None
    
    token = request.headers.get('X-CSRF-Token', '')
    session_id = getattr(g, 'session_id', None)
    
    if not session_id or not session_store.validate_csrf_token(session_id, token):
        return jsonify({'error': 'Invalid or missing CSRF token'}), 403
    
    return None


# Endpoints blocked for read-only shared sessions
_READ_ONLY_BLOCKED_ENDPOINTS = {
    'api_upload_files',
    'api_create_job',
    'api_cancel_job',
    'api_cancel_after_current',
    'api_delete_job',
    'api_mark_job_finished',
    'api_bulk_delete_jobs',
    'api_clear_jobs',
    'api_clear_duplicate_jobs',
    'api_rerun_job',
    'api_put_review_state',
    'api_update_speaker_labels',
    'api_share_session',  # Can't create shares from shared session
    'api_revoke_share',
}


def _check_read_only():
    """Block write operations for read-only shared sessions."""
    if not getattr(g, 'is_shared_session', False):
        return None
    
    if getattr(g, 'share_mode', None) != 'read':
        return None  # Edit mode allowed
    
    if request.endpoint in _READ_ONLY_BLOCKED_ENDPOINTS:
        return jsonify({'error': 'This session is read-only'}), 403
    
    return None


@app.before_request
def before_request_csrf():
    """Validate CSRF token for protected endpoints."""
    result = _require_csrf()
    if result:
        return result


@app.before_request
def before_request_read_only():
    """Block write operations for read-only shared sessions."""
    result = _check_read_only()
    if result:
        return result


def choose_folder_dialog(prompt="Select a folder"):
    """Open native macOS folder picker dialog."""
    script = f'POSIX path of (choose folder with prompt "{prompt}")'
    try:
        result = subprocess.run(['osascript', '-e', script], capture_output=True, text=True, timeout=120)
        if result.returncode == 0:
            return result.stdout.strip()
    except Exception:
        pass
    return None


def find_audio_files(input_folder: str) -> list:
    """Find all supported audio files in the input folder."""
    audio_files = []
    input_path = Path(input_folder)
    
    if not input_path.exists():
        return []
    
    for file_path in input_path.rglob('*'):
        if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_FORMATS:
            audio_files.append(file_path)
    
    return sorted(audio_files)


def get_whisper_cache_dir():
    """Get the whisper model cache directory."""
    return os.path.expanduser('~/.cache/whisper')


def get_preferences_file():
    """Get the path to the user preferences file."""
    config_dir = os.path.expanduser('~/.config/bulk-transcribe')
    os.makedirs(config_dir, exist_ok=True)
    return os.path.join(config_dir, 'preferences.json')


def load_preferences():
    """Load user preferences from file."""
    prefs_file = get_preferences_file()
    default_prefs = {
        'default_model': None,
        'default_language': '',
        'include_segments': True,
        'include_full': False,
        'include_timestamps': True,
        'word_timestamps': False,
    }
    try:
        if os.path.exists(prefs_file):
            with open(prefs_file, 'r') as f:
                saved = json.load(f)
                default_prefs.update(saved)
    except Exception:
        pass
    return default_prefs


def save_preferences(prefs):
    """Save user preferences to file."""
    prefs_file = get_preferences_file()
    try:
        with open(prefs_file, 'w') as f:
            json.dump(prefs, f, indent=2)
        return True
    except Exception:
        return False


# Mapping from model name to actual filename in cache
MODEL_FILENAMES = {
    'tiny': 'tiny.pt',
    'base': 'base.pt',
    'small': 'small.pt',
    'medium': 'medium.pt',
    'large': 'large.pt',
    'large-v3': 'large-v3.pt',
}


def get_model_filepath(model: str) -> str:
    """Get the full path to a model file."""
    cache_dir = get_whisper_cache_dir()
    filename = MODEL_FILENAMES.get(model, f'{model}.pt')
    return os.path.join(cache_dir, filename)


def check_model_available(model: str) -> bool:
    """Check if a whisper model is already downloaded."""
    model_file = get_model_filepath(model)
    return os.path.exists(model_file)


def get_model_file_size(model: str) -> str:
    """Get the actual file size of a downloaded model."""
    model_file = get_model_filepath(model)
    if os.path.exists(model_file):
        size_bytes = os.path.getsize(model_file)
        if size_bytes >= 1024 * 1024 * 1024:
            return f"{size_bytes / (1024 * 1024 * 1024):.1f} GB"
        else:
            return f"{size_bytes / (1024 * 1024):.0f} MB"
    return None


def get_all_models_status():
    """Get status of all available models."""
    models = []
    for model_name in AVAILABLE_MODELS:
        info = MODEL_INFO.get(model_name, {})
        is_downloaded = check_model_available(model_name)
        actual_size = get_model_file_size(model_name) if is_downloaded else None
        models.append({
            'name': model_name,
            'downloaded': is_downloaded,
            'size': actual_size or info.get('size', 'Unknown'),
            'ram': info.get('ram', 'Unknown'),
            'speed': info.get('speed', 'Unknown'),
            'accuracy': info.get('accuracy', 'Unknown'),
        })
    return models


def transcribe_file(audio_file: Path, model: str, language: str, output_folder: Path, 
                    word_timestamps: bool = False, status_callback=None,
                    transcribe_options: TranscribeOptions = None) -> dict:
    """Transcribe a single audio file using Whisper Python library with progress callback."""
    global transcription_status
    
    import time
    import threading
    
    start_time = time.time()
    
    transcription_status['current_file_progress'] = f'Loading model ({model})...'
    transcription_status['current_file_percent'] = 0
    # Do not clear last_transcribed_text here; keep the last snippet visible during processing
    # and update it when we have new text.
    
    try:
        import whisper
        
        whisper_model = None
        if status_callback is not None and isinstance(status_callback, dict):
            whisper_model = status_callback.get('whisper_model')

        if whisper_model is None:
            transcription_status['current_file_progress'] = f'Loading {model} model...'
            transcription_status['current_file_percent'] = 5
            # MPS has bugs with some audio files, use CPU for stability
            whisper_model = whisper.load_model(model)
        
        transcription_status['current_file_progress'] = 'Transcribing audio...'
        transcription_status['current_file_percent'] = 10
        
        # Check for cancellation before transcription
        if transcription_status.get('cancelled'):
            return {'error': 'Cancelled by user'}
        
        # Get audio duration for progress estimation
        import whisper.audio
        audio_array = whisper.audio.load_audio(str(audio_file))
        audio_duration = len(audio_array) / whisper.audio.SAMPLE_RATE
        
        # Background thread for progress updates based on segments
        stop_progress = threading.Event()
        segments_received = []
        
        def progress_update():
            while not stop_progress.is_set():
                if transcription_status.get('cancelled'):
                    return
                elapsed = time.time() - start_time
                elapsed_str = f"{int(elapsed // 60)}m {int(elapsed % 60)}s" if elapsed >= 60 else f"{int(elapsed)}s"
                
                # Estimate progress based on audio duration
                # CPU on Apple Silicon: ~0.5-1x realtime
                # With parallel workers, each file is slower due to CPU contention
                if audio_duration > 0:
                    estimated_total_time = audio_duration * 1.5  # Assume 0.67x realtime on CPU
                    estimated_pct = min(95, int(10 + (elapsed / estimated_total_time) * 85))
                else:
                    # Fallback: assume 30 seconds per percent after initial 10%
                    estimated_pct = min(95, 10 + int(elapsed / 30))
                
                # Only update if we're the current/only file being processed
                # In parallel mode, don't fight over the progress bar
                with status_lock:
                    if len(transcription_status.get('current_files', [])) <= 1:
                        transcription_status['current_file_percent'] = estimated_pct
                        transcription_status['current_file_progress'] = f'Transcribing... ({elapsed_str} elapsed)'
                
                stop_progress.wait(1.0)  # Check less frequently
        
        progress_thread = threading.Thread(target=progress_update, daemon=True)
        progress_thread.start()
        
        # Build transcription options (clone per call to avoid cross-thread mutation)
        if transcribe_options is None:
            transcribe_options = get_preset('balanced')
        else:
            transcribe_options = replace(transcribe_options)

        # Override with function parameters
        transcribe_options.language = language if language else None
        transcribe_options.word_timestamps = word_timestamps

        # Transcribe with explicit parameters
        whisper_kwargs = transcribe_options.to_whisper_kwargs()
        result = whisper.transcribe(
            whisper_model,
            str(audio_file),
            **whisper_kwargs
        )
        
        # Post-process segments to improve quality
        if result.get('segments'):
            result['segments'] = postprocess_segments(
                result['segments'],
                merge_short=transcribe_options.merge_short_segments,
                min_duration=transcribe_options.min_segment_duration,
                max_duration=transcribe_options.max_segment_duration,
                word_timestamps_available=word_timestamps and any(
                    s.get('words') for s in result['segments']
                ),
            )
        
        stop_progress.set()
        
        # Check for cancellation after transcription
        if transcription_status.get('cancelled'):
            return {'error': 'Cancelled by user'}
        
        transcription_status['current_file_progress'] = 'Saving results...'
        transcription_status['current_file_percent'] = 98
        
        # Update last transcribed text with final segment (thread-safe)
        if result.get('segments') and len(result['segments']) > 0:
            last_segment = result['segments'][-1]
            snippet = last_segment.get('text', '').strip()[:150]
            print(f"[DEBUG] Setting snippet: '{snippet[:50]}...' from {len(result['segments'])} segments")
            if snippet:
                with status_lock:
                    transcription_status['last_transcribed_text'] = snippet
        else:
            print(f"[DEBUG] No segments in result: {result.keys() if isinstance(result, dict) else type(result)}")
        
        # Save JSON output
        json_output_path = output_folder / (audio_file.stem + '.json')
        with open(json_output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        
        elapsed = time.time() - start_time
        elapsed_str = f"{int(elapsed // 60)}m {int(elapsed % 60)}s" if elapsed >= 60 else f"{int(elapsed)}s"
        transcription_status['current_file_progress'] = f'Complete ({elapsed_str})'
        transcription_status['current_file_percent'] = 100
        
        return result
        
    except Exception as e:
        import traceback
        error_msg = str(e)
        print(f"Transcription error: {error_msg}")
        print(traceback.format_exc())
        if 'CUDA' in error_msg or 'cuda' in error_msg:
            error_msg = 'GPU/CUDA error - try restarting the app'
        elif 'out of memory' in error_msg.lower():
            error_msg = 'Out of memory - try a smaller model'
        elif 'No such file' in error_msg or 'FileNotFoundError' in error_msg:
            error_msg = f'Audio file not found: {error_msg[:100]}'
        return {'error': error_msg[:200]}


def format_timestamp(seconds: float) -> str:
    """Convert seconds to HH:MM:SS format."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = seconds % 60
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{secs:05.2f}"
    return f"{minutes:02d}:{secs:05.2f}"


def create_markdown(audio_file: Path, transcription_data: dict, model: str, 
                   language: str, output_folder: Path, options: dict = None) -> Path:
    """Create a markdown file from transcription data."""
    if options is None:
        options = {}
    
    include_full = options.get('include_full', True)
    include_segments = options.get('include_segments', True)
    include_timestamps = options.get('include_timestamps', True)
    include_word_timestamps = options.get('word_timestamps', False)
    
    output_name = audio_file.stem + '_transcription.md'
    output_path = output_folder / output_name
    
    content = f"# Transcription: {audio_file.name}\n\n"
    content += f"**Source File:** `{audio_file.name}`\n"
    content += f"**Model Used:** `{model}`\n"
    
    if language:
        content += f"**Language:** `{language}`\n"
    
    if 'language' in transcription_data:
        content += f"**Detected Language:** `{transcription_data['language']}`\n"
    
    # Add segment count for reference
    segment_count = len(transcription_data.get('segments', []))
    if segment_count > 0:
        content += f"**Segments:** `{segment_count}`\n"
    
    content += f"\n---\n\n"
    
    if include_full and 'text' in transcription_data:
        content += "## Full Transcription\n\n"
        content += transcription_data['text'].strip() + "\n\n"
    
    if include_segments and 'segments' in transcription_data:
        content += "## Segmented Transcription\n\n"
        for i, segment in enumerate(transcription_data['segments'], 1):
            start_time = segment.get('start', 0)
            end_time = segment.get('end', 0)
            text = segment.get('text', '').strip()
            
            if include_timestamps:
                content += f"**[{format_timestamp(start_time)} - {format_timestamp(end_time)}]**\n\n"
            
            content += f"{text}\n\n"
            
            # Include word-level timestamps if available and requested
            if include_word_timestamps and segment.get('words'):
                words = segment['words']
                word_line = ' '.join(
                    f"`{w.get('word', '')}@{w.get('start', 0):.2f}`" 
                    for w in words if w.get('word')
                )
                if word_line:
                    content += f"<details><summary>Word timestamps</summary>\n\n{word_line}\n\n</details>\n\n"
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(content)
    
    return output_path


def run_transcription(input_folder: str, output_folder: str, model: str, language: str, options: dict = None):
    """Background task to run transcription with optional parallel workers."""
    global transcription_status
    
    if options is None:
        options = {}
    
    processing_mode = options.get('processing_mode', 'metal')
    word_timestamps = options.get('word_timestamps', False)
    
    # Compute effective mode: Metal + word_timestamps forces CPU (MPS doesn't support DTW float64)
    effective_mode = processing_mode
    if processing_mode == 'metal' and word_timestamps:
        effective_mode = 'cpu'
    
    # Only force single worker when actually using Metal GPU; allow multiple for CPU
    if effective_mode == 'metal':
        num_workers = 1
    else:
        num_workers = max(1, min(2, options.get('workers', 1)))  # Max 2 workers for CPU mode
    
    import time as time_module
    transcription_status['running'] = True
    transcription_status['active_jobs'] = 1
    transcription_status['active_workers'] = 0
    transcription_status['cancelled'] = False
    transcription_status['cancel_after_file'] = False
    transcription_status['start_time'] = time_module.time()
    transcription_status['elapsed_seconds'] = 0
    transcription_status['last_transcribed_text'] = ''
    transcription_status['completed'] = 0
    transcription_status['results'] = []
    transcription_status['error'] = None
    transcription_status['current_files'] = []
    
    try:
        output_path = Path(output_folder)
        output_path.mkdir(parents=True, exist_ok=True)
        
        audio_files = find_audio_files(input_folder)
        transcription_status['total'] = len(audio_files)
        
        if not audio_files:
            transcription_status['error'] = 'No audio files found in the input folder'
            transcription_status['running'] = False
            return
        
        overwrite_existing = options.get('overwrite_existing', False)
        
        # Build TranscribeOptions from advanced settings
        quality_preset = options.get('quality_preset', 'balanced')
        transcribe_opts = get_preset(quality_preset)
        
        # Override with specific settings from UI
        if options.get('no_speech_threshold') is not None:
            transcribe_opts.no_speech_threshold = float(options['no_speech_threshold'])
        if word_timestamps and options.get('max_segment_duration') is not None:
            transcribe_opts.max_segment_duration = float(options['max_segment_duration'])
        else:
            transcribe_opts.max_segment_duration = None
        
        transcribe_opts.merge_short_segments = True
        transcribe_opts.min_segment_duration = 0.5
        
        # Filter out files that already have output (if not overwriting)
        files_to_process = []
        for audio_file in audio_files:
            expected_output = output_path / f"{audio_file.stem}_transcription.md"
            if expected_output.exists() and not overwrite_existing:
                with status_lock:
                    transcription_status['results'].append({
                        'file': audio_file.name,
                        'status': 'skipped',
                        'message': 'Output file already exists',
                        'elapsed_seconds': 0.0,
                        'audio_seconds': None,
                        'speed_x': None,
                        'segments_count': None,
                    })
                    transcription_status['completed'] += 1
            else:
                files_to_process.append(audio_file)
        
        if not files_to_process:
            # All files were skipped
            return
        
        # Worker function for parallel processing
        def process_file(audio_file, worker_model):
            """Process a single file with the given model instance."""
            if transcription_status.get('cancelled'):
                return None
            
            # Track this file as being processed
            import time
            file_start_time = time.time()
            with status_lock:
                transcription_status['current_files'].append(audio_file.name)
                transcription_status['active_workers'] = len(transcription_status['current_files'])
                if len(transcription_status['current_files']) == 1:
                    transcription_status['current_file'] = audio_file.name
                    transcription_status['current_file_start'] = file_start_time
                else:
                    transcription_status['current_file'] = f"{len(transcription_status['current_files'])} files"
            
            try:
                result = transcribe_file(
                    audio_file,
                    model,
                    language,
                    output_path,
                    word_timestamps,
                    status_callback={'whisper_model': worker_model},
                    transcribe_options=transcribe_opts,
                )
            except Exception as e:
                result = {'error': str(e)[:200]}
            
            file_elapsed_seconds = max(0.0, time.time() - file_start_time)
            audio_seconds = None
            segments_count = None
            speed_x = None
            
            try:
                segments = result.get('segments') if isinstance(result, dict) else None
                if segments:
                    segments_count = len(segments)
                    last_end = segments[-1].get('end')
                    if isinstance(last_end, (int, float)):
                        audio_seconds = float(last_end)
            except Exception:
                pass
            
            if audio_seconds and file_elapsed_seconds > 0:
                speed_x = audio_seconds / file_elapsed_seconds
            
            # Record result
            with status_lock:
                if 'error' in result:
                    transcription_status['results'].append({
                        'file': audio_file.name,
                        'status': 'error',
                        'message': result['error'],
                        'elapsed_seconds': file_elapsed_seconds,
                        'audio_seconds': audio_seconds,
                        'speed_x': speed_x,
                        'segments_count': segments_count,
                    })
                else:
                    md_path = create_markdown(audio_file, result, model, language, output_path, options)
                    transcription_status['results'].append({
                        'file': audio_file.name,
                        'status': 'success',
                        'output': md_path.name,
                        'elapsed_seconds': file_elapsed_seconds,
                        'audio_seconds': audio_seconds,
                        'speed_x': speed_x,
                        'segments_count': segments_count,
                    })
                
                transcription_status['completed'] += 1
                
                # Remove from current files list
                if audio_file.name in transcription_status['current_files']:
                    transcription_status['current_files'].remove(audio_file.name)
                transcription_status['active_workers'] = len(transcription_status['current_files'])
                if transcription_status['current_files']:
                    if len(transcription_status['current_files']) == 1:
                        transcription_status['current_file'] = transcription_status['current_files'][0]
                    else:
                        transcription_status['current_file'] = f"{len(transcription_status['current_files'])} files"
                else:
                    transcription_status['current_file'] = ''
                    transcription_status['current_file_start'] = None
            
            return result
        
        if processing_mode in ('metal', 'cuda') or num_workers == 1:
            # Single-threaded mode with GPU (Metal on macOS, CUDA on Linux)
            transcription_status['current_file_progress'] = f'Loading {model} model...'
            transcription_status['current_file_percent'] = 1
            try:
                import whisper
                import torch
                
                # Determine device based on processing_mode and availability
                # Note: MPS doesn't support float64 required by word timestamp alignment (DTW),
                # so we must fall back to CPU when word timestamps are enabled.
                device = "cpu"
                if processing_mode == 'cuda' and torch.cuda.is_available():
                    device = "cuda"
                    transcription_status['current_file_progress'] = f'Loading {model} model (CUDA GPU)...'
                elif processing_mode == 'metal' and torch.backends.mps.is_available() and not word_timestamps:
                    device = "mps"
                    transcription_status['current_file_progress'] = f'Loading {model} model (Metal GPU)...'
                else:
                    if word_timestamps and processing_mode == 'metal':
                        transcription_status['current_file_progress'] = f'Loading {model} model (CPU - required for word timestamps)...'
                    else:
                        transcription_status['current_file_progress'] = f'Loading {model} model (CPU)...'
                whisper_model = whisper.load_model(model, device=device)
            except Exception as e:
                transcription_status['error'] = f'Failed to load model: {str(e)[:200]}'
                return
            
            for audio_file in files_to_process:
                if transcription_status['cancelled']:
                    transcription_status['error'] = 'Cancelled by user'
                    break
                if transcription_status['cancel_after_file']:
                    transcription_status['error'] = 'Stopped after completing file'
                    break
                process_file(audio_file, whisper_model)
        else:
            # Parallel mode: each worker loads its own model
            from concurrent.futures import ThreadPoolExecutor, as_completed
            import whisper
            
            transcription_status['current_file_progress'] = f'Loading {num_workers} model instances...'
            transcription_status['current_file_percent'] = 1
            
            # Pre-load models for each worker
            worker_models = []
            try:
                for i in range(num_workers):
                    if transcription_status['cancelled']:
                        break
                    transcription_status['current_file_progress'] = f'Loading model {i+1}/{num_workers}...'
                    # Use CPU for parallel workers - MPS has issues with multiple models
                    worker_models.append(whisper.load_model(model, device='cpu'))
            except Exception as e:
                transcription_status['error'] = f'Failed to load models: {str(e)[:200]}'
                return
            
            if transcription_status['cancelled']:
                transcription_status['error'] = 'Cancelled by user'
                return
            
            # Process files in parallel
            with ThreadPoolExecutor(max_workers=num_workers) as executor:
                futures = {}
                model_idx = 0
                
                for audio_file in files_to_process:
                    if transcription_status['cancelled']:
                        break
                    if transcription_status['cancel_after_file']:
                        break
                    # Round-robin assign models to workers
                    worker_model = worker_models[model_idx % len(worker_models)]
                    model_idx += 1
                    future = executor.submit(process_file, audio_file, worker_model)
                    futures[future] = audio_file
                
                # Wait for all to complete (or be cancelled)
                for future in as_completed(futures):
                    if transcription_status['cancelled']:
                        # Don't wait for remaining futures
                        break
                    try:
                        future.result()
                    except Exception:
                        pass
                    # Check cancel_after_file after each file completes
                    if transcription_status['cancel_after_file']:
                        break
            
            if transcription_status['cancelled']:
                transcription_status['error'] = 'Cancelled by user'
            elif transcription_status['cancel_after_file']:
                transcription_status['error'] = 'Stopped after completing file(s)'
        
    except Exception as e:
        transcription_status['error'] = str(e)
    finally:
        transcription_status['running'] = False
        transcription_status['active_jobs'] = 0
        transcription_status['active_workers'] = 0
        transcription_status['current_file'] = ''
        transcription_status['current_files'] = []


@app.route('/')
def index():
    return render_template('index.html', 
                         models=AVAILABLE_MODELS, 
                         languages=COMMON_LANGUAGES,
                         app_mode=session_store.get_app_mode(),
                         is_server_mode=session_store.is_server_mode())


@app.route('/admin')
def admin_page():
    """Admin settings page for remote worker configuration."""
    return render_template('admin.html')


@app.route('/start', methods=['POST'])
def start_transcription():
    """Start transcription with folder paths. Blocked in server mode."""
    global transcription_status
    
    if session_store.is_server_mode():
        return jsonify({'error': 'Folder-based transcription not available in server mode. Use /api/jobs.'}), 400
    
    if transcription_status['running']:
        return jsonify({'error': 'Transcription already in progress'}), 400
    
    data = request.json
    input_folder = data.get('input_folder', '').strip()
    output_folder = data.get('output_folder', '').strip()
    model = data.get('model', 'base')
    language = data.get('language', '')
    
    options = {
        'include_segments': data.get('include_segments', True),
        'include_full': data.get('include_full', True),
        'include_timestamps': data.get('include_timestamps', True),
        'word_timestamps': data.get('word_timestamps', False),
        'overwrite_existing': data.get('overwrite_existing', False),
        'processing_mode': data.get('processing_mode', 'metal'),
        'workers': data.get('workers', 1),
        'quality_preset': data.get('quality_preset', 'balanced'),
        'no_speech_threshold': data.get('no_speech_threshold'),
        'max_segment_duration': data.get('max_segment_duration'),
        'vad_enabled': data.get('vad_enabled', False),
        'hallucination_detection': data.get('hallucination_detection', True),
        'hallucination_sensitivity': data.get('hallucination_sensitivity', 50),
    }
    
    if not input_folder:
        return jsonify({'error': 'Input folder is required'}), 400
    if not output_folder:
        return jsonify({'error': 'Output folder is required'}), 400
    if not Path(input_folder).exists():
        return jsonify({'error': 'Input folder does not exist'}), 400
    
    # Start transcription in background thread
    thread = threading.Thread(
        target=run_transcription,
        args=(input_folder, output_folder, model, language, options)
    )
    thread.daemon = True
    thread.start()
    
    return jsonify({'status': 'started'})


@app.route('/status')
def get_status():
    import time as time_module
    # Update elapsed time if running
    if transcription_status['running'] and transcription_status.get('start_time'):
        transcription_status['elapsed_seconds'] = time_module.time() - transcription_status['start_time']
    # Update current file elapsed time
    if transcription_status.get('current_file_start'):
        transcription_status['current_file_elapsed'] = time_module.time() - transcription_status['current_file_start']
    else:
        transcription_status['current_file_elapsed'] = 0
    return jsonify(transcription_status)


@app.route('/cancel', methods=['POST'])
def cancel_transcription():
    """Cancel the current transcription job.
    
    Modes:
    - 'force': Stop immediately (current file may be incomplete)
    - 'after_file': Complete current file(s), then stop
    """
    global transcription_status
    
    if not transcription_status['running']:
        return jsonify({'error': 'No transcription in progress'}), 400
    
    data = request.get_json() or {}
    mode = data.get('mode', 'force')
    
    if mode == 'after_file':
        transcription_status['cancel_after_file'] = True
        return jsonify({'status': 'cancelling_after_file'})
    else:
        # Force stop - set cancelled and immediately mark as not running
        # The background thread will clean up when it checks the flag
        transcription_status['cancelled'] = True
        transcription_status['running'] = False
        transcription_status['error'] = 'Force stopped by user'
        transcription_status['current_file'] = ''
        transcription_status['current_file_progress'] = ''
        transcription_status['current_file_percent'] = 0
        return jsonify({'status': 'stopped'})


@app.route('/healthz')
def healthz():
    """Health check endpoint for Docker/orchestration."""
    return jsonify({'ok': True})


@app.route('/preview')
def preview_files():
    """Preview files in a folder. Blocked in server mode."""
    if session_store.is_server_mode():
        return jsonify({'error': 'Folder preview not available in server mode. Use /api/uploads.'}), 404
    
    input_folder = request.args.get('folder', '')
    if not input_folder or not Path(input_folder).exists():
        return jsonify({'files': [], 'error': 'Folder not found'})
    
    files = find_audio_files(input_folder)
    return jsonify({
        'files': [f.name for f in files],
        'count': len(files)
    })


@app.route('/browse', methods=['POST'])
def browse_folder():
    """Open native folder picker. Blocked in server mode."""
    if session_store.is_server_mode():
        return jsonify({'error': 'Folder browser not available in server mode.'}), 404
    
    data = request.json or {}
    prompt = data.get('prompt', 'Select a folder')
    folder = choose_folder_dialog(prompt)
    if folder:
        return jsonify({'folder': folder})
    return jsonify({'folder': None, 'cancelled': True})


def get_recommended_model():
    """Get the recommended model to use (user default or best downloaded)."""
    prefs = load_preferences()
    
    # If user has set a default and it's downloaded, use it
    if prefs.get('default_model'):
        if check_model_available(prefs['default_model']):
            return prefs['default_model']
    
    # Otherwise, prefer large-v3 if downloaded (best quality)
    if check_model_available('large-v3'):
        return 'large-v3'
    
    # Fallback to any downloaded model
    priority_order = ['large-v3', 'large', 'medium', 'small', 'base', 'tiny']
    for model in priority_order:
        if check_model_available(model):
            return model
    
    # Fallback to large-v3 if nothing downloaded (will download on first use)
    return 'large-v3'


@app.route('/models')
def list_models():
    """Get list of all models with their status."""
    prefs = load_preferences()
    recommended = get_recommended_model()
    return jsonify({
        'models': get_all_models_status(),
        'recommended': recommended,
        'user_default': prefs.get('default_model')
    })


@app.route('/preferences')
def get_preferences():
    """Get user preferences."""
    return jsonify(load_preferences())


@app.route('/preferences', methods=['POST'])
def set_preferences():
    """Save user preferences."""
    data = request.json or {}
    prefs = load_preferences()
    
    # Update only provided fields
    for key in ['default_model', 'default_language', 'include_segments', 
                'include_full', 'include_timestamps', 'word_timestamps']:
        if key in data:
            prefs[key] = data[key]
    
    if save_preferences(prefs):
        return jsonify({'status': 'saved', 'preferences': prefs})
    return jsonify({'error': 'Failed to save preferences'}), 500


@app.route('/models/set-default', methods=['POST'])
def set_default_model():
    """Set the default model."""
    data = request.json or {}
    model_name = data.get('model', '')
    
    if model_name and model_name not in AVAILABLE_MODELS:
        return jsonify({'error': f'Invalid model: {model_name}'}), 400
    
    prefs = load_preferences()
    prefs['default_model'] = model_name if model_name else None
    
    if save_preferences(prefs):
        return jsonify({'status': 'saved', 'default_model': prefs['default_model']})
    return jsonify({'error': 'Failed to save preference'}), 500


model_download_status = {
    'downloading': False,
    'model': '',
    'progress': ''
}


@app.route('/models/download', methods=['POST'])
def download_model():
    """Download a whisper model."""
    global model_download_status
    
    if model_download_status['downloading']:
        return jsonify({'error': 'A download is already in progress'}), 400
    
    data = request.json or {}
    model_name = data.get('model', '')
    
    if model_name not in AVAILABLE_MODELS:
        return jsonify({'error': f'Invalid model: {model_name}'}), 400
    
    if check_model_available(model_name):
        return jsonify({'error': f'Model {model_name} is already downloaded'}), 400
    
    def do_download():
        global model_download_status
        model_download_status['downloading'] = True
        model_download_status['model'] = model_name
        model_download_status['progress'] = 'Initializing download...'
        
        model_sizes = {
            'tiny': '75 MB',
            'base': '145 MB', 
            'small': '465 MB',
            'medium': '1.5 GB',
            'large': '3 GB',
            'large-v3': '3 GB'
        }
        size_str = model_sizes.get(model_name, 'Unknown size')
        
        try:
            import time
            start_time = time.time()
            
            model_download_status['progress'] = f'Connecting to download server...'
            time.sleep(0.5)
            
            model_download_status['progress'] = f'Downloading {model_name} model ({size_str})... This may take several minutes.'
            
            # Monitor the cache directory for file size changes
            model_file = get_model_filepath(model_name)
            
            def update_progress():
                while model_download_status['downloading']:
                    elapsed = time.time() - start_time
                    elapsed_str = f"{int(elapsed // 60)}m {int(elapsed % 60)}s" if elapsed >= 60 else f"{int(elapsed)}s"
                    
                    if os.path.exists(model_file):
                        current_size = os.path.getsize(model_file)
                        if current_size >= 1024 * 1024 * 1024:
                            size_downloaded = f"{current_size / (1024 * 1024 * 1024):.2f} GB"
                        else:
                            size_downloaded = f"{current_size / (1024 * 1024):.0f} MB"
                        model_download_status['progress'] = f'Downloading {model_name}: {size_downloaded} of {size_str} ({elapsed_str} elapsed)'
                    else:
                        model_download_status['progress'] = f'Downloading {model_name} model ({size_str})... {elapsed_str} elapsed'
                    time.sleep(1)
            
            # Start progress monitor in separate thread
            progress_thread = threading.Thread(target=update_progress, daemon=True)
            progress_thread.start()
            
            # Do the actual download
            import whisper
            whisper.load_model(model_name)
            
            elapsed = time.time() - start_time
            elapsed_str = f"{int(elapsed // 60)}m {int(elapsed % 60)}s" if elapsed >= 60 else f"{int(elapsed)}s"
            model_download_status['progress'] = f'Complete! Downloaded {model_name} in {elapsed_str}'
            
        except Exception as e:
            model_download_status['progress'] = f'Error: {str(e)[:100]}'
        finally:
            model_download_status['downloading'] = False
    
    thread = threading.Thread(target=do_download, daemon=True)
    thread.start()
    
    return jsonify({'status': 'downloading', 'model': model_name})


@app.route('/models/download/status')
def download_status():
    """Get the current model download status."""
    return jsonify(model_download_status)


@app.route('/models/delete', methods=['POST'])
def delete_model():
    """Delete a downloaded whisper model."""
    data = request.json or {}
    model_name = data.get('model', '')
    
    if model_name not in AVAILABLE_MODELS:
        return jsonify({'error': f'Invalid model: {model_name}'}), 400
    
    model_file = get_model_filepath(model_name)
    
    if not os.path.exists(model_file):
        return jsonify({'error': f'Model {model_name} is not downloaded'}), 400
    
    try:
        os.remove(model_file)
        return jsonify({'status': 'deleted', 'model': model_name})
    except Exception as e:
        return jsonify({'error': f'Failed to delete: {str(e)}'}), 500


@app.route('/upload', methods=['POST'])
def upload_files_legacy():
    """Upload files to a folder path. Blocked in server mode."""
    if session_store.is_server_mode():
        return jsonify({'error': 'Folder-based upload not available in server mode. Use /api/uploads.'}), 400
    
    if 'files' not in request.files:
        return jsonify({'error': 'No files provided'}), 400
    
    files = request.files.getlist('files')
    if not files or all(f.filename == '' for f in files):
        return jsonify({'error': 'No files selected'}), 400
    
    # Get target folder from form data or use default
    target_folder = request.form.get('folder', os.environ.get('INPUT_DIR', '/data/input'))
    target_path = Path(target_folder)
    
    # Create folder if it doesn't exist
    target_path.mkdir(parents=True, exist_ok=True)
    
    uploaded = []
    skipped = []
    errors = []
    
    for file in files:
        if file.filename == '':
            continue
            
        filename = secure_filename(file.filename)
        if not filename:
            continue
            
        # Check if it's a supported audio format
        ext = Path(filename).suffix.lower()
        if ext not in SUPPORTED_FORMATS:
            skipped.append({'name': filename, 'reason': f'Unsupported format: {ext}'})
            continue
        
        try:
            filepath = target_path / filename
            file.save(str(filepath))
            uploaded.append(filename)
        except Exception as e:
            errors.append({'name': filename, 'error': str(e)[:100]})
    
    return jsonify({
        'uploaded': uploaded,
        'skipped': skipped,
        'errors': errors,
        'folder': str(target_path),
        'total_uploaded': len(uploaded)
    })


@app.route('/download', methods=['GET'])
def download_transcriptions_legacy():
    """Download from folder path. Blocked in server mode."""
    if session_store.is_server_mode():
        return jsonify({'error': 'Folder-based download not available in server mode. Use /api/jobs/<job_id>/download.'}), 404
    
    output_folder = request.args.get('folder', os.environ.get('OUTPUT_DIR', '/data/output'))
    output_path = Path(output_folder)
    
    if not output_path.exists():
        return jsonify({'error': 'Output folder does not exist'}), 404
    
    # Find all transcription files (markdown and json)
    files_to_zip = []
    for ext in ['*.md', '*.json', '*.txt', '*.srt']:
        files_to_zip.extend(output_path.glob(ext))
    
    if not files_to_zip:
        return jsonify({'error': 'No transcription files found'}), 404
    
    # Create zip in memory
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for filepath in files_to_zip:
            zf.write(filepath, filepath.name)
    
    zip_buffer.seek(0)
    
    # Generate filename with timestamp
    from datetime import datetime
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    zip_filename = f'transcriptions_{timestamp}.zip'
    
    return send_file(
        zip_buffer,
        mimetype='application/zip',
        as_attachment=True,
        download_name=zip_filename
    )


@app.route('/download/list', methods=['GET'])
def list_transcriptions_legacy():
    """List files in folder. Blocked in server mode."""
    if session_store.is_server_mode():
        return jsonify({'error': 'Folder-based listing not available in server mode. Use /api/jobs/<job_id>/outputs.'}), 404
    
    output_folder = request.args.get('folder', os.environ.get('OUTPUT_DIR', '/data/output'))
    output_path = Path(output_folder)
    
    if not output_path.exists():
        return jsonify({'files': [], 'folder': str(output_path), 'exists': False})
    
    files = []
    for ext in ['*.md', '*.json', '*.txt', '*.srt']:
        for filepath in output_path.glob(ext):
            stat = filepath.stat()
            files.append({
                'name': filepath.name,
                'size': stat.st_size,
                'modified': stat.st_mtime
            })
    
    # Sort by modification time, newest first
    files.sort(key=lambda x: x['modified'], reverse=True)
    
    return jsonify({
        'files': files,
        'folder': str(output_path),
        'exists': True,
        'count': len(files)
    })


# =============================================================================
# Session-Isolated API Endpoints (Server Mode)
# =============================================================================

# In-memory job runners keyed by (session_id, job_id) for cancellation
_active_jobs = {}
_active_jobs_lock = threading.Lock()


@app.route('/api/uploads', methods=['POST'])
def api_upload_files():
    """Upload audio files to session-isolated storage."""
    if 'files' not in request.files:
        return jsonify({'error': 'No files provided'}), 400
    
    files = request.files.getlist('files')
    if not files or all(f.filename == '' for f in files):
        return jsonify({'error': 'No files selected'}), 400
    
    # Enforce file count limit
    max_files = session_store.get_max_files_per_job()
    if len(files) > max_files:
        return jsonify({'error': f'Too many files. Maximum {max_files} per upload.'}), 400
    
    session_id = g.session_id
    session_store.ensure_session_dirs(session_id)
    
    # Check session disk cap (server mode only)
    if session_store.is_server_mode():
        current_usage = session_store.get_session_disk_usage_mb(session_id)
        max_session_mb = session_store.get_max_session_mb()
        if current_usage >= max_session_mb:
            return jsonify({
                'error': f'Session storage limit reached ({max_session_mb}MB). Delete old jobs or wait for cleanup.'
            }), 400
    
    uploads_path = Path(session_store.uploads_dir(session_id))
    
    uploaded = []
    skipped = []
    errors = []
    
    for file in files:
        if file.filename == '':
            continue
        
        original_name = file.filename
        sanitized = session_store.sanitize_filename(original_name)
        ext = Path(sanitized).suffix.lower()
        
        # Check supported format
        if ext not in SUPPORTED_FORMATS:
            skipped.append({'name': original_name, 'reason': f'Unsupported format: {ext}'})
            continue
        
        try:
            upload_id = session_store.new_id(8)
            stored_name = f"{upload_id}_{sanitized}"
            filepath = uploads_path / stored_name
            file.save(str(filepath))
            
            uploaded.append({
                'id': upload_id,
                'filename': original_name,
                'storedName': stored_name,
                'size': filepath.stat().st_size
            })
        except Exception as e:
            errors.append({'name': original_name, 'error': str(e)[:100]})
    
    return jsonify({
        'uploads': uploaded,
        'skipped': skipped,
        'errors': errors
    })


@app.route('/api/uploads', methods=['GET'])
def api_list_uploads():
    """List uploaded files for current session."""
    session_id = g.session_id
    uploads = session_store.list_uploads(session_id)
    return jsonify({'uploads': uploads})


@app.route('/api/jobs', methods=['POST'])
def api_create_job():
    """Create and start a transcription job."""
    session_id = g.session_id
    
    # Check for existing running job in this session
    with _active_jobs_lock:
        for key, job_info in _active_jobs.items():
            if key[0] == session_id and job_info.get('running'):
                return jsonify({'error': 'A job is already running in this session'}), 409
    
    data = request.json or {}
    upload_ids = data.get('uploadIds', [])
    options = data.get('options', {})
    
    if not upload_ids:
        return jsonify({'error': 'No upload IDs provided'}), 400
    
    # Validate model is in allowed list (no aliasing, no "turbo")
    ALLOWED_MODELS = {'tiny', 'base', 'small', 'medium', 'large', 'large-v3'}
    requested_model = options.get('model', 'large-v3')
    if requested_model not in ALLOWED_MODELS:
        return jsonify({
            'error': f"Unsupported model '{requested_model}'. Allowed models: {', '.join(sorted(ALLOWED_MODELS))}",
            'code': 'UNSUPPORTED_MODEL'
        }), 400
    
    # Validate and determine backend
    env = compute_backend.get_cached_environment()
    requested_backend = options.get('backend', env['recommendedBackend'])
    
    is_valid, error = compute_backend.validate_backend(requested_backend)
    if not is_valid:
        return jsonify({'error': error}), 400
    
    # Validate diarization requirements if enabled
    diarization_enabled = options.get('diarizationEnabled', False)
    if diarization_enabled:
        hf_token = os.environ.get('HF_TOKEN') or os.environ.get('HUGGINGFACE_TOKEN')
        if not hf_token:
            return jsonify({
                'error': 'Diarization requires HF_TOKEN environment variable. Get a token from https://huggingface.co/settings/tokens',
                'code': 'HF_TOKEN_MISSING'
            }), 400
        
        try:
            import pyannote.audio
        except ImportError:
            return jsonify({
                'error': 'Diarization requires pyannote.audio which is not installed on this server.',
                'code': 'DIARIZATION_UNAVAILABLE'
            }), 400
    
    # Get diarization duration limit from env
    max_diarization_duration = int(os.environ.get('MAX_DIARIZATION_DURATION_SECONDS', '180'))
    
    # Resolve upload IDs to file paths and probe audio metadata
    inputs = []
    for upload_id in upload_ids:
        filepath = session_store.find_upload_by_id(session_id, upload_id)
        if not filepath:
            return jsonify({'error': f'Upload not found: {upload_id}'}), 404
        
        # Extract original filename from stored name
        stored_name = os.path.basename(filepath)
        parts = stored_name.split('_', 1)
        original_name = parts[1] if len(parts) > 1 else stored_name
        
        # Probe audio metadata
        audio_info = None
        try:
            from audio_utils import get_audio_info, AudioProbeError
            audio_info = get_audio_info(filepath)
        except AudioProbeError as e:
            return jsonify({
                'error': f'Failed to probe audio file: {original_name}',
                'code': 'AUDIO_PROBE_FAILED',
                'details': str(e)
            }), 400
        except Exception as e:
            return jsonify({
                'error': f'Unexpected error probing audio: {original_name}',
                'code': 'AUDIO_PROBE_FAILED',
                'details': str(e)[:200]
            }), 400
        
        inputs.append({
            'uploadId': upload_id,
            'path': filepath,
            'originalFilename': original_name,
            'durationSec': audio_info.get('durationSec') if audio_info else None,
            'codec': audio_info.get('codec') if audio_info else None,
            'sampleRate': audio_info.get('sampleRate') if audio_info else None,
            'channels': audio_info.get('channels') if audio_info else None
        })
    
    # Compute effective diarization policy using single source of truth
    diarization_effective = None
    diarization_warnings = []
    if diarization_enabled:
        from diarization_policy import compute_diarization_policy, get_server_policy_config, get_clamping_warnings
        
        server_config = get_server_policy_config()
        diarization_effective = compute_diarization_policy(
            diarization_enabled=diarization_enabled,
            diarization_auto_split=options.get('diarizationAutoSplit', False),
            diarization_fast_switching=options.get('diarizationFastSwitching', False),
            requested_max_duration_seconds=options.get('diarizationMaxDurationSeconds'),
            requested_chunk_seconds=options.get('diarizationChunkSeconds'),
            requested_overlap_seconds=options.get('diarizationOverlapSeconds'),
            server_max_duration_seconds=server_config['serverMaxDurationSeconds'],
            default_max_duration_seconds=server_config['defaultMaxDurationSeconds'],
            min_chunk_seconds=server_config['minChunkSeconds'],
            max_chunk_seconds=server_config['maxChunkSeconds'],
            overlap_ratio=server_config['overlapRatio'],
            min_overlap_seconds=server_config['minOverlapSeconds'],
            max_overlap_seconds=server_config['maxOverlapSeconds'],
        )
        diarization_warnings = get_clamping_warnings(diarization_effective)
    
    # Validate diarization duration limits on CPU (unless auto-split is enabled)
    if diarization_enabled and diarization_effective and requested_backend == 'cpu':
        effective_max = diarization_effective['maxDurationSeconds']
        auto_split_enabled = diarization_effective['autoSplit']
        
        if not auto_split_enabled:
            too_long_files = []
            for inp in inputs:
                duration = inp.get('durationSec')
                if duration and duration > effective_max:
                    from audio_utils import format_duration
                    too_long_files.append({
                        'filename': inp['originalFilename'],
                        'durationSec': duration,
                        'durationFormatted': format_duration(duration)
                    })
            
            if too_long_files:
                from audio_utils import format_duration
                from diarization_policy import get_clamping_warnings
                server_max = diarization_effective.get('serverMaxDurationSeconds', effective_max)
                clamped_info = diarization_effective.get('clamped', {})
                was_clamped = clamped_info.get('maxDurationClamped', False)
                requested_max = clamped_info.get('maxDurationOriginal')
                
                error_response = {
                    'error': f'Audio file(s) too long for CPU diarization. Max duration: {format_duration(effective_max)}',
                    'code': 'DIARIZATION_TOO_LONG_CPU',
                    'effectivePolicy': {
                        'maxDurationSeconds': effective_max,
                        'maxDurationFormatted': format_duration(effective_max),
                        'autoSplit': False,
                        'chunkSeconds': diarization_effective.get('chunkSeconds'),
                        'overlapSeconds': diarization_effective.get('overlapSeconds'),
                    },
                    'serverMaxDurationSec': server_max,
                    'serverMaxDurationFormatted': format_duration(server_max),
                    'tooLongFiles': too_long_files,
                    'suggestions': [
                        'Enable "Auto-split long audio" to process in chunks',
                        f'Increase max duration (server allows up to {format_duration(server_max)})',
                        'Split audio into shorter segments',
                        'Disable diarization for this job',
                        'Run on GPU backend if available'
                    ]
                }
                
                if was_clamped and requested_max:
                    error_response['requestedMaxDurationSec'] = requested_max
                    error_response['requestedMaxDurationFormatted'] = format_duration(requested_max)
                    error_response['wasClamped'] = True
                    error_response['clampingWarnings'] = get_clamping_warnings(diarization_effective)
                
                return jsonify(error_response), 400
    
    # Check if remote worker should be used
    use_remote_worker = options.get('useRemoteWorker', False)
    execution_mode = 'local'
    remote_fallback_reason = None
    
    try:
        from remote_worker import get_worker_config, get_remote_worker_status
        worker_config = get_worker_config()
        
        if worker_config['mode'] == 'required':
            # Mode is required - check if worker is actually reachable
            worker_status = get_remote_worker_status(force_refresh=True)
            if not worker_status['connected']:
                error_msg = worker_status.get('error') or 'Worker not reachable'
                return jsonify({
                    'error': f'Remote worker is required but offline: {error_msg}',
                    'code': 'REMOTE_WORKER_REQUIRED_OFFLINE',
                    'workerStatus': worker_status
                }), 503
            use_remote_worker = True
            execution_mode = 'remote'
        elif worker_config['mode'] == 'optional' and use_remote_worker:
            # User requested remote - check if available
            worker_status = get_remote_worker_status()
            if not worker_status['connected']:
                # Fall back to local - track reason for UI
                execution_mode = 'local'
                use_remote_worker = False
                remote_fallback_reason = worker_status.get('error') or 'Worker unreachable'
            else:
                execution_mode = 'remote'
    except ImportError:
        pass
    
    # Track if user requested remote but we fell back
    remote_requested = options.get('useRemoteWorker', False)
    
    # Create job
    job_id = session_store.new_id(12)
    dirs = session_store.ensure_job_dirs(session_id, job_id)
    
    now = datetime.now(timezone.utc).isoformat()
    manifest = {
        'jobId': job_id,
        'sessionId': session_id,
        'createdAt': now,
        'updatedAt': now,
        'startedAt': None,
        'finishedAt': None,
        'status': 'queued',
        'backend': requested_backend,
        'executionMode': execution_mode,
        'environment': {
            'os': env['os'],
            'arch': env['arch'],
            'inDocker': env['inDocker']
        },
        'options': {
            'model': options.get('model', 'base'),
            'language': options.get('language', ''),
            'includeSegments': options.get('includeSegments', True),
            'includeFull': options.get('includeFull', True),
            'includeTimestamps': options.get('includeTimestamps', True),
            'wordTimestamps': options.get('wordTimestamps', False),
            'qualityPreset': options.get('qualityPreset', 'balanced'),
            'diarizationEnabled': diarization_enabled,
            'minSpeakers': options.get('minSpeakers') if diarization_enabled else None,
            'maxSpeakers': options.get('maxSpeakers') if diarization_enabled else None,
            'numSpeakers': options.get('numSpeakers') if diarization_enabled else None,
            'diarizationMaxDurationSeconds': options.get('diarizationMaxDurationSeconds') if diarization_enabled else None,
            'diarizationAutoSplit': options.get('diarizationAutoSplit', False) if diarization_enabled else False,
            'diarizationFastSwitching': options.get('diarizationFastSwitching', False) if diarization_enabled else False,
            'diarizationChunkSeconds': options.get('diarizationChunkSeconds') if diarization_enabled else None,
            'diarizationOverlapSeconds': options.get('diarizationOverlapSeconds') if diarization_enabled else None,
            'diarizationEffective': diarization_effective,
            'diarizationWarnings': diarization_warnings,
            'vadEnabled': options.get('vadEnabled', False),
            'hallucinationDetection': options.get('hallucinationDetection', True),
            'hallucinationSensitivity': options.get('hallucinationSensitivity', 50),
            'noSpeechThreshold': options.get('noSpeechThreshold', 0.6),
        },
        'inputs': inputs,
        'outputs': [],
        'progress': {
            'totalFiles': len(inputs),
            'currentFileIndex': 0,
            'currentFile': '',
            'percent': 0,
            'stage': 'queued',
            'stageStartedAt': now
        },
        'error': None,
        'worker': None,  # Will be populated if using remote worker
        'remoteRequested': remote_requested,
        'remoteFallbackReason': remote_fallback_reason if (remote_requested and execution_mode == 'local') else None
    }
    
    session_store.atomic_write_json(session_store.job_manifest_path(session_id, job_id), manifest)
    
    with _active_jobs_lock:
        _active_jobs[(session_id, job_id)] = {'running': True, 'cancel_requested': False}
    
    # Start job in background thread
    if execution_mode == 'remote':
        def run_remote_job():
            _run_remote_session_job(session_id, job_id, inputs, manifest['options'], dirs['outputs'])
        job_thread = threading.Thread(target=run_remote_job, daemon=True)
    else:
        def run_job():
            _run_session_job(session_id, job_id, inputs, manifest['options'], dirs['outputs'], requested_backend)
        job_thread = threading.Thread(target=run_job, daemon=True)
    
    job_thread.start()
    
    return jsonify({'jobId': job_id, 'executionMode': execution_mode})


def _run_session_job(session_id: str, job_id: str, inputs: list, options: dict, outputs_dir: str, backend: str):
    """Run transcription job and update manifest with progress."""
    from logger import log_event, with_timer
    
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    
    # Log job start
    log_event('info', 'job_started', 
             jobId=job_id, 
             sessionId=session_id,
             backend=backend,
             numFiles=len(inputs),
             options=options)
    
    def update_manifest(**updates):
        manifest = session_store.read_json(manifest_path) or {}
        now = datetime.now(timezone.utc).isoformat()
        manifest['updatedAt'] = now
        for key, value in updates.items():
            if key == 'progress':
                # Check if stage is changing
                new_stage = value.get('stage')
                old_stage = manifest.get('progress', {}).get('stage')
                if new_stage and new_stage != old_stage:
                    value['stageStartedAt'] = now
                manifest.setdefault('progress', {}).update(value)
            elif key == 'debug':
                manifest.setdefault('debug', {}).update(value)
            else:
                manifest[key] = value
        session_store.atomic_write_json(manifest_path, manifest)
        return manifest
    
    def is_cancelled():
        with _active_jobs_lock:
            job_info = _active_jobs.get((session_id, job_id), {})
            return job_info.get('cancel_requested', False)
    
    def is_cancel_after_current():
        with _active_jobs_lock:
            job_info = _active_jobs.get((session_id, job_id), {})
            return job_info.get('cancel_after_current', False)
    
    def write_error_artifact(code: str, message: str, phase: str, suggestion: str = '', outputs: list = None):
        """Write a job_error.md artifact for reviewer-friendly error reporting."""
        error_filename = 'job_error.md'
        error_path = os.path.join(outputs_dir, error_filename)
        
        # Get current manifest for effective policy
        current_manifest = session_store.read_json(manifest_path) or {}
        effective_policy = current_manifest.get('options', {}).get('diarizationEffective')
        current_stage = current_manifest.get('progress', {}).get('stage', 'unknown')
        
        content = f"""# Job Error Report

**Job ID:** {job_id}
**Status:** failed
**Backend:** {backend}
**Phase:** {phase}
**Stage at failure:** {current_stage}
**Timestamp:** {datetime.now(timezone.utc).isoformat()}

## Error

**Code:** `{code}`

**Message:** {message}

"""
        if suggestion:
            content += f"""## Suggestion

{suggestion}

"""
        
        if effective_policy:
            content += f"""## Effective Diarization Policy

- **Max Duration:** {effective_policy.get('maxDurationSeconds')}s
- **Auto-Split:** {effective_policy.get('autoSplit')}
- **Chunk Size:** {effective_policy.get('chunkSeconds')}s
- **Overlap:** {effective_policy.get('overlapSeconds')}s
- **Server Max:** {effective_policy.get('serverMaxDurationSeconds')}s

"""
        
        try:
            with open(error_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            error_output = {
                'id': session_store.new_id(8),
                'filename': error_filename,
                'path': error_path,
                'type': 'error-report',
                'sizeBytes': os.path.getsize(error_path)
            }
            
            if outputs is not None:
                outputs.append(error_output)
            return error_output
        except Exception:
            return None
    
    def fail_job(code: str, message: str, phase: str, suggestion: str = '', outputs: list = None):
        """Terminate job with error, write artifact, and update manifest."""
        error_output = write_error_artifact(code, message, phase, suggestion, outputs)
        final_outputs = outputs if outputs else []
        if error_output and error_output not in final_outputs:
            final_outputs.append(error_output)
        
        update_manifest(
            status='failed',
            finishedAt=datetime.now(timezone.utc).isoformat(),
            error={'code': code, 'message': message},
            outputs=final_outputs
        )
    
    def cancel_job(outputs: list = None):
        """Terminate job as canceled, preserving outputs."""
        update_manifest(
            status='canceled',
            finishedAt=datetime.now(timezone.utc).isoformat(),
            error={'code': 'USER_CANCELED', 'message': 'Job canceled by user'},
            outputs=outputs or []
        )
    
    try:
        update_manifest(status='running', startedAt=datetime.now(timezone.utc).isoformat())
        
        # Load whisper model
        model_name = options.get('model', 'base')
        language = options.get('language', '') or None
        
        update_manifest(progress={
            'currentFile': f'Loading {model_name} model on {backend}...',
            'stage': 'loading_model'
        })
        
        # Use explicit backend - NO auto-detection
        device = compute_backend.get_torch_device(backend)
        
        # Validate backend is actually available at runtime
        try:
            import torch
            import whisper
        except ImportError as e:
            fail_job('DEPENDENCY_ERROR', f'Required library not available: {e}',
                     phase='initialization',
                     suggestion='Check that torch and openai-whisper are installed.')
            return
        
        try:
            if device == 'cuda' and not torch.cuda.is_available():
                raise RuntimeError('CUDA is not available on this system')
            if device == 'mps':
                if not hasattr(torch.backends, 'mps') or not torch.backends.mps.is_available():
                    raise RuntimeError('Metal (MPS) is not available on this system')
                if options.get('wordTimestamps', False):
                    raise RuntimeError('Metal backend does not support word timestamps. Use CPU instead.')
        except RuntimeError as e:
            fail_job('BACKEND_RUNTIME_FAILURE', str(e),
                     phase='backend validation',
                     suggestion='Try using CPU backend instead.')
            return
        
        try:
            whisper_model = whisper.load_model(model_name, device=device)
        except torch.cuda.OutOfMemoryError:
            fail_job('OUT_OF_MEMORY', f'Not enough GPU memory to load {model_name} model.',
                     phase='model loading',
                     suggestion='Try a smaller model (e.g., base or small) or use CPU backend.')
            return
        except Exception as e:
            fail_job('MODEL_LOAD_ERROR', f'Failed to load {model_name} model: {str(e)[:200]}',
                     phase='model loading',
                     suggestion='Check model name is valid and sufficient memory is available.')
            return
        
        outputs = []
        total = len(inputs)
        
        for idx, input_info in enumerate(inputs):
            # Check for immediate cancel
            if is_cancelled():
                cancel_job(outputs)
                return
            
            # Check for graceful cancel after previous file
            if idx > 0 and is_cancel_after_current():
                update_manifest(
                    status='canceled',
                    finishedAt=datetime.now(timezone.utc).isoformat(),
                    outputs=outputs,
                    progress={'currentFile': 'Canceled after file ' + str(idx), 'percent': int((idx / total) * 100)}
                )
                return
            
            filepath = input_info['path']
            original_name = input_info['originalFilename']
            upload_id = input_info.get('uploadId', '')
            base_name = Path(original_name).stem
            
            update_manifest(progress={
                'currentFile': original_name,
                'currentFileIndex': idx,
                'percent': int((idx / total) * 100),
                'stage': 'transcribing'
            })
            
            try:
                # Transcribe
                result = whisper.transcribe(
                    whisper_model,
                    filepath,
                    language=language,
                    word_timestamps=options.get('wordTimestamps', False)
                )
                
                # Check for cancellation between transcription and diarization
                if is_cancelled():
                    cancel_job(outputs)
                    return
                
                # Run diarization if enabled
                diarization_enabled = options.get('diarizationEnabled', False)
                speaker_segments = None
                merged_segments = None
                
                if diarization_enabled:
                    update_manifest(progress={
                        'currentFile': f'{original_name} (loading diarization pipeline...)',
                        'currentFileIndex': idx,
                        'percent': int((idx / total) * 100),
                        'stage': 'diarizing'
                    })
                    
                    try:
                        import diarization
                        import time
                        from threading import Thread
                        from queue import Queue, Empty
                        from audio_utils import get_audio_info, split_audio_to_wav_chunks, cleanup_chunks
                        
                        # Get effective diarization policy from manifest (single source of truth)
                        effective_policy = options.get('diarizationEffective') or {}
                        auto_split = effective_policy.get('autoSplit', False)
                        chunk_seconds = effective_policy.get('chunkSeconds', 150)
                        overlap_seconds = effective_policy.get('overlapSeconds', 5)
                        default_max = int(os.environ.get('DIARIZATION_DEFAULT_MAX_DURATION_SECONDS', '180'))
                        effective_max_duration = effective_policy.get('maxDurationSeconds', default_max)
                        
                        # Get file duration to decide if chunking is needed
                        file_duration = None
                        try:
                            audio_info = get_audio_info(filepath)
                            file_duration = audio_info.get('durationSec', 0)
                        except Exception:
                            pass
                        
                        # Determine if we should use chunked diarization
                        use_chunked = auto_split and file_duration and file_duration > effective_max_duration
                        
                        # Log effective policy once per file
                        log_event('diarization_policy_effective', {
                            'filename': original_name,
                            'fileDurationSec': file_duration,
                            'maxDurationSeconds': effective_max_duration,
                            'autoSplit': auto_split,
                            'chunkSeconds': chunk_seconds,
                            'overlapSeconds': overlap_seconds,
                            'useChunked': use_chunked,
                            'derived': effective_policy.get('derived', True),
                        }, stage='diarization')
                        
                        # Add watchdog timeout for diarization
                        max_diarization_minutes = int(os.environ.get('MAX_DIARIZATION_MINUTES', '30'))
                        # Increase timeout for chunked diarization
                        if use_chunked:
                            estimated_chunks = int(file_duration / (chunk_seconds - overlap_seconds)) + 1
                            max_diarization_minutes = max(max_diarization_minutes, estimated_chunks * 5)
                        
                        diarization_result = Queue()
                        diarization_error = Queue()
                        
                        def run_diarization_with_timeout():
                            try:
                                # Check for cancellation before starting diarization
                                if is_cancelled():
                                    return
                                
                                if use_chunked:
                                    # Chunked diarization for long files
                                    update_manifest(progress={
                                        'currentFile': f'{original_name} (splitting audio for diarization...)',
                                        'currentFileIndex': idx,
                                        'percent': int((idx / total) * 100)
                                    })
                                    
                                    # Create chunk directory
                                    chunk_dir = os.path.join(outputs_dir, 'tmp', 'chunks', upload_id)
                                    chunks = split_audio_to_wav_chunks(
                                        filepath, chunk_dir,
                                        chunk_seconds=chunk_seconds,
                                        overlap_seconds=overlap_seconds
                                    )
                                    
                                    total_chunks = len(chunks)
                                    
                                    def chunk_progress_callback(chunk_idx, total, status):
                                        if is_cancelled():
                                            return
                                        update_manifest(progress={
                                            'currentFile': f'{original_name} ({status})',
                                            'currentFileIndex': idx,
                                            'percent': int((idx / total) * 100),
                                            'stage': 'diarizing',
                                            'chunkIndex': chunk_idx,
                                            'totalChunks': total
                                        })
                                    
                                    # Load pipeline once
                                    update_manifest(progress={
                                        'currentFile': f'{original_name} (loading diarization pipeline...)',
                                        'currentFileIndex': idx,
                                        'percent': int((idx / total) * 100),
                                        'totalChunks': total_chunks
                                    })
                                    
                                    pipeline = diarization.load_pipeline(device=device, job_id=job_id, session_id=session_id)
                                    
                                    # Run chunked diarization
                                    speaker_segments = diarization.run_chunked_diarization(
                                        filepath, chunks, pipeline,
                                        min_speakers=options.get('minSpeakers'),
                                        max_speakers=options.get('maxSpeakers'),
                                        num_speakers=options.get('numSpeakers'),
                                        job_id=job_id,
                                        session_id=session_id,
                                        progress_callback=chunk_progress_callback
                                    )
                                    
                                    # Cleanup chunks unless KEEP_CHUNKS is set
                                    keep_chunks = os.environ.get('KEEP_CHUNKS', '0') == '1'
                                    if keep_chunks:
                                        # Store chunk debug info in manifest
                                        chunk_debug_info = [{
                                            'index': c.index,
                                            'startSec': c.start_sec,
                                            'endSec': c.end_sec,
                                            'path': os.path.basename(c.path)
                                        } for c in chunks]
                                        update_manifest(debug={'chunks': chunk_debug_info})
                                    else:
                                        cleanup_chunks(chunk_dir)
                                else:
                                    # Standard diarization for short files
                                    update_manifest(progress={
                                        'currentFile': f'{original_name} (running diarization...)',
                                        'currentFileIndex': idx,
                                        'percent': int((idx / total) * 100)
                                    })
                                    
                                    speaker_segments = diarization.run_diarization(
                                        filepath,
                                        min_speakers=options.get('minSpeakers'),
                                        max_speakers=options.get('maxSpeakers'),
                                        num_speakers=options.get('numSpeakers'),
                                        device=device,
                                        job_id=job_id,
                                        session_id=session_id,
                                        temp_dir=os.path.join(outputs_dir, 'tmp')
                                    )
                                
                                diarization_result.put(speaker_segments)
                            except Exception as e:
                                diarization_error.put(e)
                        
                        # Start diarization in thread
                        diarization_thread = Thread(target=run_diarization_with_timeout)
                        diarization_thread.start()
                        
                        # Wait for completion or timeout
                        diarization_thread.join(timeout=max_diarization_minutes * 60)
                        
                        # Check if still running (timeout)
                        if diarization_thread.is_alive():
                            fail_job(
                                'DIARIZATION_TIMEOUT',
                                f'Diarization timed out after {max_diarization_minutes} minutes',
                                'diarization',
                                'Try with a shorter audio file or increase MAX_DIARIZATION_MINUTES',
                                outputs
                            )
                            return
                        
                        # Check for cancellation
                        if is_cancelled():
                            cancel_job(outputs)
                            return
                        
                        # Get result or error
                        try:
                            speaker_segments = diarization_result.get_nowait()
                        except Empty:
                            try:
                                error = diarization_error.get_nowait()
                                raise error
                            except Empty:
                                fail_job(
                                    'DIARIZATION_ERROR',
                                    'Diarization failed without error message',
                                    'diarization',
                                    outputs=outputs
                                )
                                return
                        
                        update_manifest(progress={
                            'currentFile': f'{original_name} (merging transcript + speakers...)',
                            'currentFileIndex': idx,
                            'percent': int((idx / total) * 100),
                            'stage': 'merging'
                        })
                        
                        merged_segments = diarization.merge_transcript_with_speakers(
                            result.get('segments', []),
                            speaker_segments,
                            job_id=job_id,
                            session_id=session_id
                        )
                    except ImportError as e:
                        # pyannote not installed - record error but continue without diarization
                        outputs.append({
                            'id': session_store.new_id(8),
                            'forUploadId': upload_id,
                            'inputFilename': original_name,
                            'error': {'code': 'DIARIZATION_UNAVAILABLE', 'message': str(e)[:200]}
                        })
                    except Exception as e:
                        # Diarization failed - record error but continue with transcript
                        outputs.append({
                            'id': session_store.new_id(8),
                            'forUploadId': upload_id,
                            'inputFilename': original_name,
                            'error': {'code': 'DIARIZATION_ERROR', 'message': str(e)[:200]}
                        })
                
                # Check for cancellation before writing outputs
                if is_cancelled():
                    cancel_job(outputs)
                    return
                
                update_manifest(progress={
                    'currentFile': f'{original_name} (writing outputs...)',
                    'currentFileIndex': idx,
                    'percent': int((idx / total) * 100),
                    'stage': 'writing_outputs'
                })
                
                # Generate output files with upload_id suffix to prevent collisions
                output_id = session_store.new_id(8)
                suffix = f"_{upload_id[:6]}" if upload_id else ""
                
                # Markdown output
                md_filename = f"{base_name}{suffix}.md"
                md_path = os.path.join(outputs_dir, md_filename)
                
                content = f"# {original_name}\n\n"
                if options.get('includeFull', True):
                    content += f"## Full Transcript\n\n{result['text'].strip()}\n\n"
                
                if options.get('includeSegments', True) and result.get('segments'):
                    content += "## Segments\n\n"
                    for seg in result['segments']:
                        if options.get('includeTimestamps', True):
                            start = seg.get('start', 0)
                            end = seg.get('end', 0)
                            content += f"**[{start:.2f} - {end:.2f}]** "
                        content += f"{seg.get('text', '').strip()}\n\n"
                
                with open(md_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                
                outputs.append({
                    'id': output_id,
                    'forUploadId': upload_id,
                    'inputFilename': original_name,
                    'filename': md_filename,
                    'path': md_path,
                    'type': 'markdown',
                    'sizeBytes': os.path.getsize(md_path)
                })
                
                # JSON output
                json_filename = f"{base_name}{suffix}.json"
                json_path = os.path.join(outputs_dir, json_filename)
                
                json_output = {
                    'source': original_name,
                    'text': result['text'],
                    'language': result.get('language', ''),
                    'segments': result.get('segments', [])
                }
                
                with open(json_path, 'w', encoding='utf-8') as f:
                    json.dump(json_output, f, indent=2, ensure_ascii=False)
                
                outputs.append({
                    'id': session_store.new_id(8),
                    'forUploadId': upload_id,
                    'inputFilename': original_name,
                    'filename': json_filename,
                    'path': json_path,
                    'type': 'json',
                    'sizeBytes': os.path.getsize(json_path)
                })
                
                # Generate diarization outputs if diarization was run (even if no speakers detected)
                if diarization_enabled and speaker_segments is not None:
                    import diarization as diarization_module
                    
                    # Speaker markdown (primary reviewer artifact)
                    speaker_md_filename = f"{base_name}{suffix}.speaker.md"
                    speaker_md_path = os.path.join(outputs_dir, speaker_md_filename)
                    speaker_md_content = diarization_module.format_speaker_markdown(
                        merged_segments, original_name, transcript_segments=result.get('segments', [])
                    )
                    with open(speaker_md_path, 'w', encoding='utf-8') as f:
                        f.write(speaker_md_content)
                    
                    outputs.append({
                        'id': session_store.new_id(8),
                        'forUploadId': upload_id,
                        'inputFilename': original_name,
                        'filename': speaker_md_filename,
                        'path': speaker_md_path,
                        'type': 'speaker-markdown',
                        'sizeBytes': os.path.getsize(speaker_md_path)
                    })
                    
                    # Diarization JSON (structured data)
                    diar_json_filename = f"{base_name}{suffix}.diarization.json"
                    diar_json_path = os.path.join(outputs_dir, diar_json_filename)
                    diar_json_output = diarization_module.format_diarization_json(
                        speaker_segments, merged_segments, original_name
                    )
                    with open(diar_json_path, 'w', encoding='utf-8') as f:
                        json.dump(diar_json_output, f, indent=2, ensure_ascii=False)
                    
                    outputs.append({
                        'id': session_store.new_id(8),
                        'forUploadId': upload_id,
                        'inputFilename': original_name,
                        'filename': diar_json_filename,
                        'path': diar_json_path,
                        'type': 'diarization-json',
                        'sizeBytes': os.path.getsize(diar_json_path)
                    })
                    
                    # RTTM output (interop/debug)
                    rttm_filename = f"{base_name}{suffix}.rttm"
                    rttm_path = os.path.join(outputs_dir, rttm_filename)
                    rttm_content = diarization_module.format_rttm(speaker_segments, original_name)
                    with open(rttm_path, 'w', encoding='utf-8') as f:
                        f.write(rttm_content)
                    
                    outputs.append({
                        'id': session_store.new_id(8),
                        'forUploadId': upload_id,
                        'inputFilename': original_name,
                        'filename': rttm_filename,
                        'path': rttm_path,
                        'type': 'rttm',
                        'sizeBytes': os.path.getsize(rttm_path)
                    })
                
            except Exception as e:
                # Record error but continue with other files
                outputs.append({
                    'id': session_store.new_id(8),
                    'forUploadId': upload_id,
                    'inputFilename': original_name,
                    'error': {'code': 'TRANSCRIBE_ERROR', 'message': str(e)[:200]}
                })
        
        # Determine final status - complete_with_errors if some files failed
        has_errors = any(o.get('error') for o in outputs)
        final_status = 'complete_with_errors' if has_errors else 'complete'
        
        # Log job completion
        log_event('info', 'job_finished', 
                 jobId=job_id,
                 sessionId=session_id,
                 status=final_status,
                 numOutputs=len(outputs),
                 numErrors=sum(1 for o in outputs if o.get('error')))
        
        update_manifest(
            status=final_status,
            finishedAt=datetime.now(timezone.utc).isoformat(),
            outputs=outputs,
            progress={'currentFileIndex': total, 'percent': 100, 'currentFile': '', 'stage': 'complete'}
        )
        
    except Exception as e:
        # Catch-all for unexpected errors
        log_event('error', 'job_failed', 
                 jobId=job_id,
                 sessionId=session_id,
                 error=str(e),
                 phase='unknown')
        
        fail_job('WORKER_EXCEPTION', f'Unexpected error: {str(e)[:400]}',
                 phase='unknown',
                 suggestion='Check server logs for details.',
                 outputs=outputs if 'outputs' in dir() else None)
    
    finally:
        with _active_jobs_lock:
            if (session_id, job_id) in _active_jobs:
                _active_jobs[(session_id, job_id)]['running'] = False


def _run_remote_session_job(session_id: str, job_id: str, inputs: list, options: dict, outputs_dir: str):
    """Run transcription job on remote GPU worker."""
    from logger import log_event
    from remote_worker import dispatch_to_remote_worker, get_worker_config
    
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    
    log_event('info', 'remote_job_started',
              jobId=job_id, sessionId=session_id,
              numFiles=len(inputs))
    
    def update_manifest(**updates):
        manifest = session_store.read_json(manifest_path) or {}
        now = datetime.now(timezone.utc).isoformat()
        manifest['updatedAt'] = now
        for key, value in updates.items():
            if key == 'progress':
                new_stage = value.get('stage')
                current_progress = manifest.get('progress') or {}
                old_stage = current_progress.get('stage')
                if new_stage and new_stage != old_stage:
                    value['stageStartedAt'] = now
                if manifest.get('progress') is None:
                    manifest['progress'] = {}
                manifest['progress'].update(value)
            elif key == 'worker':
                if manifest.get('worker') is None:
                    manifest['worker'] = {}
                manifest['worker'].update(value)
            elif key == 'remote':
                # Remote job metadata - merge to preserve existing fields
                if manifest.get('remote') is None:
                    manifest['remote'] = {}
                manifest['remote'].update(value)
            else:
                manifest[key] = value
        session_store.atomic_write_json(manifest_path, manifest)
        return manifest
    
    def is_cancelled():
        with _active_jobs_lock:
            job_info = _active_jobs.get((session_id, job_id), {})
            return job_info.get('cancel_requested', False)
    
    try:
        update_manifest(
            status='running',
            startedAt=datetime.now(timezone.utc).isoformat(),
            progress={'stage': 'dispatching', 'currentFile': 'Connecting to GPU worker...'}
        )
        
        # Get controller base URL from request context or environment
        controller_base_url = os.environ.get('CONTROLLER_BASE_URL', '')
        if not controller_base_url:
            # Try to construct from common patterns
            controller_base_url = os.environ.get('BASE_URL', 'http://localhost:8476')
        
        # Dispatch to remote worker (may return 'retry' for auto-retry on 404)
        result = dispatch_to_remote_worker(
            session_id=session_id,
            job_id=job_id,
            inputs=inputs,
            options=options,
            controller_base_url=controller_base_url,
            update_manifest_callback=update_manifest,
            is_cancelled_callback=is_cancelled
        )
        
        # Handle auto-retry on worker 404 (job lost)
        if result == 'retry':
            # Check current retry count from manifest
            manifest = session_store.read_json(manifest_path)
            remote_info = manifest.get('remote', {}) if manifest else {}
            retry_count = remote_info.get('retryCount', 0)
            
            if retry_count < 1:
                log_event('info', 'remote_job_auto_retry', jobId=job_id, retryCount=retry_count + 1)
                # Update retry count and reset status for retry
                update_manifest(
                    status='running',
                    remote={'retryCount': retry_count + 1},
                    error=None  # Clear error for retry
                )
                # Retry dispatch
                result = dispatch_to_remote_worker(
                    session_id=session_id,
                    job_id=job_id,
                    inputs=inputs,
                    options=options,
                    controller_base_url=controller_base_url,
                    update_manifest_callback=update_manifest,
                    is_cancelled_callback=is_cancelled
                )
        
        if result is True:
            log_event('info', 'remote_job_complete', jobId=job_id)
            # Final status is set by worker callback or polling
        else:
            # Job was cancelled or failed - status already updated
            log_event('info', 'remote_job_ended', jobId=job_id, success=False)
            
    except Exception as e:
        log_event('error', 'remote_job_exception', jobId=job_id, error=str(e))
        update_manifest(
            status='failed',
            finishedAt=datetime.now(timezone.utc).isoformat(),
            error={'code': 'REMOTE_DISPATCH_ERROR', 'message': str(e)[:400]}
        )
    
    finally:
        with _active_jobs_lock:
            if (session_id, job_id) in _active_jobs:
                _active_jobs[(session_id, job_id)]['running'] = False


def _repair_output_ids(outputs: list, inputs: list) -> list:
    """
    Repair outputs missing forUploadId by matching filename patterns.
    
    Output filenames are typically: <input_basename>_suffix.ext
    e.g., Call_27-05-2025_transcript.json from Call_27-05-2025.m4a
    """
    import re
    
    if not outputs or not inputs:
        return outputs
    
    # Build lookup: input basename -> uploadId
    input_lookup = {}
    for inp in inputs:
        original = inp.get('originalFilename') or inp.get('filename') or ''
        if original:
            basename = os.path.splitext(original)[0]
            input_lookup[basename.lower()] = inp.get('uploadId')
    
    repaired = []
    for output in outputs:
        out = dict(output)  # Copy to avoid mutating original
        
        # Already has forUploadId or inputId - use it
        existing_id = out.get('forUploadId') or out.get('inputId')
        if existing_id:
            out['forUploadId'] = existing_id
            repaired.append(out)
            continue
        
        # Try to match by filename pattern
        out_filename = out.get('filename') or ''
        if out_filename:
            # Strip common suffixes to get base name
            # Patterns: _transcript.json, _diarization.json, _speaker.md, .md, .json, .rttm
            base = re.sub(r'(_transcript|_diarization|_speaker|\.speaker)?(\.(json|md|rttm))$', '', out_filename, flags=re.IGNORECASE)
            # Also strip upload ID suffix if present (e.g., _abc123)
            base = re.sub(r'_[a-zA-Z0-9]{6}$', '', base)
            
            # Look up in inputs
            matched_id = input_lookup.get(base.lower())
            if matched_id:
                out['forUploadId'] = matched_id
        
        repaired.append(out)
    
    return repaired


@app.route('/api/jobs/<job_id>', methods=['GET'])
def api_get_job(job_id):
    """Get job status and details."""
    session_id = g.session_id
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    # Check for stale running jobs and mark them as failed
    if session_store.is_job_stale(manifest):
        session_store.mark_job_stale(session_id, job_id)
        manifest = session_store.read_json(manifest_path)
    
    # Repair outputs missing forUploadId using filename matching
    raw_outputs = manifest.get('outputs', [])
    raw_inputs = manifest.get('inputs', [])
    repaired_outputs = _repair_output_ids(raw_outputs, raw_inputs)
    
    # Don't expose internal paths
    safe_manifest = {
        'jobId': manifest.get('jobId'),
        'createdAt': manifest.get('createdAt'),
        'updatedAt': manifest.get('updatedAt'),
        'startedAt': manifest.get('startedAt'),
        'finishedAt': manifest.get('finishedAt'),
        'finished': manifest.get('finished', False),
        'status': manifest.get('status'),
        'backend': manifest.get('backend'),
        'executionMode': manifest.get('executionMode', 'local'),
        'environment': manifest.get('environment'),
        'options': manifest.get('options'),
        'inputs': [{'uploadId': i.get('uploadId'), 'filename': i.get('originalFilename'), 'durationSec': i.get('durationSec')} for i in manifest.get('inputs', [])],
        'outputs': [{'id': o.get('id'), 'forUploadId': o.get('forUploadId'), 'filename': o.get('filename'), 'type': o.get('type'), 'sizeBytes': o.get('sizeBytes'), 'error': o.get('error')} for o in repaired_outputs],
        'progress': manifest.get('progress'),
        'error': manifest.get('error')
    }
    
    # Include worker info if remote execution
    if manifest.get('executionMode') == 'remote' and manifest.get('worker'):
        worker = manifest['worker']
        safe_manifest['worker'] = {
            'workerJobId': worker.get('workerJobId'),
            'lastSeenAt': worker.get('lastSeenAt'),
            'gpu': worker.get('gpu')  # Include GPU status for accurate badge
        }
    
    # Include fallback info if remote was requested but ran locally
    if manifest.get('remoteRequested'):
        safe_manifest['remoteRequested'] = True
        if manifest.get('remoteFallbackReason'):
            safe_manifest['remoteFallbackReason'] = manifest['remoteFallbackReason']
    
    # Include last error info for operator visibility
    if manifest.get('lastErrorCode'):
        safe_manifest['lastErrorCode'] = manifest['lastErrorCode']
        safe_manifest['lastErrorMessage'] = manifest.get('lastErrorMessage', '')[:200]
        safe_manifest['lastErrorAt'] = manifest.get('lastErrorAt')
    
    return jsonify(safe_manifest)


@app.route('/api/jobs/<job_id>/cancel', methods=['POST'])
def api_cancel_job(job_id):
    """Cancel a running job immediately."""
    session_id = g.session_id
    
    # Check if job exists first
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    # Check terminal states (note: 'canceled' is the canonical spelling)
    if manifest.get('status') in ('complete', 'complete_with_errors', 'failed', 'canceled'):
        return jsonify({'status': manifest.get('status'), 'message': 'Job already finished'})
    
    with _active_jobs_lock:
        job_info = _active_jobs.get((session_id, job_id))
        if job_info and job_info.get('running'):
            job_info['cancel_requested'] = True
            
            # Immediately update manifest to show canceling state
            # This ensures UI sees the state change right away
            manifest['cancelRequested'] = True
            manifest['progress'] = manifest.get('progress', {})
            manifest['progress']['stage'] = 'canceling'
            manifest['progress']['currentFile'] = 'Canceling job...'
            manifest['updatedAt'] = datetime.now(timezone.utc).isoformat()
            session_store.atomic_write_json(manifest_path, manifest)
            
            return jsonify({'status': 'canceling', 'message': 'Cancel requested'})
    
    # Job not in active jobs but manifest shows running - might be stale
    if manifest.get('status') == 'running':
        # Force update to canceled since job thread is gone
        manifest['status'] = 'canceled'
        manifest['finishedAt'] = datetime.now(timezone.utc).isoformat()
        manifest['error'] = {'code': 'STALE_JOB', 'message': 'Job was orphaned and has been canceled'}
        session_store.atomic_write_json(manifest_path, manifest)
        return jsonify({'status': 'canceled', 'message': 'Orphaned job canceled'})
    
    return jsonify({'status': 'not_running'})


@app.route('/api/jobs/<job_id>/cancel-after-current', methods=['POST'])
def api_cancel_after_current(job_id):
    """Request graceful cancel after current file completes."""
    session_id = g.session_id
    
    with _active_jobs_lock:
        job_info = _active_jobs.get((session_id, job_id))
        if job_info and job_info.get('running'):
            job_info['cancel_after_current'] = True
            # Also update manifest so UI can see it
            manifest_path = session_store.job_manifest_path(session_id, job_id)
            manifest = session_store.read_json(manifest_path)
            if manifest:
                manifest['cancelAfterCurrentFile'] = True
                session_store.atomic_write_json(manifest_path, manifest)
            return jsonify({'status': 'pending_cancel', 'message': 'Will cancel after current file'})
    
    # Check if job exists
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    if manifest.get('status') in ('complete', 'complete_with_errors', 'failed', 'canceled'):
        return jsonify({'status': manifest.get('status'), 'message': 'Job already finished'})
    
    return jsonify({'status': 'not_running'})


@app.route('/api/jobs/<job_id>', methods=['DELETE'])
def api_delete_job(job_id):
    """Delete a job and its outputs."""
    session_id = g.session_id
    
    # Check if job is running
    with _active_jobs_lock:
        if (session_id, job_id) in _active_jobs:
            job_info = _active_jobs[(session_id, job_id)]
            if job_info.get('running'):
                return jsonify({'error': 'Cannot delete a running job'}), 409
    
    # Get job directory
    job_dir = session_store.job_dir(session_id, job_id)
    if not job_dir or not os.path.exists(job_dir):
        return jsonify({'error': 'Job not found'}), 404
    
    # Delete job directory
    import shutil
    try:
        shutil.rmtree(job_dir)
        return jsonify({'status': 'deleted', 'jobId': job_id})
    except Exception as e:
        return jsonify({'error': f'Failed to delete job: {str(e)[:200]}'}), 500


@app.route('/api/jobs/<job_id>/mark-finished', methods=['POST'])
def api_mark_job_finished(job_id):
    """Mark a job as finished or unfinished."""
    session_id = g.session_id
    data = request.get_json() or {}
    finished = data.get('finished', True)
    
    job_dir = session_store.job_dir(session_id, job_id)
    if not job_dir or not os.path.exists(job_dir):
        return jsonify({'error': 'Job not found'}), 404
    
    # Load existing job.json
    job_json_path = os.path.join(job_dir, 'job.json')
    job_data = session_store.read_json(job_json_path) or {}
    
    # Update finished flag
    job_data['finished'] = finished
    
    # Save back
    session_store.atomic_write_json(job_json_path, job_data)
    
    return jsonify({'status': 'ok', 'finished': finished})

_CSRF_PROTECTED_ENDPOINTS.add('api_mark_job_finished')


@app.route('/api/jobs/bulk-delete', methods=['POST'])
def api_bulk_delete_jobs():
    """Bulk delete multiple jobs."""
    session_id = g.session_id
    data = request.get_json() or {}
    job_ids = data.get('jobIds', [])
    
    if not job_ids:
        return jsonify({'error': 'No job IDs provided'}), 400
    
    import shutil
    deleted = []
    skipped_running = []
    errors = []
    
    for job_id in job_ids:
        # Check if job is running
        with _active_jobs_lock:
            if (session_id, job_id) in _active_jobs:
                job_info = _active_jobs[(session_id, job_id)]
                if job_info.get('running'):
                    skipped_running.append(job_id)
                    continue
        
        # Get job directory
        job_dir = session_store.job_dir(session_id, job_id)
        if not job_dir or not os.path.exists(job_dir):
            errors.append({'jobId': job_id, 'error': 'Job not found'})
            continue
        
        # Delete job directory
        try:
            shutil.rmtree(job_dir)
            # Verify deletion actually happened
            if os.path.exists(job_dir):
                errors.append({'jobId': job_id, 'error': f'Directory still exists after rmtree: {job_dir}'})
            else:
                deleted.append(job_id)
        except Exception as e:
            errors.append({'jobId': job_id, 'error': str(e)[:200]})
    
    return jsonify({
        'deleted': deleted,
        'skippedRunning': skipped_running,
        'errors': errors
    })


@app.route('/api/jobs/clear', methods=['POST'])
def api_clear_jobs():
    """Clear old jobs from the session."""
    session_id = g.session_id
    data = request.get_json() or {}
    
    older_than_hours = data.get('olderThanHours', 24)
    
    # Get all jobs for this session
    jobs_dir = os.path.join(session_store.session_dir(session_id), 'jobs')
    if not jobs_dir or not os.path.exists(jobs_dir):
        return jsonify({'cleared': 0})
    
    import shutil
    from datetime import datetime, timezone, timedelta
    
    cutoff = datetime.now(timezone.utc) - timedelta(hours=older_than_hours)
    cleared = 0
    skipped_running = 0
    
    for job_id in os.listdir(jobs_dir):
        job_dir = os.path.join(jobs_dir, job_id)
        if not os.path.isdir(job_dir):
            continue
        
        # Check if running
        with _active_jobs_lock:
            if (session_id, job_id) in _active_jobs:
                job_info = _active_jobs[(session_id, job_id)]
                if job_info.get('running'):
                    skipped_running += 1
                    continue
        
        # Check job age - use job.json (not manifest.json)
        manifest_path = os.path.join(job_dir, 'job.json')
        if os.path.exists(manifest_path):
            try:
                manifest = session_store.read_json(manifest_path)
                if manifest:
                    created_at = manifest.get('createdAt')
                    if created_at:
                        job_time = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                        if job_time >= cutoff:
                            continue  # Job is too recent
            except Exception:
                pass  # Skip unreadable jobs safely
        
        # Delete job
        try:
            shutil.rmtree(job_dir)
            cleared += 1
        except Exception:
            pass
    
    return jsonify({
        'cleared': cleared,
        'skippedRunning': skipped_running,
        'olderThanHours': older_than_hours
    })


@app.route('/api/jobs/clear-duplicates', methods=['POST'])
def api_clear_duplicate_jobs():
    """Clear duplicate jobs, keeping most recent per filename."""
    session_id = g.session_id
    
    session_dir = session_store.session_dir(session_id)
    if not session_dir:
        return jsonify({'kept': 0, 'deleted': 0, 'byFilename': []})
    
    jobs_dir = os.path.join(session_dir, 'jobs')
    if not os.path.exists(jobs_dir):
        return jsonify({'kept': 0, 'deleted': 0, 'byFilename': []})
    
    import shutil
    from collections import defaultdict
    
    # Group jobs by primary filename
    jobs_by_filename = defaultdict(list)
    
    for job_id in os.listdir(jobs_dir):
        job_dir = os.path.join(jobs_dir, job_id)
        if not os.path.isdir(job_dir):
            continue
        
        manifest_path = os.path.join(job_dir, 'job.json')
        if not os.path.exists(manifest_path):
            continue
        
        manifest = session_store.read_json(manifest_path)
        if not manifest:
            continue
        
        # Get primary filename (first input)
        inputs = manifest.get('inputs', [])
        if inputs:
            filename = inputs[0].get('originalFilename') or inputs[0].get('filename') or inputs[0].get('storedName') or ''
        else:
            filename = ''
        
        # Normalize filename for grouping (case-insensitive)
        filename_key = filename.lower().strip()
        
        jobs_by_filename[filename_key].append({
            'jobId': job_id,
            'jobDir': job_dir,
            'filename': filename,
            'createdAt': manifest.get('createdAt', ''),
            'status': manifest.get('status', 'unknown')
        })
    
    kept = 0
    deleted = 0
    by_filename = []
    
    for filename_key, jobs in jobs_by_filename.items():
        if len(jobs) <= 1:
            kept += len(jobs)
            continue
        
        # Sort by createdAt descending (most recent first)
        jobs.sort(key=lambda j: j['createdAt'] or '', reverse=True)
        
        # Keep the first (most recent), delete the rest if not running
        kept_job = jobs[0]
        deleted_job_ids = []
        
        for job in jobs[1:]:
            # Never delete running jobs
            is_running = False
            with _active_jobs_lock:
                if (session_id, job['jobId']) in _active_jobs:
                    job_info = _active_jobs[(session_id, job['jobId'])]
                    if job_info.get('running'):
                        is_running = True
            
            if is_running:
                kept += 1
                continue
            
            # Delete this duplicate
            try:
                shutil.rmtree(job['jobDir'])
                deleted += 1
                deleted_job_ids.append(job['jobId'])
            except Exception:
                kept += 1
        
        kept += 1  # The kept job
        
        if deleted_job_ids:
            by_filename.append({
                'filename': kept_job['filename'],
                'keptJobId': kept_job['jobId'],
                'deletedJobIds': deleted_job_ids
            })
    
    return jsonify({
        'kept': kept,
        'deleted': deleted,
        'byFilename': by_filename
    })


@app.route('/api/jobs/<job_id>/rerun', methods=['POST'])
def api_rerun_job(job_id):
    """Create a new job using the same input files from a previous job."""
    if not session_store.is_server_mode():
        return jsonify({'error': 'Rerun only available in server mode'}), 400
    
    session_id = g.session_id
    
    # Check for existing running job
    with _active_jobs_lock:
        for key, job_info in _active_jobs.items():
            if key[0] == session_id and job_info.get('running'):
                return jsonify({'error': 'A job is already running in this session'}), 409
    
    # Get original job
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    original_manifest = session_store.read_json(manifest_path)
    
    if not original_manifest:
        return jsonify({'error': 'Original job not found'}), 404
    
    # Verify input files still exist
    inputs = []
    for input_info in original_manifest.get('inputs', []):
        filepath = input_info.get('path')
        if not filepath or not os.path.exists(filepath):
            return jsonify({'error': f'Input file no longer exists: {input_info.get("originalFilename")}'}), 400
        inputs.append(input_info)
    
    if not inputs:
        return jsonify({'error': 'No input files found in original job'}), 400
    
    # Get new options from request or use original
    data = request.json or {}
    options = data.get('options', original_manifest.get('options', {}))
    
    # Determine backend with proper precedence:
    # 1. If REMOTE_WORKER_MODE=required, always use remote
    # 2. If options specify a backend, use that
    # 3. Otherwise, use original job's backend
    # 4. Fall back to environment's recommended backend
    env = compute_backend.get_cached_environment()
    original_backend = original_manifest.get('backend', '')
    
    # Check if remote worker is required
    try:
        from remote_worker import get_worker_config
        worker_config = get_worker_config()
        remote_required = worker_config.get('mode') == 'required'
        remote_available = worker_config.get('url') and worker_config.get('token')
    except ImportError:
        remote_required = False
        remote_available = False
    
    if remote_required and remote_available:
        # Force remote when mode is required
        requested_backend = 'remote'
    elif 'backend' in options:
        # Explicit backend in request options
        requested_backend = options['backend']
    elif original_backend:
        # Use original job's backend (preserve remote if it was remote)
        requested_backend = original_backend
    else:
        # Fall back to environment recommendation
        requested_backend = env['recommendedBackend']
    
    # Map UI concept "remote" to a real backend for validation
    actual_backend = requested_backend
    if requested_backend == 'remote':
        actual_backend = env['recommendedBackend']
    
    is_valid, error = compute_backend.validate_backend(actual_backend)
    if not is_valid:
        return jsonify({'error': error}), 400
    
    # Keep 'remote' as the backend value if that's what we're using
    # This ensures the job shows as "Remote GPU" in the UI
    if requested_backend == 'remote':
        actual_backend = 'remote'
    
    # Create new job
    new_job_id = session_store.new_id(12)
    dirs = session_store.ensure_job_dirs(session_id, new_job_id)
    
    now = datetime.now(timezone.utc).isoformat()
    manifest = {
        'jobId': new_job_id,
        'sessionId': session_id,
        'createdAt': now,
        'updatedAt': now,
        'startedAt': None,
        'finishedAt': None,
        'status': 'queued',
        'rerunOf': job_id,
        'backend': requested_backend,
        'environment': {
            'os': env['os'],
            'arch': env['arch'],
            'inDocker': env['inDocker']
        },
        'options': {
            'model': options.get('model', 'base'),
            'language': options.get('language', ''),
            'includeSegments': options.get('includeSegments', True),
            'includeFull': options.get('includeFull', True),
            'includeTimestamps': options.get('includeTimestamps', True),
            'wordTimestamps': options.get('wordTimestamps', False),
            'qualityPreset': options.get('qualityPreset', 'balanced'),
            # Diarization options - preserve from original job
            'diarizationEnabled': options.get('diarizationEnabled', False),
            'minSpeakers': options.get('minSpeakers'),
            'maxSpeakers': options.get('maxSpeakers'),
            'numSpeakers': options.get('numSpeakers'),
            'diarizationMaxDurationSeconds': options.get('diarizationMaxDurationSeconds'),
            'diarizationAutoSplit': options.get('diarizationAutoSplit', False),
            'diarizationFastSwitching': options.get('diarizationFastSwitching', False),
            'diarizationChunkSeconds': options.get('diarizationChunkSeconds'),
            'diarizationOverlapSeconds': options.get('diarizationOverlapSeconds'),
            'diarizationEffective': options.get('diarizationEffective'),
            # VAD and hallucination detection - preserve from original job
            'vadEnabled': options.get('vadEnabled', False),
            'hallucinationDetection': options.get('hallucinationDetection', True),
            'hallucinationSensitivity': options.get('hallucinationSensitivity', 50),
            'noSpeechThreshold': options.get('noSpeechThreshold', 0.6),
        },
        'inputs': inputs,
        'outputs': [],
        'progress': {
            'totalFiles': len(inputs),
            'currentFileIndex': 0,
            'currentFile': '',
            'percent': 0
        },
        'error': None
    }
    
    session_store.atomic_write_json(session_store.job_manifest_path(session_id, new_job_id), manifest)
    
    with _active_jobs_lock:
        _active_jobs[(session_id, new_job_id)] = {'running': True, 'cancel_requested': False}
    
    # Start job in background thread - dispatch to remote or local based on backend
    if requested_backend == 'remote':
        def run_remote_job():
            _run_remote_session_job(session_id, new_job_id, inputs, manifest['options'], dirs['outputs'])
        job_thread = threading.Thread(target=run_remote_job, daemon=True)
    else:
        def run_job():
            _run_session_job(session_id, new_job_id, inputs, manifest['options'], dirs['outputs'], requested_backend)
        job_thread = threading.Thread(target=run_job, daemon=True)
    
    job_thread.start()
    
    return jsonify({'jobId': new_job_id, 'rerunOf': job_id, 'executionMode': requested_backend})


# Add rerun to CSRF protected endpoints
_CSRF_PROTECTED_ENDPOINTS.add('api_rerun_job')


@app.route('/api/jobs/<job_id>/outputs', methods=['GET'])
def api_list_job_outputs(job_id):
    """List outputs for a job."""
    session_id = g.session_id
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    outputs = [
        {'id': o.get('id'), 'filename': o.get('filename'), 'type': o.get('type'), 'size': o.get('size')}
        for o in manifest.get('outputs', [])
        if not o.get('error')
    ]
    
    return jsonify({'outputs': outputs})


@app.route('/api/jobs/<job_id>/download', methods=['GET'])
def api_download_job(job_id):
    """Download all job outputs as a zip file."""
    session_id = g.session_id
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    outputs = manifest.get('outputs', [])
    files_to_zip = [o for o in outputs if o.get('path') and not o.get('error')]
    
    if not files_to_zip:
        return jsonify({'error': 'No output files available'}), 404
    
    # Verify all files exist and are within job directory
    job_outputs_path = session_store.job_outputs_dir(session_id, job_id)
    
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for output in files_to_zip:
            filepath = output.get('path')
            if filepath and os.path.exists(filepath) and session_store.is_safe_path(job_outputs_path, filepath):
                zf.write(filepath, output.get('filename'))
    
    zip_buffer.seek(0)
    
    timestamp = datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')
    zip_filename = f'transcriptions_{job_id[:8]}_{timestamp}.zip'
    
    return send_file(
        zip_buffer,
        mimetype='application/zip',
        as_attachment=True,
        download_name=zip_filename
    )


@app.route('/api/jobs/<job_id>/outputs/<output_id>', methods=['GET'])
def api_download_output(job_id, output_id):
    """Download a single output file."""
    session_id = g.session_id
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    # Find output by ID
    output = None
    for o in manifest.get('outputs', []):
        if o.get('id') == output_id:
            output = o
            break
    
    if not output or output.get('error'):
        return jsonify({'error': 'Output not found'}), 404
    
    filepath = output.get('path')
    job_outputs_path = session_store.job_outputs_dir(session_id, job_id)
    
    # Security: verify file is within job outputs directory
    if not filepath or not os.path.exists(filepath) or not session_store.is_safe_path(job_outputs_path, filepath):
        return jsonify({'error': 'Output file not accessible'}), 404
    
    return send_file(
        filepath,
        as_attachment=True,
        download_name=output.get('filename')
    )


@app.route('/api/jobs/<job_id>/audio/<input_id>', methods=['GET'])
def api_stream_audio(job_id, input_id):
    """
    Stream audio file for in-browser playback with Range support.
    
    input_id must match an uploadId in the job's inputs list.
    Supports HTTP Range requests for seeking in audio player.
    """
    session_id = g.session_id
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    # Find input by uploadId
    input_info = None
    for inp in manifest.get('inputs', []):
        if inp.get('uploadId') == input_id:
            input_info = inp
            break
    
    if not input_info:
        return jsonify({'error': 'Input not found'}), 404
    
    filepath = input_info.get('path')
    if not filepath or not os.path.exists(filepath):
        return jsonify({'error': 'Audio file not accessible'}), 404
    
    # Security: verify file is within session uploads directory
    session_uploads_dir = session_store.uploads_dir(session_id)
    if not session_store.is_safe_path(session_uploads_dir, filepath):
        return jsonify({'error': 'Access denied'}), 403
    
    # Determine content type from extension
    ext = os.path.splitext(filepath)[1].lower()
    content_types = {
        '.mp3': 'audio/mpeg',
        '.wav': 'audio/wav',
        '.m4a': 'audio/mp4',
        '.ogg': 'audio/ogg',
        '.webm': 'audio/webm',
        '.mp4': 'audio/mp4',
        '.flac': 'audio/flac',
        '.aac': 'audio/aac',
    }
    content_type = content_types.get(ext, 'application/octet-stream')
    
    file_size = os.path.getsize(filepath)
    
    # Handle Range requests for seeking
    range_header = request.headers.get('Range')
    if range_header:
        # Parse Range header (e.g., "bytes=0-1023")
        try:
            range_match = re.match(r'bytes=(\d+)-(\d*)', range_header)
            if range_match:
                start = int(range_match.group(1))
                end_str = range_match.group(2)
                end = int(end_str) if end_str else file_size - 1
                
                # Clamp to file size
                end = min(end, file_size - 1)
                
                if start > end or start >= file_size:
                    return Response('Range not satisfiable', status=416)
                
                length = end - start + 1
                
                def generate_range():
                    with open(filepath, 'rb') as f:
                        f.seek(start)
                        remaining = length
                        chunk_size = 64 * 1024  # 64KB chunks
                        while remaining > 0:
                            read_size = min(chunk_size, remaining)
                            data = f.read(read_size)
                            if not data:
                                break
                            remaining -= len(data)
                            yield data
                
                response = Response(
                    generate_range(),
                    status=206,
                    mimetype=content_type,
                    direct_passthrough=True
                )
                response.headers['Content-Range'] = f'bytes {start}-{end}/{file_size}'
                response.headers['Content-Length'] = length
                response.headers['Accept-Ranges'] = 'bytes'
                response.headers['Cache-Control'] = 'no-cache'
                return response
        except (ValueError, AttributeError):
            pass  # Fall through to full file response
    
    # Full file response (no Range or invalid Range)
    def generate_full():
        with open(filepath, 'rb') as f:
            chunk_size = 64 * 1024
            while True:
                data = f.read(chunk_size)
                if not data:
                    break
                yield data
    
    response = Response(
        generate_full(),
        status=200,
        mimetype=content_type,
        direct_passthrough=True
    )
    response.headers['Content-Length'] = file_size
    response.headers['Accept-Ranges'] = 'bytes'
    response.headers['Cache-Control'] = 'no-cache'
    return response


@app.route('/api/jobs/<job_id>/outputs/<output_id>/text', methods=['GET'])
def api_get_output_text(job_id, output_id):
    """
    Get text content of an output file for inline preview.
    
    Only allows text-like file types (.md, .txt, .json, .rttm, .srt, .vtt).
    Returns JSON with content and mime type.
    Supports ?maxBytes= query param to limit response size.
    """
    session_id = g.session_id
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    # Find output by ID
    output = None
    for o in manifest.get('outputs', []):
        if o.get('id') == output_id:
            output = o
            break
    
    if not output or output.get('error'):
        return jsonify({'error': 'Output not found'}), 404
    
    filepath = output.get('path')
    job_outputs_path = session_store.job_outputs_dir(session_id, job_id)
    
    # Security: verify file is within job outputs directory
    if not filepath or not os.path.exists(filepath) or not session_store.is_safe_path(job_outputs_path, filepath):
        return jsonify({'error': 'Output file not accessible'}), 404
    
    # Only allow text-like file types
    ext = os.path.splitext(filepath)[1].lower()
    allowed_extensions = {
        '.md': 'text/markdown',
        '.txt': 'text/plain',
        '.json': 'application/json',
        '.rttm': 'text/plain',
        '.srt': 'text/plain',
        '.vtt': 'text/vtt',
    }
    
    if ext not in allowed_extensions:
        return jsonify({'error': 'File type not supported for text preview'}), 400
    
    mime_type = allowed_extensions[ext]
    
    # Get max bytes limit (default 2MB, max 5MB)
    max_bytes = min(
        int(request.args.get('maxBytes', 2 * 1024 * 1024)),
        5 * 1024 * 1024
    )
    
    file_size = os.path.getsize(filepath)
    truncated = file_size > max_bytes
    
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read(max_bytes)
    except UnicodeDecodeError:
        return jsonify({'error': 'File is not valid UTF-8 text'}), 400
    except Exception as e:
        return jsonify({'error': f'Failed to read file: {str(e)[:100]}'}), 500
    
    return jsonify({
        'content': content,
        'mime': mime_type,
        'sizeBytes': file_size,
        'truncated': truncated,
        'filename': output.get('filename')
    })


@app.route('/api/jobs/<job_id>/speakers', methods=['GET'])
def api_get_speakers(job_id):
    """
    Get detected speakers and current label mapping for a job.
    
    Returns speakers detected from diarization.json or inferred from speaker.md,
    along with any custom labels that have been set.
    """
    session_id = g.session_id
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    # Get current label mapping from manifest
    speaker_labels = manifest.get('speakerLabels', {})
    
    # Try to detect speakers from diarization.json outputs
    detected_speakers = set()
    
    for output in manifest.get('outputs', []):
        if output.get('type') == 'diarization_json' and not output.get('error'):
            filepath = output.get('path')
            if filepath and os.path.exists(filepath):
                try:
                    diarization_data = session_store.read_json(filepath)
                    if diarization_data and 'segments' in diarization_data:
                        for seg in diarization_data['segments']:
                            speaker = seg.get('speaker')
                            if speaker:
                                detected_speakers.add(speaker)
                except Exception:
                    pass
    
    # If no diarization.json, try to infer from speaker.md content
    if not detected_speakers:
        for output in manifest.get('outputs', []):
            if output.get('type') == 'speaker_markdown' and not output.get('error'):
                filepath = output.get('path')
                if filepath and os.path.exists(filepath):
                    try:
                        with open(filepath, 'r', encoding='utf-8') as f:
                            content = f.read()
                        # Look for patterns like "SPEAKER_00:" or "Speaker 1:"
                        import re
                        # Match SPEAKER_XX format
                        matches = re.findall(r'\b(SPEAKER_\d+)\b', content)
                        detected_speakers.update(matches)
                        # Also match "Speaker N:" format
                        speaker_n_matches = re.findall(r'\bSpeaker\s+(\d+):', content)
                        for n in speaker_n_matches:
                            detected_speakers.add(f'SPEAKER_{int(n)-1:02d}')
                    except Exception:
                        pass
    
    # Sort speakers for consistent ordering
    speakers = sorted(list(detected_speakers))
    
    # Build response with default labels for any speakers without custom labels
    speaker_info = []
    for speaker_id in speakers:
        speaker_info.append({
            'id': speaker_id,
            'label': speaker_labels.get(speaker_id),
            'defaultLabel': f'Speaker {int(speaker_id.split("_")[1]) + 1}' if '_' in speaker_id else speaker_id
        })
    
    return jsonify({
        'speakers': speaker_info,
        'labels': speaker_labels
    })


@app.route('/api/jobs/<job_id>/speakers', methods=['PUT'])
def api_put_speakers(job_id):
    """
    Update speaker label mapping for a job.
    
    Body: { "labels": { "SPEAKER_00": "Child", "SPEAKER_01": "Parent" } }
    Labels are persisted in the job manifest.
    """
    session_id = g.session_id
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    data = request.get_json() or {}
    labels = data.get('labels', {})
    
    if not isinstance(labels, dict):
        return jsonify({'error': 'labels must be an object'}), 400
    
    # Validate labels
    MAX_LABEL_LENGTH = 40
    ALLOWED_CHARS = re.compile(r'^[\w\s\-\'\.]+$', re.UNICODE)
    
    validated_labels = {}
    for speaker_id, label in labels.items():
        # Validate speaker ID format
        if not re.match(r'^SPEAKER_\d+$', speaker_id):
            return jsonify({'error': f'Invalid speaker ID format: {speaker_id}'}), 400
        
        # Validate label
        if label is None or label == '':
            # Empty label means remove custom label (use default)
            continue
        
        if not isinstance(label, str):
            return jsonify({'error': f'Label for {speaker_id} must be a string'}), 400
        
        label = label.strip()
        if len(label) > MAX_LABEL_LENGTH:
            return jsonify({'error': f'Label for {speaker_id} exceeds {MAX_LABEL_LENGTH} characters'}), 400
        
        if not ALLOWED_CHARS.match(label):
            return jsonify({'error': f'Label for {speaker_id} contains invalid characters'}), 400
        
        validated_labels[speaker_id] = label
    
    # Update manifest with new labels
    manifest['speakerLabels'] = validated_labels
    manifest['updatedAt'] = datetime.now(timezone.utc).isoformat()
    
    # Atomic write
    session_store.atomic_write_json(manifest_path, manifest)
    
    return jsonify({
        'labels': validated_labels,
        'message': 'Speaker labels updated'
    })

_CSRF_PROTECTED_ENDPOINTS.add('api_put_speakers')


# =============================================================================
# REVIEW STATE + TIMELINE ENDPOINTS
# =============================================================================

def _invalidate_timeline_cache(job_dir: str, input_id: Optional[str] = None):
    """
    Remove cached timeline files so subsequent loads rebuild from outputs.

    Args:
        job_dir: Path to the job directory
        input_id: If provided, only invalidate cache for this input.
                  If None, invalidate all cached timelines in the job.
    """
    review_dir = os.path.join(job_dir, 'review')
    if not os.path.isdir(review_dir):
        return

    if input_id:
        # Invalidate specific input's cache
        cached_dir = os.path.join(review_dir, input_id)
        if os.path.isdir(cached_dir):
            for cached_file in ['timeline_merged.json', 'timeline_raw.json']:
                cached_path = os.path.join(cached_dir, cached_file)
                try:
                    if os.path.exists(cached_path):
                        os.remove(cached_path)
                except Exception:
                    pass
    else:
        # Invalidate all inputs' caches
        for entry in os.listdir(review_dir):
            entry_path = os.path.join(review_dir, entry)
            if os.path.isdir(entry_path):
                for cached_file in ['timeline_merged.json', 'timeline_raw.json']:
                    cached_path = os.path.join(entry_path, cached_file)
                    try:
                        if os.path.exists(cached_path):
                            os.remove(cached_path)
                    except Exception:
                        pass


@app.route('/api/jobs/<job_id>/review/state', methods=['GET'])
def api_get_review_state(job_id):
    """
    Get the current review state for a job.
    
    Query params:
    - inputId: (optional) specific input to get state for
    
    Returns speaker label map, chunk edits, and UI preferences.
    New format uses perInput[inputId] scoping.
    """
    session_id = g.session_id
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    input_id = request.args.get('inputId')
    
    # Review state is stored in {job}/review/review_state.json
    job_dir = session_store.job_dir(session_id, job_id)
    review_dir = os.path.join(job_dir, 'review')
    state_path = os.path.join(review_dir, 'review_state.json')
    
    state = session_store.read_json(state_path)
    
    # Handle new perInput format
    if state and 'perInput' in state:
        # New format - return input-specific state merged with uiPrefs
        input_state = state.get('perInput', {}).get(input_id, {}) if input_id else {}
        return jsonify({
            'speakerLabelMap': input_state.get('speakerLabelMap', {}),
            'speakerColorMap': input_state.get('speakerColorMap', {}),
            'speakerNumberMap': input_state.get('speakerNumberMap', {}),
            'chunkEdits': input_state.get('chunkEdits', {}),
            'manualSpeakers': input_state.get('manualSpeakers', []),
            'uiPrefs': state.get('uiPrefs', {}),
            '_format': 'perInput'
        })
    
    # Legacy format - return as-is for backward compatibility
    if not state:
        state = {
            'speakerLabelMap': {},
            'chunkEdits': {},
            'uiPrefs': {}
        }
    
    return jsonify(state)


@app.route('/api/jobs/<job_id>/review/highlight-counts', methods=['GET'])
def api_get_highlight_counts(job_id):
    """
    Get highlight counts for all inputs in a job.
    Returns { inputId: count, ... }
    """
    session_id = g.session_id
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    # Review state is stored in {job}/review/review_state.json
    job_dir = session_store.job_dir(session_id, job_id)
    review_dir = os.path.join(job_dir, 'review')
    state_path = os.path.join(review_dir, 'review_state.json')
    
    state = session_store.read_json(state_path)
    counts = {}
    
    if state and 'perInput' in state:
        # New perInput format - count highlights per input
        for input_id, input_state in state.get('perInput', {}).items():
            chunk_edits = input_state.get('chunkEdits', {})
            highlight_count = sum(1 for edit in chunk_edits.values() if edit.get('highlighted'))
            if highlight_count > 0:
                counts[input_id] = highlight_count
    
    return jsonify(counts)


@app.route('/api/jobs/<job_id>/review/state', methods=['PUT'])
def api_put_review_state(job_id):
    """
    Update the review state for a job.
    
    Query params:
    - inputId: (required for perInput format) specific input to save state for
    
    Accepts:
    {
        "speakerLabelMap": {"SPEAKER_00": "Matt", ...},
        "chunkEdits": {"t_000001": {"speakerId": "SPEAKER_02"}, ...},
        "uiPrefs": {...}
    }
    
    New perInput format stores edits scoped per input file.
    """
    session_id = g.session_id
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    
    input_id = request.args.get('inputId')
    
    # Validate speaker labels
    ALLOWED_CHARS = re.compile(r'^[\w\s\-\'\.]+$', re.UNICODE)
    MAX_LABEL_LENGTH = 40
    
    speaker_labels = data.get('speakerLabelMap', {})
    for speaker_id, label in speaker_labels.items():
        if not re.match(r'^SPEAKER_\d+$', speaker_id):
            return jsonify({'error': f'Invalid speaker ID: {speaker_id}'}), 400
        if len(label) > MAX_LABEL_LENGTH:
            return jsonify({'error': f'Label for {speaker_id} exceeds {MAX_LABEL_LENGTH} chars'}), 400
        if not ALLOWED_CHARS.match(label):
            return jsonify({'error': f'Label for {speaker_id} contains invalid characters'}), 400
    
    # Validate chunk edits - allow split suffixes like t_000001_a, t_000001_b, etc.
    chunk_edits = data.get('chunkEdits', {})
    CHUNK_ID_PATTERN = re.compile(r'^[tr]_\d+(_[a-z])*$')
    for chunk_id, edit in chunk_edits.items():
        if not CHUNK_ID_PATTERN.match(chunk_id):
            return jsonify({'error': f'Invalid chunk ID: {chunk_id}'}), 400
        if 'speakerId' in edit:
            sid = edit['speakerId']
            if sid and not re.match(r'^SPEAKER_\d+$', sid):
                return jsonify({'error': f'Invalid speaker ID in edit: {sid}'}), 400
    
    # Validate speaker colors (hex format)
    speaker_colors = data.get('speakerColorMap', {})
    HEX_COLOR = re.compile(r'^#[0-9a-fA-F]{6}$')
    for speaker_id, color in speaker_colors.items():
        if not re.match(r'^SPEAKER_\d+$', speaker_id):
            return jsonify({'error': f'Invalid speaker ID in colorMap: {speaker_id}'}), 400
        if color and not HEX_COLOR.match(color):
            return jsonify({'error': f'Invalid color for {speaker_id}: {color}'}), 400
    
    # Ensure review directory exists
    job_dir = session_store.job_dir(session_id, job_id)
    review_dir = os.path.join(job_dir, 'review')
    os.makedirs(review_dir, exist_ok=True)
    state_path = os.path.join(review_dir, 'review_state.json')
    
    # Load existing state to preserve other inputs' data
    existing_state = session_store.read_json(state_path) or {}
    
    # Build input-specific state
    input_state = {
        'speakerLabelMap': speaker_labels,
        'speakerColorMap': speaker_colors,
        'speakerNumberMap': data.get('speakerNumberMap', {}),
        'chunkEdits': chunk_edits,
        'manualSpeakers': data.get('manualSpeakers', []),
    }
    
    # Use perInput format if inputId provided
    if input_id:
        # Migrate existing state to perInput format if needed
        if 'perInput' not in existing_state:
            existing_state = {
                'perInput': {},
                'uiPrefs': existing_state.get('uiPrefs', {})
            }
        
        # Update only this input's state
        existing_state['perInput'][input_id] = input_state
        existing_state['uiPrefs'] = data.get('uiPrefs', existing_state.get('uiPrefs', {}))
        existing_state['updatedAt'] = datetime.now(timezone.utc).isoformat()

        session_store.atomic_write_json(state_path, existing_state)

        # Invalidate cached timelines for this input so edits take effect
        _invalidate_timeline_cache(job_dir, input_id)

        return jsonify({
            'message': 'Review state saved',
            'inputId': input_id
        })
    else:
        # Legacy format (no inputId) - save flat structure
        state = {
            'speakerLabelMap': speaker_labels,
            'speakerColorMap': speaker_colors,
            'chunkEdits': chunk_edits,
            'uiPrefs': data.get('uiPrefs', {}),
            'updatedAt': datetime.now(timezone.utc).isoformat()
        }
        session_store.atomic_write_json(state_path, state)

        # Invalidate all cached timelines for this job (legacy mode affects all inputs)
        _invalidate_timeline_cache(job_dir, None)

        return jsonify({
            'message': 'Review state saved',
            'state': state
        })

_CSRF_PROTECTED_ENDPOINTS.add('api_put_review_state')


@app.route('/api/jobs/<job_id>/review/timeline', methods=['GET'])
def api_get_review_timeline(job_id):
    """
    Get the review timeline for a job input.
    
    Query params:
    - inputId: (optional) specific input to get timeline for, defaults to first
    
    Returns a ReviewTimeline with chunks, speakers, and edits applied.
    """
    from review_timeline import TimelineParser, apply_review_state
    from logger import log_event
    
    session_id = g.session_id
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    # Get input ID
    input_id = request.args.get('inputId')
    inputs = manifest.get('inputs', [])
    
    if not inputs:
        return jsonify({'error': 'No inputs in job'}), 400
    
    # Find the requested input or use first
    target_input = None
    if input_id:
        for inp in inputs:
            if inp.get('uploadId') == input_id:
                target_input = inp
                break
        if not target_input:
            return jsonify({'error': 'Input not found'}), 404
    else:
        target_input = inputs[0]
        input_id = target_input.get('uploadId')
    
    filename = target_input.get('originalFilename') or target_input.get('filename') or 'unknown'
    
    log_event('debug', 'review_timeline_input_lookup',
              jobId=job_id, inputId=input_id,
              targetInputKeys=list(target_input.keys()),
              filename=filename,
              inputsCount=len(inputs))
    
    # Find outputs for this input
    outputs = manifest.get('outputs', [])
    
    # Output type aliases for robustness
    TRANSCRIPT_JSON_TYPES = {'json', 'transcript_json', 'whisper_json', 'segments_json'}
    TRANSCRIPT_MD_TYPES = {'markdown', 'transcript_markdown', 'transcript_md'}
    SPEAKER_MD_TYPES = {'speaker-markdown', 'speaker_markdown', 'speaker_md'}
    DIARIZATION_TYPES = {'diarization-json', 'diarization_json'}
    
    # Read available output files
    transcript_json = None
    diarization_json = None
    speaker_md = None
    transcript_md = None
    
    # Track what we find for debugging
    available_outputs = []
    selected_outputs = {}
    matched_outputs = []
    skipped_outputs = []
    
    for output in outputs:
        output_info = {
            'id': output.get('id'),
            'type': output.get('type'),
            'filename': output.get('filename'),
            'forUploadId': output.get('forUploadId')
        }
        available_outputs.append(output_info)
        
        # Check forUploadId match - STRICT: require forUploadId to match
        # If forUploadId is missing, skip to fallback (filename-based matching)
        output_for_id = output.get('forUploadId')
        if not output_for_id:
            # No forUploadId - will be handled by filename-based fallback
            skipped_outputs.append({'reason': 'no_forUploadId', **output_info})
            continue
        if output_for_id != input_id:
            skipped_outputs.append({'reason': 'forUploadId_mismatch', **output_info})
            continue
        
        output_type = output.get('type', '')
        output_path = output.get('path')
        
        if not output_path:
            skipped_outputs.append({'reason': 'no_path', **output_info})
            continue
        
        if not os.path.exists(output_path):
            skipped_outputs.append({'reason': 'path_not_exists', 'path': output_path, **output_info})
            continue
        
        try:
            with open(output_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except Exception as e:
            skipped_outputs.append({'reason': 'read_error', 'error': str(e)[:50], **output_info})
            continue
        
        # Match output types with aliases
        if output_type in TRANSCRIPT_JSON_TYPES and not transcript_json:
            transcript_json = content
            selected_outputs['transcript_json'] = output_info
            matched_outputs.append(output_info)
        elif output_type in DIARIZATION_TYPES and not diarization_json:
            diarization_json = content
            selected_outputs['diarization_json'] = output_info
            matched_outputs.append(output_info)
        elif output_type in SPEAKER_MD_TYPES and not speaker_md:
            speaker_md = content
            selected_outputs['speaker_md'] = output_info
            matched_outputs.append(output_info)
        elif output_type in TRANSCRIPT_MD_TYPES and not transcript_md:
            transcript_md = content
            selected_outputs['transcript_md'] = output_info
            matched_outputs.append(output_info)
    
    # Fallback: if no outputs matched by forUploadId but outputs exist, try filename-based matching
    if not matched_outputs and outputs:
        import re
        
        # Extract base filename from input (e.g., "Call_27-05-2025.m4a" -> "Call_27-05-2025")
        input_basename = os.path.splitext(filename)[0] if filename else ''
        
        log_event('warning', 'review_timeline_fallback',
                  jobId=job_id, inputId=input_id,
                  reason='no_outputs_matched_inputId_trying_filename_match',
                  inputBasename=input_basename)
        
        def output_matches_input(output_filename: str, input_base: str) -> bool:
            """Check if output filename matches input basename (legacy job matching)."""
            if not output_filename or not input_base:
                return False
            # Output filenames are like "Call_27-05-2025_transcript.json" or "Call_27-05-2025_speaker.md"
            # Extract base by removing known suffixes
            output_base = re.sub(r'_(transcript|diarization|speaker)\.(json|md)$', '', output_filename)
            output_base = re.sub(r'\.rttm$', '', output_base)
            matches = output_base == input_base
            log_event('debug', 'review_timeline_filename_match',
                      outputFilename=output_filename, outputBase=output_base,
                      inputBase=input_base, matches=matches)
            return matches
        
        for output in outputs:
            if output.get('error'):
                continue
            
            output_type = output.get('type', '')
            output_path = output.get('path')
            output_filename = output.get('filename', '')
            
            # For legacy jobs, match by filename pattern
            if input_basename and not output_matches_input(output_filename, input_basename):
                continue
            
            if not output_path or not os.path.exists(output_path):
                continue
            
            try:
                with open(output_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            except Exception:
                continue
            
            output_info = {
                'id': output.get('id'),
                'type': output_type,
                'filename': output_filename,
                'forUploadId': output.get('forUploadId')
            }
            
            if output_type in TRANSCRIPT_JSON_TYPES and not transcript_json:
                transcript_json = content
                selected_outputs['transcript_json'] = output_info
            elif output_type in DIARIZATION_TYPES and not diarization_json:
                diarization_json = content
                selected_outputs['diarization_json'] = output_info
            elif output_type in SPEAKER_MD_TYPES and not speaker_md:
                speaker_md = content
                selected_outputs['speaker_md'] = output_info
            elif output_type in TRANSCRIPT_MD_TYPES and not transcript_md:
                transcript_md = content
                selected_outputs['transcript_md'] = output_info
    
    # Determine source used
    source_used = 'none'
    if diarization_json:
        source_used = 'diarization_json'
    elif transcript_json:
        source_used = 'transcript_json'
    elif speaker_md:
        source_used = 'speaker_md'
    elif transcript_md:
        source_used = 'transcript_md'
    
    # Parse into timeline
    parser = TimelineParser(job_id, input_id, filename)
    timeline = parser.parse(
        transcript_json=transcript_json,
        diarization_json=diarization_json,
        speaker_md=speaker_md,
        transcript_md=transcript_md
    )
    
    # Log debug info including dedupe stats
    dedupe_stats = getattr(timeline, 'dedupe_stats', {})
    log_event('info', 'review_timeline_debug',
              jobId=job_id, inputId=input_id,
              availableOutputs=available_outputs,
              selectedOutputs=selected_outputs,
              sourceUsed=source_used,
              numChunks=len(timeline.chunks),
              dedupeStats=dedupe_stats,
              skippedCount=len(skipped_outputs))
    
    # Hard guarantee: if we have text but no chunks, create a fallback chunk
    if len(timeline.chunks) == 0:
        fallback_text = None
        fallback_source = None
        
        # Try to extract any text from available content
        if transcript_json:
            try:
                data = json.loads(transcript_json)
                fallback_text = data.get('text', '')
                if not fallback_text and data.get('segments'):
                    fallback_text = ' '.join(s.get('text', '') for s in data['segments'])
                fallback_source = 'transcript_json_fallback'
            except Exception:
                pass
        
        if not fallback_text and diarization_json:
            try:
                data = json.loads(diarization_json)
                if data.get('segments'):
                    fallback_text = ' '.join(s.get('text', '') for s in data['segments'])
                    fallback_source = 'diarization_json_fallback'
            except Exception:
                pass
        
        if not fallback_text and speaker_md:
            fallback_text = speaker_md
            fallback_source = 'speaker_md_fallback'
        
        if not fallback_text and transcript_md:
            fallback_text = transcript_md
            fallback_source = 'transcript_md_fallback'
        
        if fallback_text and fallback_text.strip():
            from review_timeline import Chunk
            fallback_chunk = Chunk(
                chunk_id='t_000000',
                start=0.0,
                end=0.0,
                speaker_id='SPEAKER_00',
                text=fallback_text.strip()[:10000],  # Limit size
                origin={'source': fallback_source, 'fallback': True}
            )
            timeline.chunks.append(fallback_chunk)
            timeline.add_speaker('SPEAKER_00', 'Speaker 1')
            
            log_event('warning', 'review_timeline_fallback_chunk',
                      jobId=job_id, inputId=input_id,
                      fallbackSource=fallback_source,
                      textLength=len(fallback_text))
    
    # Apply saved review state if exists
    job_dir = session_store.job_dir(session_id, job_id)
    state_path = os.path.join(job_dir, 'review', 'review_state.json')
    review_state = session_store.read_json(state_path)
    
    if review_state:
        timeline = apply_review_state(timeline, review_state)
    
    return jsonify(timeline.to_dict())


@app.route('/api/jobs/<job_id>/review/regenerate', methods=['POST'])
def api_regenerate_review_timeline(job_id):
    """
    Regenerate timeline by scanning output files on disk and matching by filename.
    
    This is useful for legacy jobs where forUploadId wasn't stored, or when
    the manifest is out of sync with actual files.
    
    Query params:
    - inputId: (optional) specific input to regenerate for
    
    Body (optional):
    - filename: explicit filename to match outputs against
    """
    from review_timeline import TimelineParser, apply_review_state
    from logger import log_event
    import re
    
    session_id = g.session_id
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    input_id = request.args.get('inputId')
    body = request.get_json(silent=True) or {}
    explicit_filename = body.get('filename')
    
    inputs = manifest.get('inputs', [])
    outputs = manifest.get('outputs', [])
    
    if not inputs:
        return jsonify({'error': 'No inputs in job'}), 400
    
    # Find target input
    target_input = None
    if input_id:
        for inp in inputs:
            if inp.get('uploadId') == input_id:
                target_input = inp
                break
        if not target_input:
            return jsonify({'error': 'Input not found'}), 404
    else:
        target_input = inputs[0]
        input_id = target_input.get('uploadId')
    
    # Get filename to match against
    filename = explicit_filename or target_input.get('originalFilename') or target_input.get('filename') or 'unknown'
    input_basename = os.path.splitext(filename)[0] if filename else ''
    
    log_event('info', 'regenerate_timeline_start',
              jobId=job_id, inputId=input_id, filename=filename, inputBasename=input_basename)
    
    # Scan job directory for output files
    job_dir = session_store.job_dir(session_id, job_id)
    outputs_dir = os.path.join(job_dir, 'outputs')
    
    # Helper to match output filename to input
    def output_matches_input(output_filename: str, input_base: str) -> bool:
        if not output_filename or not input_base:
            return False
        output_base = re.sub(r'_(transcript|diarization|speaker)\.(json|md)$', '', output_filename)
        output_base = re.sub(r'\.rttm$', '', output_base)
        return output_base == input_base
    
    # Find matching files
    transcript_json = None
    diarization_json = None
    speaker_md = None
    transcript_md = None
    matched_files = []
    
    # First try manifest outputs with path
    for output in outputs:
        output_filename = output.get('filename', '')
        output_path = output.get('path')
        output_type = output.get('type', '')
        
        if not output_matches_input(output_filename, input_basename):
            continue
        
        if not output_path or not os.path.exists(output_path):
            continue
        
        try:
            with open(output_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except Exception:
            continue
        
        matched_files.append({'filename': output_filename, 'type': output_type, 'source': 'manifest'})
        
        if output_type in {'json', 'transcript_json'} and not transcript_json:
            transcript_json = content
        elif output_type in {'diarization-json', 'diarization_json'} and not diarization_json:
            diarization_json = content
        elif output_type in {'speaker-markdown', 'speaker_markdown'} and not speaker_md:
            speaker_md = content
        elif output_type in {'markdown', 'transcript_markdown'} and not transcript_md:
            transcript_md = content
    
    # If no matches from manifest, scan outputs directory directly
    if not matched_files and os.path.isdir(outputs_dir):
        for fname in os.listdir(outputs_dir):
            if not output_matches_input(fname, input_basename):
                continue
            
            fpath = os.path.join(outputs_dir, fname)
            if not os.path.isfile(fpath):
                continue
            
            try:
                with open(fpath, 'r', encoding='utf-8') as f:
                    content = f.read()
            except Exception:
                continue
            
            matched_files.append({'filename': fname, 'source': 'disk_scan'})
            
            if fname.endswith('_transcript.json') and not transcript_json:
                transcript_json = content
            elif fname.endswith('_diarization.json') and not diarization_json:
                diarization_json = content
            elif fname.endswith('_speaker.md') and not speaker_md:
                speaker_md = content
            elif fname.endswith('_transcript.md') and not transcript_md:
                transcript_md = content
    
    if not any([transcript_json, diarization_json, speaker_md, transcript_md]):
        return jsonify({
            'error': 'No matching output files found',
            'inputBasename': input_basename,
            'scannedOutputsDir': outputs_dir,
            'hint': 'Check that output files exist and match the input filename pattern'
        }), 404
    
    # Parse into timeline
    parser = TimelineParser(job_id, input_id, filename)
    timeline = parser.parse(
        transcript_json=transcript_json,
        diarization_json=diarization_json,
        speaker_md=speaker_md,
        transcript_md=transcript_md
    )
    
    log_event('info', 'regenerate_timeline_complete',
              jobId=job_id, inputId=input_id,
              matchedFiles=matched_files,
              numChunks=len(timeline.chunks),
              numSpeakers=len(timeline.speakers))

    # Clear cached timeline files so future loads use fresh data
    _invalidate_timeline_cache(job_dir, input_id)

    # Apply saved review state if exists
    state_path = os.path.join(job_dir, 'review', 'review_state.json')
    review_state = session_store.read_json(state_path)
    
    if review_state:
        timeline = apply_review_state(timeline, review_state)
    
    return jsonify({
        'success': True,
        'timeline': timeline.to_dict(),
        'matchedFiles': matched_files,
        'inputBasename': input_basename
    })


@app.route('/api/jobs/<job_id>/review/export', methods=['GET'])
def api_export_review_project(job_id):
    """
    Export a review project as a zip bundle.
    
    Query params:
    - inputId: (optional) specific input to export, defaults to first
    
    Returns a .btproj.zip containing:
    - manifest_snapshot.json
    - timeline.json
    - review_state.json
    - outputs/ (text outputs used for review)
    """
    from review_timeline import TimelineParser, apply_review_state
    
    session_id = g.session_id
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    input_id = request.args.get('inputId')
    inputs = manifest.get('inputs', [])
    
    if not inputs:
        return jsonify({'error': 'No inputs in job'}), 400
    
    target_input = None
    if input_id:
        for inp in inputs:
            if inp.get('uploadId') == input_id:
                target_input = inp
                break
        if not target_input:
            return jsonify({'error': 'Input not found'}), 404
    else:
        target_input = inputs[0]
        input_id = target_input.get('uploadId')
    
    filename = target_input.get('originalFilename', 'unknown')
    safe_name = secure_filename(Path(filename).stem) or 'project'
    
    # Build zip in memory
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        # Manifest snapshot (sanitized - no paths)
        safe_manifest = {
            'jobId': manifest.get('jobId'),
            'createdAt': manifest.get('createdAt'),
            'status': manifest.get('status'),
            'options': manifest.get('options'),
            'inputs': [{'uploadId': i.get('uploadId'), 'filename': i.get('originalFilename')} 
                      for i in manifest.get('inputs', [])],
            'outputs': [{'id': o.get('id'), 'filename': o.get('filename'), 'type': o.get('type')} 
                       for o in manifest.get('outputs', []) if not o.get('error')],
            'exportedAt': datetime.now(timezone.utc).isoformat(),
            'buildCommit': os.environ.get('BUILD_COMMIT', 'unknown'),
        }
        zf.writestr('manifest_snapshot.json', json.dumps(safe_manifest, indent=2))
        
        # Review state
        job_dir = session_store.job_dir(session_id, job_id)
        state_path = os.path.join(job_dir, 'review', 'review_state.json')
        review_state = session_store.read_json(state_path) or {}
        zf.writestr('review_state.json', json.dumps(review_state, indent=2))
        
        # Timeline
        outputs = manifest.get('outputs', [])
        transcript_json = None
        diarization_json = None
        speaker_md = None
        transcript_md = None
        
        for output in outputs:
            if output.get('forUploadId') != input_id and output.get('forUploadId'):
                continue
            
            output_type = output.get('type', '')
            output_path = output.get('path')
            output_filename = output.get('filename', '')
            
            if not output_path or not os.path.exists(output_path):
                continue
            
            try:
                with open(output_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                # Add to zip
                zf.writestr(f'outputs/{output_filename}', content)
                
                # Track for timeline parsing
                if output_type == 'json':
                    transcript_json = content
                elif output_type == 'diarization-json':
                    diarization_json = content
                elif output_type == 'speaker-markdown':
                    speaker_md = content
                elif output_type == 'markdown':
                    transcript_md = content
            except Exception:
                continue
        
        # Generate and include timeline
        parser = TimelineParser(job_id, input_id, filename)
        timeline = parser.parse(
            transcript_json=transcript_json,
            diarization_json=diarization_json,
            speaker_md=speaker_md,
            transcript_md=transcript_md
        )
        if review_state:
            timeline = apply_review_state(timeline, review_state)
        
        zf.writestr('timeline.json', timeline.to_json())
    
    zip_buffer.seek(0)
    
    return send_file(
        zip_buffer,
        mimetype='application/zip',
        as_attachment=True,
        download_name=f'{safe_name}.btproj.zip'
    )


@app.route('/api/jobs/<job_id>/review/export-md', methods=['GET'])
def api_export_review_md(job_id):
    """Export review as speaker markdown with timestamps."""
    from review_timeline import TimelineParser, apply_review_state
    
    session_id = g.session_id
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    input_id = request.args.get('inputId')
    inputs = manifest.get('inputs', [])
    if not inputs:
        return jsonify({'error': 'No inputs in job'}), 400
    
    target_input = inputs[0] if not input_id else next((i for i in inputs if i.get('uploadId') == input_id), inputs[0])
    input_id = target_input.get('uploadId')
    filename = target_input.get('originalFilename', target_input.get('filename', 'unknown'))
    safe_name = secure_filename(Path(filename).stem) or 'transcript'
    
    # Build timeline with review state (respect view mode)
    view_mode = request.args.get('viewMode', 'merged')
    timeline, _ = _build_review_timeline(session_id, job_id, input_id, filename, manifest, view_mode)
    
    # Format as markdown
    lines = []
    for chunk in timeline.chunks:
        speaker = next((s for s in timeline.speakers if s.id == chunk.speaker_id), None)
        label = speaker.label if speaker else chunk.speaker_id or 'Speaker'
        
        mins, secs = divmod(int(chunk.start), 60)
        hours, mins = divmod(mins, 60)
        timestamp = f'{hours:02d}:{mins:02d}:{secs:02d}' if hours else f'{mins:02d}:{secs:02d}'
        
        lines.append(f'[{timestamp}] {label}: {chunk.text}')
    
    content = '\n\n'.join(lines)
    
    return Response(
        content,
        mimetype='text/markdown',
        headers={'Content-Disposition': f'attachment; filename="{safe_name}_review.md"'}
    )


@app.route('/api/jobs/<job_id>/review/export-txt', methods=['GET'])
def api_export_review_txt(job_id):
    """Export review as plain text."""
    from review_timeline import TimelineParser, apply_review_state
    
    session_id = g.session_id
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    input_id = request.args.get('inputId')
    inputs = manifest.get('inputs', [])
    if not inputs:
        return jsonify({'error': 'No inputs in job'}), 400
    
    target_input = inputs[0] if not input_id else next((i for i in inputs if i.get('uploadId') == input_id), inputs[0])
    input_id = target_input.get('uploadId')
    filename = target_input.get('originalFilename', target_input.get('filename', 'unknown'))
    safe_name = secure_filename(Path(filename).stem) or 'transcript'
    
    # Build timeline with review state (respect view mode)
    view_mode = request.args.get('viewMode', 'merged')
    timeline, _ = _build_review_timeline(session_id, job_id, input_id, filename, manifest, view_mode)
    
    # Format as plain text
    lines = []
    for chunk in timeline.chunks:
        speaker = next((s for s in timeline.speakers if s.id == chunk.speaker_id), None)
        label = speaker.label if speaker else chunk.speaker_id or 'Speaker'
        lines.append(f'{label}: {chunk.text}')
    
    content = '\n\n'.join(lines)
    
    return Response(
        content,
        mimetype='text/plain',
        headers={'Content-Disposition': f'attachment; filename="{safe_name}_review.txt"'}
    )


@app.route('/api/jobs/<job_id>/review/export-timeline-json', methods=['GET'])
def api_export_review_timeline_json(job_id):
    """Export review timeline as JSON."""
    from review_timeline import TimelineParser, apply_review_state
    
    session_id = g.session_id
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    input_id = request.args.get('inputId')
    inputs = manifest.get('inputs', [])
    if not inputs:
        return jsonify({'error': 'No inputs in job'}), 400
    
    target_input = inputs[0] if not input_id else next((i for i in inputs if i.get('uploadId') == input_id), inputs[0])
    input_id = target_input.get('uploadId')
    filename = target_input.get('originalFilename', target_input.get('filename', 'unknown'))
    safe_name = secure_filename(Path(filename).stem) or 'transcript'
    
    # Build timeline with review state (respect view mode)
    view_mode = request.args.get('viewMode', 'merged')
    timeline, _ = _build_review_timeline(session_id, job_id, input_id, filename, manifest, view_mode)
    
    content = timeline.to_json()
    
    return Response(
        content,
        mimetype='application/json',
        headers={'Content-Disposition': f'attachment; filename="{safe_name}_review_timeline.json"'}
    )


def _build_review_timeline(session_id: str, job_id: str, input_id: str, filename: str, manifest: dict, view_mode: str = 'merged', force_rebuild: bool = False):
    """Build timeline with review state applied. Returns (timeline, review_state).

    Args:
        view_mode: 'merged' (default) or 'raw' - controls which chunk view to use for export
        force_rebuild: If True, skip cached timeline and rebuild from outputs
    """
    from review_timeline import TimelineParser, ReviewTimeline, Chunk, apply_review_state

    job_dir = session_store.job_dir(session_id, job_id)
    state_path = os.path.join(job_dir, 'review', 'review_state.json')
    review_state = session_store.read_json(state_path) or {}

    # Check for cached timeline file (from session import or previous export)
    if not force_rebuild and input_id:
        timeline_filename = 'timeline_merged.json' if view_mode == 'merged' else 'timeline_raw.json'
        cached_path = os.path.join(job_dir, 'review', input_id, timeline_filename)
        if os.path.exists(cached_path):
            try:
                with open(cached_path, 'r', encoding='utf-8') as f:
                    cached_json = f.read()
                timeline = ReviewTimeline.from_json(cached_json)
                # Normalize chunks_raw/chunks_merged to Chunk objects
                if timeline.chunks_raw:
                    timeline.chunks_raw = [Chunk.from_dict(c) if isinstance(c, dict) else c for c in timeline.chunks_raw]
                if timeline.chunks_merged:
                    timeline.chunks_merged = [Chunk.from_dict(c) if isinstance(c, dict) else c for c in timeline.chunks_merged]
                # Apply view mode from cached data
                if view_mode == 'raw' and timeline.chunks_raw:
                    timeline.chunks = list(timeline.chunks_raw)
                elif timeline.chunks_merged:
                    timeline.chunks = list(timeline.chunks_merged)
                # Re-apply review state in case edits were made after cache was created
                if review_state:
                    timeline = apply_review_state(timeline, review_state)
                return timeline, review_state
            except Exception:
                # Fall through to rebuild if cache is invalid
                pass

    # Build from output files
    outputs = manifest.get('outputs', [])
    transcript_json = None
    diarization_json = None
    speaker_md = None
    transcript_md = None

    for output in outputs:
        if output.get('forUploadId') != input_id and output.get('forUploadId'):
            continue

        output_type = output.get('type', '')
        output_path = output.get('path')

        if not output_path or not os.path.exists(output_path):
            continue

        try:
            with open(output_path, 'r', encoding='utf-8') as f:
                content = f.read()

            if output_type == 'json':
                transcript_json = content
            elif output_type == 'diarization-json':
                diarization_json = content
            elif output_type == 'speaker-markdown':
                speaker_md = content
            elif output_type == 'markdown':
                transcript_md = content
        except Exception:
            continue

    parser = TimelineParser(job_id, input_id, filename)
    timeline = parser.parse(
        transcript_json=transcript_json,
        diarization_json=diarization_json,
        speaker_md=speaker_md,
        transcript_md=transcript_md
    )

    if review_state:
        timeline = apply_review_state(timeline, review_state)

    # Apply view mode: use raw or merged chunks
    if view_mode == 'raw' and hasattr(timeline, 'chunks_raw') and timeline.chunks_raw:
        timeline.chunks = timeline.chunks_raw
    elif hasattr(timeline, 'chunks_merged') and timeline.chunks_merged:
        timeline.chunks = timeline.chunks_merged

    return timeline, review_state


@app.route('/api/jobs/<job_id>/review/export-docx', methods=['GET'])
def api_export_review_docx(job_id):
    """Export review timeline as DOCX with timestamps."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    import io
    
    # Safe access helper for dict or object
    def val(obj, key, default=None):
        if obj is None:
            return default
        if isinstance(obj, dict):
            return obj.get(key, default)
        return getattr(obj, key, default)
    
    # Convert hex color to RGB tuple
    def hex_to_rgb(hex_color):
        hex_color = (hex_color or '').lstrip('#')
        if len(hex_color) != 6:
            return (107, 114, 128)  # fallback gray
        try:
            return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
        except ValueError:
            return (107, 114, 128)
    
    try:
        session_id = g.session_id
        manifest_path = session_store.job_manifest_path(session_id, job_id)
        manifest = session_store.read_json(manifest_path)
        
        if not manifest:
            return jsonify({'error': 'Job not found'}), 404
        
        input_id = request.args.get('inputId')
        inputs = manifest.get('inputs', [])
        if not inputs:
            return jsonify({'error': 'No inputs in job'}), 400
        
        target_input = inputs[0] if not input_id else next((i for i in inputs if i.get('uploadId') == input_id), inputs[0])
        input_id = target_input.get('uploadId')
        filename = target_input.get('originalFilename', target_input.get('filename', 'unknown'))
        safe_name = secure_filename(Path(filename).stem) or 'transcript'
        
        # Build timeline with review state (respect view mode)
        view_mode = request.args.get('viewMode', 'merged')
        include_highlights = request.args.get('includeHighlights', 'true').lower() == 'true'
        timeline, review_state = _build_review_timeline(session_id, job_id, input_id, filename, manifest, view_mode)
        
        # Get chunk edits for highlight info
        chunk_edits = review_state.get('chunkEdits', {}) if review_state else {}
        
        # Format time helper
        def fmt_time(seconds):
            m, s = divmod(int(seconds), 60)
            h, m = divmod(m, 60)
            if h > 0:
                return f"{h}:{m:02d}:{s:02d}"
            return f"{m}:{s:02d}"
        
        # Build highlight sections (group consecutive highlighted chunks)
        def get_highlight_sections():
            highlighted = []
            for chunk in (timeline.chunks or []):
                cid = val(chunk, 'chunk_id', '')
                if chunk_edits.get(cid, {}).get('highlighted', False):
                    highlighted.append(chunk)
            if not highlighted:
                return []
            
            # Sort by start time
            sorted_chunks = sorted(highlighted, key=lambda c: val(c, 'start', 0))
            sections = []
            current = {'start': val(sorted_chunks[0], 'start', 0), 'end': val(sorted_chunks[0], 'end', 0), 'count': 1}
            
            for i in range(1, len(sorted_chunks)):
                chunk = sorted_chunks[i]
                prev = sorted_chunks[i - 1]
                # Consecutive if within 1 second
                if val(chunk, 'start', 0) - val(prev, 'end', 0) < 1.0:
                    current['end'] = val(chunk, 'end', 0)
                    current['count'] += 1
                else:
                    sections.append(current)
                    current = {'start': val(chunk, 'start', 0), 'end': val(chunk, 'end', 0), 'count': 1}
            sections.append(current)
            return sections
        
        highlight_sections = get_highlight_sections() if include_highlights else []
        
        # Create DOCX document
        doc = Document()
        doc.add_heading(filename, 0)
        
        # Add highlight summary if there are highlights
        if highlight_sections:
            summary_para = doc.add_paragraph()
            summary_para.add_run(f"🔆 {len(highlight_sections)} Highlighted Section(s): ").bold = True
            times = [f"[{fmt_time(s['start'])}]" for s in highlight_sections]
            summary_para.add_run(", ".join(times))
            doc.add_paragraph()  # Spacer
        
        # Build speaker lookup (handles both dict and object)
        speakers = {}
        for s in (timeline.speakers or []):
            sid = val(s, 'id')
            if sid:
                speakers[sid] = s
        
        # Timestamped mode: [mm:ss] Speaker: text for each chunk
        for chunk in (timeline.chunks or []):
            chunk_id = val(chunk, 'chunk_id', '')
            sid = val(chunk, 'speaker_id')
            speaker = speakers.get(sid)
            label = val(speaker, 'label', sid or 'Unknown') or sid or 'Unknown'
            color = val(speaker, 'color', '#6b7280') or '#6b7280'
            start = val(chunk, 'start', 0)
            text = val(chunk, 'text', '') or ''
            timestamp = fmt_time(start)
            
            # Check if chunk is highlighted
            is_highlighted = chunk_edits.get(chunk_id, {}).get('highlighted', False) if include_highlights else False
            
            para = doc.add_paragraph()
            run = para.add_run(f"[{timestamp}] {label}: ")
            run.bold = True
            red, green, blue = hex_to_rgb(color)
            run.font.color.rgb = RGBColor(red, green, blue)
            
            # Add text with highlight if applicable
            text_run = para.add_run(text)
            if is_highlighted:
                from docx.shared import RGBColor as DocxRGBColor
                from docx.oxml.ns import nsdecls
                from docx.oxml import parse_xml
                # Apply yellow highlight
                text_run.font.highlight_color = 7  # WD_COLOR_INDEX.YELLOW = 7
        
        # Write to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return Response(
            buffer.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename="{safe_name}_review.docx"'}
        )
    except Exception as e:
        log_event('error', 'review_export_failed', jobId=job_id, exportType='docx', error=str(e))
        return jsonify({'error': 'Export failed', 'details': str(e)[:200]}), 500


@app.route('/api/jobs/<job_id>/export-all', methods=['GET'])
def api_export_all_formats(job_id):
    """Export all formats (btproj, md, docx, pdf, timeline.json) as a single zip."""
    session_id = g.session_id
    
    job_dir = session_store.job_dir(session_id, job_id)
    if not job_dir or not os.path.exists(job_dir):
        return jsonify({'error': 'Job not found'}), 404
    
    try:
        # Get base files from btproj export helper
        files, filename = _build_job_export_files(session_id, job_id)
        safe_name = secure_filename(Path(filename).stem) or 'export'
        
        # Get timeline and review state for additional exports
        timeline, review_state = _get_timeline_and_state_for_export(session_id, job_id)
        
        if timeline:
            # Generate markdown
            md_lines = []
            speakers = {s.id: s for s in (timeline.speakers or [])}
            for chunk in (timeline.chunks or []):
                sid = chunk.speaker_id
                speaker = speakers.get(sid)
                label = speaker.label if speaker else (sid or 'Unknown')
                start = chunk.start or 0
                mins, secs = divmod(int(start), 60)
                timestamp = f"{mins:02d}:{secs:02d}"
                text = chunk.text or ''
                md_lines.append(f"[{timestamp}] **{label}**: {text}")
            files[f'{safe_name}.md'] = '\n\n'.join(md_lines).encode('utf-8')
            
            # Generate DOCX
            try:
                from docx import Document
                from docx.shared import RGBColor
                
                def hex_to_rgb(hex_color):
                    hex_color = hex_color.lstrip('#')
                    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
                
                doc = Document()
                doc.add_heading(safe_name, 0)
                
                for chunk in (timeline.chunks or []):
                    sid = chunk.speaker_id
                    speaker = speakers.get(sid)
                    label = speaker.label if speaker else (sid or 'Unknown')
                    color = speaker.color if speaker else '#6b7280'
                    start = chunk.start or 0
                    mins, secs = divmod(int(start), 60)
                    timestamp = f"{mins:02d}:{secs:02d}"
                    text = chunk.text or ''
                    
                    para = doc.add_paragraph()
                    run = para.add_run(f"[{timestamp}] {label}: ")
                    run.bold = True
                    r, g, b = hex_to_rgb(color)
                    run.font.color.rgb = RGBColor(r, g, b)
                    para.add_run(text)
                
                docx_buffer = io.BytesIO()
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                files[f'{safe_name}.docx'] = docx_buffer.getvalue()
            except ImportError:
                pass  # python-docx not available
            
            # Generate PDF
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.platypus import SimpleDocTemplate, Paragraph
                from reportlab.lib.styles import getSampleStyleSheet
                
                pdf_buffer = io.BytesIO()
                doc = SimpleDocTemplate(pdf_buffer, pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                
                for chunk in (timeline.chunks or []):
                    sid = chunk.speaker_id
                    speaker = speakers.get(sid)
                    label = speaker.label if speaker else (sid or 'Unknown')
                    start = chunk.start or 0
                    mins, secs = divmod(int(start), 60)
                    timestamp = f"{mins:02d}:{secs:02d}"
                    text = chunk.text or ''
                    
                    story.append(Paragraph(f"[{timestamp}] <b>{label}</b>: {text}", styles['Normal']))
                
                if story:
                    doc.build(story)
                    pdf_buffer.seek(0)
                    files[f'{safe_name}.pdf'] = pdf_buffer.getvalue()
            except ImportError:
                pass  # reportlab not available
        
        # Create zip
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            for file_path, content in files.items():
                zf.writestr(f'{safe_name}_all/{file_path}', content)
        
        zip_buffer.seek(0)
        
        return send_file(
            zip_buffer,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f'{safe_name}_all_formats.zip'
        )
    except Exception as e:
        log_event('error', 'export_all_failed', jobId=job_id, error=str(e))
        return jsonify({'error': 'Export failed', 'details': str(e)[:200]}), 500


@app.route('/api/jobs/<job_id>/review/export-pdf', methods=['GET'])
def api_export_review_pdf(job_id):
    """Export review timeline as PDF with timestamps."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    import io
    
    # Safe access helper for dict or object
    def val(obj, key, default=None):
        if obj is None:
            return default
        if isinstance(obj, dict):
            return obj.get(key, default)
        return getattr(obj, key, default)
    
    try:
        session_id = g.session_id
        manifest_path = session_store.job_manifest_path(session_id, job_id)
        manifest = session_store.read_json(manifest_path)
        
        if not manifest:
            return jsonify({'error': 'Job not found'}), 404
        
        input_id = request.args.get('inputId')
        inputs = manifest.get('inputs', [])
        if not inputs:
            return jsonify({'error': 'No inputs in job'}), 400
        
        target_input = inputs[0] if not input_id else next((i for i in inputs if i.get('uploadId') == input_id), inputs[0])
        input_id = target_input.get('uploadId')
        filename = target_input.get('originalFilename', target_input.get('filename', 'unknown'))
        safe_name = secure_filename(Path(filename).stem) or 'transcript'
        
        # Build timeline with review state (respect view mode)
        view_mode = request.args.get('viewMode', 'merged')
        include_highlights = request.args.get('includeHighlights', 'true').lower() == 'true'
        timeline, review_state = _build_review_timeline(session_id, job_id, input_id, filename, manifest, view_mode)
        
        # Get chunk edits for highlight info
        chunk_edits = review_state.get('chunkEdits', {}) if review_state else {}
        
        # Format time helper
        def fmt_time(seconds):
            m, s = divmod(int(seconds), 60)
            h, m = divmod(m, 60)
            if h > 0:
                return f"{h}:{m:02d}:{s:02d}"
            return f"{m}:{s:02d}"
        
        # Build highlight sections (group consecutive highlighted chunks)
        def get_highlight_sections():
            highlighted = []
            for chunk in (timeline.chunks or []):
                cid = val(chunk, 'chunk_id', '')
                if chunk_edits.get(cid, {}).get('highlighted', False):
                    highlighted.append(chunk)
            if not highlighted:
                return []
            
            sorted_chunks = sorted(highlighted, key=lambda c: val(c, 'start', 0))
            sections = []
            current = {'start': val(sorted_chunks[0], 'start', 0), 'end': val(sorted_chunks[0], 'end', 0), 'count': 1}
            
            for i in range(1, len(sorted_chunks)):
                chunk = sorted_chunks[i]
                prev = sorted_chunks[i - 1]
                if val(chunk, 'start', 0) - val(prev, 'end', 0) < 1.0:
                    current['end'] = val(chunk, 'end', 0)
                    current['count'] += 1
                else:
                    sections.append(current)
                    current = {'start': val(chunk, 'start', 0), 'end': val(chunk, 'end', 0), 'count': 1}
            sections.append(current)
            return sections
        
        highlight_sections = get_highlight_sections() if include_highlights else []
        
        # Build speaker lookup (handles both dict and object)
        speakers = {}
        for s in (timeline.speakers or []):
            sid = val(s, 'id')
            if sid:
                speakers[sid] = s
        
        # Create PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
        
        styles = getSampleStyleSheet()
        title_style = styles['Heading1']
        body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10, leading=14, spaceAfter=8)
        highlight_style = ParagraphStyle('Highlight', parent=styles['Normal'], fontSize=10, leading=14, spaceAfter=8, backColor='#FFFF99')
        summary_style = ParagraphStyle('Summary', parent=styles['Normal'], fontSize=10, leading=14, spaceAfter=12, textColor='#B8860B')
        
        story = []
        story.append(Paragraph(filename, title_style))
        
        # Add highlight summary if there are highlights
        if highlight_sections:
            times = [f"[{fmt_time(s['start'])}]" for s in highlight_sections]
            summary_text = f"<b>🔆 {len(highlight_sections)} Highlighted Section(s):</b> {', '.join(times)}"
            story.append(Paragraph(summary_text, summary_style))
        
        story.append(Spacer(1, 0.25 * inch))
        
        def escape_xml(text):
            return (text or '').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        
        # Timestamped mode: [mm:ss] Speaker: text for each chunk
        for chunk in (timeline.chunks or []):
            chunk_id = val(chunk, 'chunk_id', '')
            sid = val(chunk, 'speaker_id')
            speaker = speakers.get(sid)
            label = val(speaker, 'label', sid or 'Unknown') or sid or 'Unknown'
            color = val(speaker, 'color', '#6b7280') or '#6b7280'
            start = val(chunk, 'start', 0)
            text = escape_xml(val(chunk, 'text', '') or '')
            timestamp = fmt_time(start)
            
            # Check if chunk is highlighted
            is_highlighted = chunk_edits.get(chunk_id, {}).get('highlighted', False) if include_highlights else False
            
            line = f'<font color="{color}"><b>[{timestamp}] {escape_xml(label)}:</b></font> {text}'
            style = highlight_style if is_highlighted else body_style
            story.append(Paragraph(line, style))
        
        doc.build(story)
        buffer.seek(0)
        
        return Response(
            buffer.getvalue(),
            mimetype='application/pdf',
            headers={'Content-Disposition': f'attachment; filename="{safe_name}_review.pdf"'}
        )
    except Exception as e:
        log_event('error', 'review_export_failed', jobId=job_id, exportType='pdf', error=str(e))
        return jsonify({'error': 'Export failed', 'details': str(e)[:200]}), 500


def _build_job_export_files(session_id: str, job_id: str, input_id: str = None) -> dict:
    """
    Build export files for a single job. Returns dict of {filename: content_bytes}.
    Raises exception on failure.
    """
    from review_timeline import TimelineParser, apply_review_state
    
    manifest_path = session_store.job_manifest_path(session_id, job_id)
    manifest = session_store.read_json(manifest_path)
    
    if not manifest:
        raise ValueError('Job not found')
    
    inputs = manifest.get('inputs', [])
    if not inputs:
        raise ValueError('No inputs in job')
    
    target_input = None
    if input_id:
        for inp in inputs:
            if inp.get('uploadId') == input_id:
                target_input = inp
                break
        if not target_input:
            raise ValueError('Input not found')
    else:
        target_input = inputs[0]
        input_id = target_input.get('uploadId')
    
    filename = target_input.get('originalFilename', target_input.get('filename', 'unknown'))
    
    files = {}
    
    # Manifest snapshot
    safe_manifest = {
        'jobId': manifest.get('jobId'),
        'createdAt': manifest.get('createdAt'),
        'status': manifest.get('status'),
        'options': manifest.get('options'),
        'inputs': [{'uploadId': i.get('uploadId'), 'filename': i.get('originalFilename', i.get('filename'))} 
                  for i in manifest.get('inputs', [])],
        'outputs': [{'id': o.get('id'), 'filename': o.get('filename'), 'type': o.get('type')} 
                   for o in manifest.get('outputs', []) if not o.get('error')],
        'exportedAt': datetime.now(timezone.utc).isoformat(),
        'buildCommit': os.environ.get('BUILD_COMMIT', 'unknown'),
    }
    files['manifest_snapshot.json'] = json.dumps(safe_manifest, indent=2).encode('utf-8')
    
    # Review state
    job_dir = session_store.job_dir(session_id, job_id)
    state_path = os.path.join(job_dir, 'review', 'review_state.json')
    review_state = session_store.read_json(state_path) or {}
    files['review_state.json'] = json.dumps(review_state, indent=2).encode('utf-8')
    
    # Outputs and timeline parsing
    outputs = manifest.get('outputs', [])
    transcript_json = None
    diarization_json = None
    speaker_md = None
    transcript_md = None
    
    for output in outputs:
        if output.get('forUploadId') != input_id and output.get('forUploadId'):
            continue
        
        output_type = output.get('type', '')
        output_path = output.get('path')
        output_filename = output.get('filename', '')
        
        if not output_path or not os.path.exists(output_path):
            continue
        
        try:
            with open(output_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            files[f'outputs/{output_filename}'] = content.encode('utf-8')
            
            if output_type == 'json':
                transcript_json = content
            elif output_type == 'diarization-json':
                diarization_json = content
            elif output_type == 'speaker-markdown':
                speaker_md = content
            elif output_type == 'markdown':
                transcript_md = content
        except Exception:
            continue
    
    # Generate timeline
    parser = TimelineParser(job_id, input_id, filename)
    timeline = parser.parse(
        transcript_json=transcript_json,
        diarization_json=diarization_json,
        speaker_md=speaker_md,
        transcript_md=transcript_md
    )
    if review_state:
        timeline = apply_review_state(timeline, review_state)
    
    files['timeline.json'] = timeline.to_json().encode('utf-8')
    
    return files, filename


# Session export job storage (in-memory for simplicity, keyed by export_id)
_session_exports = {}
_session_exports_lock = threading.Lock()


def _cleanup_old_exports():
    """Remove exports older than 24 hours."""
    now = datetime.now(timezone.utc)
    with _session_exports_lock:
        expired = [
            eid for eid, exp in _session_exports.items()
            if (now - datetime.fromisoformat(exp.get('createdAt', now.isoformat()))).total_seconds() > 86400
        ]
        for eid in expired:
            exp = _session_exports.pop(eid, None)
            if exp and exp.get('zipPath') and os.path.exists(exp['zipPath']):
                try:
                    os.remove(exp['zipPath'])
                except Exception:
                    pass


def _run_session_export(export_id: str, session_id: str):
    """Background worker for session export."""
    try:
        with _session_exports_lock:
            _session_exports[export_id]['status'] = 'running'
        
        jobs_dir = os.path.join(session_store.session_dir(session_id), 'jobs')
        job_ids = [d for d in os.listdir(jobs_dir) if os.path.isdir(os.path.join(jobs_dir, d))]
        jobs_total = len(job_ids)
        
        with _session_exports_lock:
            _session_exports[export_id]['progress'] = {'jobsDone': 0, 'jobsTotal': jobs_total, 'currentJob': ''}
        
        # Create temp file for zip
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        root_folder = f'session_export_{timestamp}'
        zip_path = os.path.join(tempfile.gettempdir(), f'session_export_{export_id}.zip')
        
        job_count = 0
        files_total = 0
        with_review_state = 0
        failed_jobs = []
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            # Export session.json
            session_meta = {
                'sessionId': session_id,
                'exportedAt': datetime.now(timezone.utc).isoformat(),
                'version': '1.0',
                'buildCommit': os.environ.get('BUILD_COMMIT', 'unknown'),
            }
            zf.writestr(f'{root_folder}/session.json', json.dumps(session_meta, indent=2).encode('utf-8'))
            
            for idx, job_id in enumerate(job_ids):
                with _session_exports_lock:
                    _session_exports[export_id]['progress'] = {
                        'jobsDone': idx,
                        'jobsTotal': jobs_total,
                        'currentJob': job_id[:8]
                    }
                
                job_folder = f'{root_folder}/jobs/{job_id}'
                try:
                    manifest_path = session_store.job_manifest_path(session_id, job_id)
                    manifest = session_store.read_json(manifest_path)
                    if manifest:
                        zf.writestr(f'{job_folder}/job.json', json.dumps(manifest, indent=2).encode('utf-8'))
                    
                    job_dir = session_store.job_dir(session_id, job_id)
                    state_path = os.path.join(job_dir, 'review', 'review_state.json')
                    review_state = session_store.read_json(state_path)
                    if review_state:
                        zf.writestr(f'{job_folder}/review_state.json', json.dumps(review_state, indent=2).encode('utf-8'))
                        with_review_state += 1
                    
                    # Export all files from review directory
                    review_dir = os.path.join(job_dir, 'review')
                    if os.path.isdir(review_dir):
                        for fname in os.listdir(review_dir):
                            fpath = os.path.join(review_dir, fname)
                            if os.path.isfile(fpath):
                                try:
                                    with open(fpath, 'rb') as f:
                                        zf.writestr(f'{job_folder}/review/{fname}', f.read())
                                        files_total += 1
                                except Exception:
                                    pass

                    # Generate and export timelines for each input file
                    # Each input gets its own subdirectory under review/
                    if manifest:
                        inputs = manifest.get('inputs', [])
                        for input_info in inputs:
                            input_id = input_info.get('uploadId', '')
                            if not input_id:
                                continue
                            input_filename = input_info.get('originalFilename', input_info.get('filename', 'unknown'))

                            try:
                                # Build and export merged timeline
                                timeline_merged, review_state_for_input = _build_review_timeline(
                                    session_id, job_id, input_id, input_filename, manifest, 'merged'
                                )
                                if timeline_merged:
                                    timeline_path = f'{job_folder}/review/{input_id}/timeline_merged.json'
                                    zf.writestr(timeline_path, timeline_merged.to_json().encode('utf-8'))
                                    files_total += 1

                                # Build and export raw timeline
                                timeline_raw, _ = _build_review_timeline(
                                    session_id, job_id, input_id, input_filename, manifest, 'raw'
                                )
                                if timeline_raw:
                                    timeline_path = f'{job_folder}/review/{input_id}/timeline_raw.json'
                                    zf.writestr(timeline_path, timeline_raw.to_json().encode('utf-8'))
                                    files_total += 1

                                # Export input metadata for easier identification
                                input_meta = {
                                    'inputId': input_id,
                                    'filename': input_filename,
                                    'exportedAt': datetime.now(timezone.utc).isoformat(),
                                }
                                zf.writestr(f'{job_folder}/review/{input_id}/input.json',
                                           json.dumps(input_meta, indent=2).encode('utf-8'))
                                files_total += 1

                            except Exception as timeline_err:
                                # Log but don't fail the whole job export
                                pass

                    # Export all files from outputs directory (scan directly, not just manifest)
                    outputs_dir = os.path.join(job_dir, 'outputs')
                    if os.path.isdir(outputs_dir):
                        for fname in os.listdir(outputs_dir):
                            fpath = os.path.join(outputs_dir, fname)
                            if os.path.isfile(fpath):
                                try:
                                    with open(fpath, 'rb') as f:
                                        zf.writestr(f'{job_folder}/outputs/{fname}', f.read())
                                        files_total += 1
                                except Exception:
                                    pass

                    job_count += 1
                        
                except Exception as e:
                    failed_jobs.append({'jobId': job_id, 'reason': str(e)[:200]})
                    error_content = f'Export failed: {str(e)[:200]}'
                    zf.writestr(f'{job_folder}/error.txt', error_content.encode('utf-8'))
            
            summary = {
                'exportedAt': datetime.now(timezone.utc).isoformat(),
                'sessionId': session_id,
                'jobCount': job_count,
                'filesTotal': files_total,
                'withReviewState': with_review_state,
                'failedJobs': failed_jobs,
            }
            zf.writestr(f'{root_folder}/session_summary.json', json.dumps(summary, indent=2).encode('utf-8'))
        
        with _session_exports_lock:
            _session_exports[export_id].update({
                'status': 'complete',
                'progress': {'jobsDone': jobs_total, 'jobsTotal': jobs_total, 'currentJob': ''},
                'zipPath': zip_path,
                'downloadFilename': f'session-export-{timestamp}.zip',
                'summary': summary,
            })
            
    except Exception as e:
        with _session_exports_lock:
            _session_exports[export_id].update({
                'status': 'failed',
                'error': str(e)[:500],
            })


@app.route('/api/session/export', methods=['POST'])
def api_export_session():
    """Start async session export job. Returns exportId for polling."""
    session_id = g.session_id
    
    # Cleanup old exports
    _cleanup_old_exports()
    
    # Check if jobs exist
    jobs_dir = os.path.join(session_store.session_dir(session_id), 'jobs')
    if not jobs_dir or not os.path.exists(jobs_dir):
        return jsonify({'error': 'No jobs found'}), 404
    
    job_ids = [d for d in os.listdir(jobs_dir) if os.path.isdir(os.path.join(jobs_dir, d))]
    if not job_ids:
        return jsonify({'error': 'No jobs found'}), 404
    
    # Create export job
    export_id = f'exp_{secrets.token_hex(8)}'
    with _session_exports_lock:
        _session_exports[export_id] = {
            'exportId': export_id,
            'sessionId': session_id,
            'status': 'queued',
            'progress': {'jobsDone': 0, 'jobsTotal': len(job_ids), 'currentJob': ''},
            'createdAt': datetime.now(timezone.utc).isoformat(),
        }
    
    # Start background thread
    thread = threading.Thread(target=_run_session_export, args=(export_id, session_id), daemon=True)
    thread.start()
    
    return jsonify({'exportId': export_id})

_CSRF_PROTECTED_ENDPOINTS.add('api_export_session')


@app.route('/api/session/export/<export_id>', methods=['GET'])
def api_get_export_status(export_id):
    """Get status of an export job."""
    with _session_exports_lock:
        export = _session_exports.get(export_id)
    
    if not export:
        return jsonify({'error': 'Export not found'}), 404
    
    # Verify session ownership
    if export.get('sessionId') != g.session_id:
        return jsonify({'error': 'Export not found'}), 404
    
    result = {
        'exportId': export['exportId'],
        'status': export['status'],
        'progress': export.get('progress', {}),
    }
    
    if export['status'] == 'complete':
        result['downloadUrl'] = f'/api/session/export/{export_id}/download'
        result['summary'] = export.get('summary', {})
    elif export['status'] == 'failed':
        result['error'] = export.get('error', 'Unknown error')
    
    return jsonify(result)


@app.route('/api/session/export/<export_id>/download', methods=['GET'])
def api_download_export(export_id):
    """Download completed export zip."""
    with _session_exports_lock:
        export = _session_exports.get(export_id)
    
    if not export:
        return jsonify({'error': 'Export not found'}), 404
    
    if export.get('sessionId') != g.session_id:
        return jsonify({'error': 'Export not found'}), 404
    
    if export['status'] != 'complete':
        return jsonify({'error': 'Export not ready'}), 400
    
    zip_path = export.get('zipPath')
    if not zip_path or not os.path.exists(zip_path):
        return jsonify({'error': 'Export file not found'}), 404
    
    return send_file(
        zip_path,
        mimetype='application/zip',
        as_attachment=True,
        download_name=export.get('downloadFilename', 'session-export.zip')
    )


@app.route('/api/session/import', methods=['POST'])
def api_import_session():
    """
    Import a session export zip, restoring all jobs with their editor state.
    
    Accepts multipart form with 'file' containing the session export zip.
    Returns summary of imported jobs.
    """
    session_id = g.session_id
    
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': 'No file selected'}), 400
    
    MAX_SESSION_SIZE = 500 * 1024 * 1024  # 500MB limit
    
    # Read file into memory
    file_content = file.read()
    if len(file_content) > MAX_SESSION_SIZE:
        return jsonify({'error': f'File too large (max {MAX_SESSION_SIZE // 1024 // 1024}MB)'}), 400
    
    imported = []
    failed = []
    
    try:
        with zipfile.ZipFile(io.BytesIO(file_content), 'r') as zf:
            # Find the root folder (session_export_TIMESTAMP or similar)
            namelist = zf.namelist()
            if not namelist:
                return jsonify({'error': 'Empty zip file'}), 400
            
            # Detect root folder
            root_folder = None
            for name in namelist:
                parts = name.split('/')
                if len(parts) >= 1 and parts[0]:
                    root_folder = parts[0]
                    break
            
            if not root_folder:
                return jsonify({'error': 'Invalid session export structure'}), 400
            
            # Find all job folders
            job_folders = set()
            for name in namelist:
                if name.startswith(f'{root_folder}/jobs/') and '/job.json' in name:
                    # Extract job_id from path like "root/jobs/job_id/job.json"
                    parts = name.split('/')
                    if len(parts) >= 4:
                        job_folders.add(parts[2])
            
            if not job_folders:
                return jsonify({'error': 'No jobs found in session export'}), 400
            
            for orig_job_id in job_folders:
                job_prefix = f'{root_folder}/jobs/{orig_job_id}/'
                
                try:
                    # Read job.json
                    job_json_path = f'{job_prefix}job.json'
                    if job_json_path not in namelist:
                        failed.append({'jobId': orig_job_id, 'reason': 'Missing job.json'})
                        continue
                    
                    job_data = json.loads(zf.read(job_json_path).decode('utf-8'))
                    
                    # Check if job ID already exists - generate new one if so
                    new_job_id = orig_job_id
                    existing_manifest = session_store.job_manifest_path(session_id, orig_job_id)
                    if os.path.exists(existing_manifest):
                        new_job_id = str(uuid.uuid4())
                        job_data['jobId'] = new_job_id
                        job_data['importedFrom'] = orig_job_id
                    
                    # Create job directory
                    job_dir = session_store.job_dir(session_id, new_job_id)
                    os.makedirs(job_dir, exist_ok=True)
                    os.makedirs(os.path.join(job_dir, 'review'), exist_ok=True)
                    os.makedirs(os.path.join(job_dir, 'outputs'), exist_ok=True)
                    
                    # Update output paths in manifest to point to new location
                    outputs = job_data.get('outputs', [])
                    for output in outputs:
                        old_path = output.get('path', '')
                        filename = output.get('filename', '')
                        if filename:
                            new_path = os.path.join(job_dir, 'outputs', filename)
                            output['path'] = new_path
                    
                    # Write job manifest
                    manifest_path = session_store.job_manifest_path(session_id, new_job_id)
                    session_store.write_json(manifest_path, job_data)
                    
                    # Restore review_state.json if present (check both old and new paths)
                    review_state_path = f'{job_prefix}review_state.json'
                    review_state_path_new = f'{job_prefix}review/review_state.json'
                    if review_state_path in namelist:
                        review_state = json.loads(zf.read(review_state_path).decode('utf-8'))
                        state_dest = os.path.join(job_dir, 'review', 'review_state.json')
                        session_store.write_json(state_dest, review_state)
                    elif review_state_path_new in namelist:
                        review_state = json.loads(zf.read(review_state_path_new).decode('utf-8'))
                        state_dest = os.path.join(job_dir, 'review', 'review_state.json')
                        session_store.write_json(state_dest, review_state)

                    # Restore timeline.json if present (check both old and new paths)
                    timeline_path = f'{job_prefix}timeline.json'
                    timeline_path_new = f'{job_prefix}review/timeline.json'
                    if timeline_path in namelist:
                        timeline_content = zf.read(timeline_path).decode('utf-8')
                        timeline_dest = os.path.join(job_dir, 'review', 'timeline.json')
                        with open(timeline_dest, 'w', encoding='utf-8') as f:
                            f.write(timeline_content)
                    elif timeline_path_new in namelist:
                        timeline_content = zf.read(timeline_path_new).decode('utf-8')
                        timeline_dest = os.path.join(job_dir, 'review', 'timeline.json')
                        with open(timeline_dest, 'w', encoding='utf-8') as f:
                            f.write(timeline_content)

                    # Restore any other review files from review/ subdirectory (including nested dirs)
                    review_prefix = f'{job_prefix}review/'
                    for name in namelist:
                        if name.startswith(review_prefix) and not name.endswith('/'):
                            # Get relative path within review/
                            rel_path = name[len(review_prefix):]
                            filename = rel_path.split('/')[-1]
                            # Skip files we already handled explicitly at root level
                            if rel_path in ('review_state.json', 'timeline.json'):
                                continue
                            if filename:
                                review_dest = os.path.join(job_dir, 'review', rel_path)
                                # Create subdirectories if needed (e.g., review/{input_id}/)
                                os.makedirs(os.path.dirname(review_dest), exist_ok=True)
                                with open(review_dest, 'wb') as f:
                                    f.write(zf.read(name))

                    # Restore outputs
                    for name in namelist:
                        if name.startswith(f'{job_prefix}outputs/') and not name.endswith('/'):
                            filename = name.split('/')[-1]
                            if filename:
                                output_dest = os.path.join(job_dir, 'outputs', filename)
                                with open(output_dest, 'wb') as f:
                                    f.write(zf.read(name))
                    
                    imported.append({
                        'originalJobId': orig_job_id,
                        'newJobId': new_job_id,
                        'remapped': new_job_id != orig_job_id
                    })
                    
                except Exception as e:
                    failed.append({'jobId': orig_job_id, 'reason': str(e)[:200]})
    
    except zipfile.BadZipFile:
        return jsonify({'error': 'Invalid zip file'}), 400
    except Exception as e:
        return jsonify({'error': f'Import failed: {str(e)[:200]}'}), 500
    
    return jsonify({
        'imported': len(imported),
        'importedJobs': imported,
        'failed': failed
    })

_CSRF_PROTECTED_ENDPOINTS.add('api_import_session')


# =============================================================================
# SESSION SHARING ENDPOINTS
# =============================================================================

@app.route('/api/session/share', methods=['POST'])
def api_share_session():
    """
    Create a share token for the current session.
    
    Body: { "mode": "read"|"edit", "password": "optional" }
    Returns: { "token": "abc123" }
    """
    import hashlib
    
    session_id = g.session_id
    data = request.get_json() or {}
    
    mode = data.get('mode', 'read')
    if mode not in ('read', 'edit'):
        return jsonify({'error': 'Invalid mode. Must be "read" or "edit"'}), 400
    
    password = data.get('password')
    password_hash = None
    if password:
        # Simple hash (not bcrypt to avoid dependency, but sufficient for share tokens)
        password_hash = hashlib.sha256(password.encode()).hexdigest()
    
    # Generate token
    token = secrets.token_urlsafe(16)
    
    # Store share
    session_store.add_share_token(session_id, token, mode, password_hash)
    
    return jsonify({'token': token, 'mode': mode})

_CSRF_PROTECTED_ENDPOINTS.add('api_share_session')


@app.route('/api/session/share', methods=['GET'])
def api_list_shares():
    """List all share tokens for the current session."""
    session_id = g.session_id
    shares = session_store.get_shared_sessions(session_id)
    
    # Don't expose password hashes
    safe_shares = []
    for s in shares:
        safe_shares.append({
            'token': s.get('token'),
            'mode': s.get('mode'),
            'hasPassword': bool(s.get('passwordHash')),
            'createdAt': s.get('createdAt')
        })
    
    return jsonify({'shares': safe_shares})


@app.route('/api/session/share/<token>', methods=['DELETE'])
def api_revoke_share(token):
    """Revoke a share token."""
    session_id = g.session_id
    
    if session_store.revoke_share_token(session_id, token):
        return jsonify({'message': 'Token revoked'})
    else:
        return jsonify({'error': 'Token not found'}), 404

_CSRF_PROTECTED_ENDPOINTS.add('api_revoke_share')


@app.route('/api/session/join', methods=['POST'])
def api_join_session():
    """
    Join a shared session using a token.
    
    Body: { "token": "abc123", "password": "optional" }
    Returns: { "sessionId": "...", "mode": "read"|"edit" }
    
    Sets session cookie to the shared session.
    """
    import hashlib
    
    data = request.get_json() or {}
    token = data.get('token', '').strip()
    password = data.get('password', '')
    
    if not token:
        return jsonify({'error': 'Token required'}), 400
    
    # Find share
    share = session_store.find_share_by_token(token)
    if not share:
        return jsonify({'error': 'Invalid token'}), 404
    
    # Verify password if required
    if share.get('passwordHash'):
        if not password:
            return jsonify({'error': 'Password required'}), 401
        provided_hash = hashlib.sha256(password.encode()).hexdigest()
        if provided_hash != share['passwordHash']:
            return jsonify({'error': 'Invalid password'}), 401
    
    # Set the shared session ID and mode
    shared_session_id = share['sessionId']
    mode = share['mode']
    
    # Store mode in a special way - we'll use a prefixed session ID
    # Format: "shared:<mode>:<original_session_id>"
    # This allows the middleware to detect shared sessions and enforce read-only
    joined_session_id = f"shared:{mode}:{shared_session_id}"
    
    # Mark that we need to set a new cookie
    g._new_session_id = joined_session_id
    g._joined_share_mode = mode
    
    return jsonify({
        'message': 'Joined session',
        'mode': mode,
        'sessionId': shared_session_id[:8] + '...'  # Partial for privacy
    })


@app.route('/api/session/info', methods=['GET'])
def api_session_info():
    """Get info about current session including share mode."""
    raw_session_id = request.cookies.get(session_store.COOKIE_NAME, '')
    
    is_shared = raw_session_id.startswith('shared:')
    mode = 'edit'  # Default for own session
    
    if is_shared:
        parts = raw_session_id.split(':', 2)
        if len(parts) >= 2:
            mode = parts[1]
    
    return jsonify({
        'isShared': is_shared,
        'mode': mode,
        'readOnly': mode == 'read'
    })


@app.route('/api/session/leave', methods=['POST'])
def api_leave_session():
    """Leave a shared session and get a fresh session."""
    # Generate new session ID
    new_session_id = session_store.new_id(32)
    g._new_session_id = new_session_id
    
    return jsonify({'message': 'Left shared session'})

_CSRF_PROTECTED_ENDPOINTS.add('api_leave_session')


@app.route('/api/jobs/bulk-export', methods=['POST'])
def api_bulk_export_jobs():
    """
    Export multiple jobs as a single zip containing subfolders per job.
    
    Request JSON:
    {
        "jobIds": ["job1", "job2", ...],
        "inputIdMode": "first"  // optional, default "first"
    }
    """
    session_id = g.session_id
    
    data = request.json or {}
    job_ids = data.get('jobIds', [])
    
    if not job_ids:
        return jsonify({'error': 'No job IDs provided'}), 400
    
    if len(job_ids) > 100:
        return jsonify({'error': 'Maximum 100 jobs per bulk export'}), 400
    
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for job_id in job_ids:
            try:
                files, filename = _build_job_export_files(session_id, job_id)
                safe_name = secure_filename(Path(filename).stem) or 'project'
                folder_name = f'{job_id}__{safe_name}'
                
                for file_path, content in files.items():
                    zf.writestr(f'bulk_export/{folder_name}/{file_path}', content)
                    
            except Exception as e:
                # Include error file for failed jobs
                error_content = f'Export failed: {str(e)[:200]}'
                zf.writestr(f'bulk_export/{job_id}__error/error.txt', error_content.encode('utf-8'))
    
    zip_buffer.seek(0)
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
    return send_file(
        zip_buffer,
        mimetype='application/zip',
        as_attachment=True,
        download_name=f'bulk-review-export-{timestamp}.zip'
    )

_CSRF_PROTECTED_ENDPOINTS.add('api_bulk_export_jobs')


@app.route('/api/review/import-bulk', methods=['POST'])
def api_bulk_import_projects():
    """
    Import multiple .btproj.zip files at once.
    
    Multipart form with 'files' containing multiple zip files.
    """
    session_id = g.session_id
    
    if 'files' not in request.files:
        return jsonify({'error': 'No files provided'}), 400
    
    files = request.files.getlist('files')
    
    if not files:
        return jsonify({'error': 'No files selected'}), 400
    
    if len(files) > 50:
        return jsonify({'error': 'Maximum 50 files per bulk import'}), 400
    
    MAX_PROJECT_SIZE = 50 * 1024 * 1024
    imported = []
    failed = []
    
    for file in files:
        filename = file.filename or 'unknown.zip'
        
        if not filename.endswith('.btproj.zip'):
            failed.append({'filename': filename, 'error': 'File must be a .btproj.zip'})
            continue
        
        try:
            zip_bytes = file.read()
            if len(zip_bytes) > MAX_PROJECT_SIZE:
                failed.append({'filename': filename, 'error': f'File too large (max {MAX_PROJECT_SIZE // (1024*1024)}MB)'})
                continue
            
            zip_data = io.BytesIO(zip_bytes)
            with zipfile.ZipFile(zip_data, 'r') as zf:
                names = zf.namelist()
                
                if 'manifest_snapshot.json' not in names:
                    failed.append({'filename': filename, 'error': 'Invalid project: missing manifest_snapshot.json'})
                    continue
                
                # Check for zip slip attacks
                for name in names:
                    if name.startswith('/') or '..' in name:
                        failed.append({'filename': filename, 'error': 'Invalid project: path traversal detected'})
                        continue
                
                manifest_data = json.loads(zf.read('manifest_snapshot.json'))
                
                project_id = session_store.new_id(16)
                project_dir = os.path.join(
                    session_store.session_dir(session_id),
                    'projects',
                    project_id
                )
                os.makedirs(project_dir, exist_ok=True)
                
                outputs_dir = os.path.join(project_dir, 'outputs')
                os.makedirs(outputs_dir, exist_ok=True)
                
                imported_outputs = []
                for name in names:
                    if name.startswith('outputs/') and not name.endswith('/'):
                        out_filename = os.path.basename(name)
                        if out_filename:
                            safe_name = secure_filename(out_filename)
                            output_path = os.path.join(outputs_dir, safe_name)
                            with open(output_path, 'wb') as f:
                                f.write(zf.read(name))
                            
                            output_type = 'text'
                            if safe_name.endswith('.json'):
                                if 'diarization' in safe_name:
                                    output_type = 'diarization-json'
                                else:
                                    output_type = 'json'
                            elif safe_name.endswith('.md'):
                                if 'speaker' in safe_name:
                                    output_type = 'speaker-markdown'
                                else:
                                    output_type = 'markdown'
                            
                            imported_outputs.append({
                                'id': session_store.new_id(8),
                                'filename': safe_name,
                                'path': output_path,
                                'type': output_type,
                                'sizeBytes': os.path.getsize(output_path)
                            })
                
                review_dir = os.path.join(project_dir, 'review')
                os.makedirs(review_dir, exist_ok=True)
                
                if 'review_state.json' in names:
                    with open(os.path.join(review_dir, 'review_state.json'), 'wb') as f:
                        f.write(zf.read('review_state.json'))
                
                if 'timeline.json' in names:
                    with open(os.path.join(review_dir, 'timeline.json'), 'wb') as f:
                        f.write(zf.read('timeline.json'))
                
                project_manifest = {
                    'projectId': project_id,
                    'jobId': project_id,
                    'sessionId': session_id,
                    'isImportedProject': True,
                    'importedFrom': manifest_data.get('jobId'),
                    'createdAt': datetime.now(timezone.utc).isoformat(),
                    'status': 'complete',
                    'inputs': manifest_data.get('inputs', []),
                    'outputs': imported_outputs,
                    'options': manifest_data.get('options', {}),
                }
                
                manifest_path = os.path.join(project_dir, 'manifest.json')
                session_store.atomic_write_json(manifest_path, project_manifest)
                
                projects_list_path = os.path.join(
                    session_store.session_dir(session_id),
                    'projects.json'
                )
                projects_list = session_store.read_json(projects_list_path) or {'projects': []}
                projects_list['projects'].insert(0, {
                    'projectId': project_id,
                    'name': manifest_data.get('inputs', [{}])[0].get('filename', 'Imported Project'),
                    'importedAt': datetime.now(timezone.utc).isoformat(),
                })
                session_store.atomic_write_json(projects_list_path, projects_list)
                
                imported.append({
                    'filename': filename,
                    'projectId': project_id,
                    'outputs': len(imported_outputs)
                })
                
        except zipfile.BadZipFile:
            failed.append({'filename': filename, 'error': 'Invalid zip file'})
        except json.JSONDecodeError:
            failed.append({'filename': filename, 'error': 'Invalid project: corrupt JSON'})
        except Exception as e:
            failed.append({'filename': filename, 'error': f'Import failed: {str(e)[:100]}'})
    
    return jsonify({
        'imported': imported,
        'failed': failed
    })

_CSRF_PROTECTED_ENDPOINTS.add('api_bulk_import_projects')


@app.route('/api/review/import', methods=['POST'])
def api_import_review_project():
    """
    Import a review project from a .btproj.zip bundle.
    
    Creates a new project in the current session that can be reviewed.
    """
    session_id = g.session_id
    
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': 'No file selected'}), 400
    
    if not file.filename.endswith('.btproj.zip'):
        return jsonify({'error': 'File must be a .btproj.zip'}), 400
    
    # Read zip into memory with size limit (50MB max for project bundles)
    MAX_PROJECT_SIZE = 50 * 1024 * 1024
    try:
        zip_bytes = file.read()
        if len(zip_bytes) > MAX_PROJECT_SIZE:
            return jsonify({'error': f'Project too large. Maximum size is {MAX_PROJECT_SIZE // (1024*1024)}MB'}), 400
        zip_data = io.BytesIO(zip_bytes)
        with zipfile.ZipFile(zip_data, 'r') as zf:
            # Validate structure
            names = zf.namelist()
            
            if 'manifest_snapshot.json' not in names:
                return jsonify({'error': 'Invalid project: missing manifest_snapshot.json'}), 400
            
            # Check for zip slip attacks
            for name in names:
                if name.startswith('/') or '..' in name:
                    return jsonify({'error': 'Invalid project: path traversal detected'}), 400
            
            # Read manifest
            manifest_data = json.loads(zf.read('manifest_snapshot.json'))
            
            # Create new project
            project_id = session_store.new_id(16)
            project_dir = os.path.join(
                session_store.session_dir(session_id),
                'projects',
                project_id
            )
            os.makedirs(project_dir, exist_ok=True)
            
            # Extract outputs
            outputs_dir = os.path.join(project_dir, 'outputs')
            os.makedirs(outputs_dir, exist_ok=True)
            
            imported_outputs = []
            for name in names:
                if name.startswith('outputs/') and not name.endswith('/'):
                    filename = os.path.basename(name)
                    if filename:
                        safe_name = secure_filename(filename)
                        output_path = os.path.join(outputs_dir, safe_name)
                        with open(output_path, 'wb') as f:
                            f.write(zf.read(name))
                        
                        # Determine type from filename
                        output_type = 'text'
                        if safe_name.endswith('.json'):
                            if 'diarization' in safe_name:
                                output_type = 'diarization-json'
                            else:
                                output_type = 'json'
                        elif safe_name.endswith('.md'):
                            if 'speaker' in safe_name:
                                output_type = 'speaker-markdown'
                            else:
                                output_type = 'markdown'
                        
                        imported_outputs.append({
                            'id': session_store.new_id(8),
                            'filename': safe_name,
                            'path': output_path,
                            'type': output_type,
                            'sizeBytes': os.path.getsize(output_path)
                        })
            
            # Extract review state
            review_dir = os.path.join(project_dir, 'review')
            os.makedirs(review_dir, exist_ok=True)
            
            if 'review_state.json' in names:
                review_state = zf.read('review_state.json')
                with open(os.path.join(review_dir, 'review_state.json'), 'wb') as f:
                    f.write(review_state)
            
            if 'timeline.json' in names:
                timeline = zf.read('timeline.json')
                with open(os.path.join(review_dir, 'timeline.json'), 'wb') as f:
                    f.write(timeline)
            
            # Create project manifest
            project_manifest = {
                'projectId': project_id,
                'jobId': project_id,  # Use same ID for compatibility
                'sessionId': session_id,
                'isImportedProject': True,
                'importedFrom': manifest_data.get('jobId'),
                'createdAt': datetime.now(timezone.utc).isoformat(),
                'status': 'complete',
                'inputs': manifest_data.get('inputs', []),
                'outputs': imported_outputs,
                'options': manifest_data.get('options', {}),
            }
            
            manifest_path = os.path.join(project_dir, 'manifest.json')
            session_store.atomic_write_json(manifest_path, project_manifest)
            
            # Register in projects list
            projects_list_path = os.path.join(
                session_store.session_dir(session_id),
                'projects.json'
            )
            projects_list = session_store.read_json(projects_list_path) or {'projects': []}
            projects_list['projects'].insert(0, {
                'projectId': project_id,
                'name': manifest_data.get('inputs', [{}])[0].get('filename', 'Imported Project'),
                'importedAt': datetime.now(timezone.utc).isoformat(),
            })
            session_store.atomic_write_json(projects_list_path, projects_list)
            
            return jsonify({
                'projectId': project_id,
                'message': 'Project imported successfully',
                'outputs': len(imported_outputs)
            })
            
    except zipfile.BadZipFile:
        return jsonify({'error': 'Invalid zip file'}), 400
    except json.JSONDecodeError:
        return jsonify({'error': 'Invalid project: corrupt JSON'}), 400
    except Exception as e:
        return jsonify({'error': f'Import failed: {str(e)[:100]}'}), 500

_CSRF_PROTECTED_ENDPOINTS.add('api_import_review_project')


@app.route('/api/projects', methods=['GET'])
def api_list_projects():
    """List imported projects for current session."""
    session_id = g.session_id
    
    projects_list_path = os.path.join(
        session_store.session_dir(session_id),
        'projects.json'
    )
    projects_list = session_store.read_json(projects_list_path) or {'projects': []}
    
    # Enrich with manifest data
    enriched = []
    for proj in projects_list.get('projects', []):
        project_id = proj.get('projectId')
        manifest_path = os.path.join(
            session_store.session_dir(session_id),
            'projects',
            project_id,
            'manifest.json'
        )
        manifest = session_store.read_json(manifest_path)
        if manifest:
            enriched.append({
                'projectId': project_id,
                'name': proj.get('name'),
                'importedAt': proj.get('importedAt'),
                'status': manifest.get('status'),
                'outputs': len(manifest.get('outputs', [])),
            })
    
    return jsonify({'projects': enriched})


# =============================================================================
# REMOTE WORKER ENDPOINTS
# These endpoints are called by the GPU worker, not by browser clients.
# They use Bearer token auth instead of session cookies.
# =============================================================================

def require_worker_auth(f):
    """Decorator to require worker token authentication."""
    from functools import wraps
    @wraps(f)
    def decorated(*args, **kwargs):
        # Support both REMOTE_WORKER_TOKEN and WORKER_TOKEN (same as remote_worker.py)
        worker_token = os.environ.get('REMOTE_WORKER_TOKEN') or os.environ.get('WORKER_TOKEN', '')
        
        # Also check saved config if env vars not set
        if not worker_token:
            try:
                from config_store import get_remote_worker_token
                worker_token = get_remote_worker_token()
            except ImportError:
                pass
        
        if not worker_token:
            return jsonify({'error': 'Worker endpoints not configured (no WORKER_TOKEN)'}), 503
        
        auth_header = request.headers.get('Authorization', '')
        if not auth_header.startswith('Bearer '):
            return jsonify({'error': 'Missing Authorization header'}), 401
        
        token = auth_header[7:]
        if not secrets.compare_digest(token, worker_token):
            return jsonify({'error': 'Invalid worker token'}), 403
        
        return f(*args, **kwargs)
    return decorated


@app.route('/api/jobs/<job_id>/inputs/<input_id>', methods=['GET'])
def api_get_job_input(job_id, input_id):
    """
    Download input file for a job. Used by remote worker in pull mode.
    
    Requires signed URL with valid signature and expiry.
    Query params:
    - expires: Unix timestamp when URL expires
    - sig: HMAC signature
    """
    from remote_worker import verify_signed_url
    
    # Verify signed URL
    expires = request.args.get('expires', '')
    sig = request.args.get('sig', '')
    secret = os.environ.get('SECRET_KEY', 'default-secret-key')
    
    if not verify_signed_url(job_id, input_id, expires, sig, secret):
        return jsonify({'error': 'Invalid or expired signature'}), 403
    
    # Find the job - check all sessions (worker doesn't have session context)
    # This is safe because the signed URL validates the job_id
    job_dir = None
    manifest = None
    
    # In server mode, we need to find which session owns this job
    if session_store.is_server_mode():
        sessions_root = Path(session_store.data_root()) / 'sessions'
        if sessions_root.exists():
            for session_dir in sessions_root.iterdir():
                if session_dir.is_dir():
                    potential_job_dir = session_dir / 'jobs' / job_id
                    manifest_path = potential_job_dir / 'job.json'
                    if manifest_path.exists():
                        manifest = session_store.read_json(str(manifest_path))
                        if manifest:
                            job_dir = str(potential_job_dir)
                            session_id = session_dir.name
                            break
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    # Find the input file
    for inp in manifest.get('inputs', []):
        if inp.get('uploadId') == input_id:
            filepath = inp.get('path')
            if filepath and os.path.exists(filepath):
                return send_file(
                    filepath,
                    mimetype='application/octet-stream',
                    as_attachment=True,
                    download_name=inp.get('originalFilename', 'input')
                )
    
    return jsonify({'error': 'Input not found'}), 404


@app.route('/api/jobs/<job_id>/worker/outputs', methods=['POST'])
@require_worker_auth
def api_worker_upload_output(job_id):
    """
    Receive output file from remote worker.
    
    Form fields:
    - workerJobId: Worker's job ID
    - inputId: Input file this output corresponds to
    - outputType: Type of output (json, markdown, diarization-json, speaker-markdown)
    
    Files:
    - file: The output file content
    """
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    worker_job_id = request.form.get('workerJobId', '')
    input_id = request.form.get('inputId', '')
    output_type = request.form.get('outputType', 'unknown')
    
    # Find the job across all sessions
    job_dir = None
    manifest_path = None
    manifest = None
    
    if session_store.is_server_mode():
        sessions_root = Path(session_store.data_root()) / 'sessions'
        if sessions_root.exists():
            for session_dir in sessions_root.iterdir():
                if session_dir.is_dir():
                    potential_job_dir = session_dir / 'jobs' / job_id
                    potential_manifest = potential_job_dir / 'job.json'
                    if potential_manifest.exists():
                        manifest = session_store.read_json(str(potential_manifest))
                        if manifest:
                            job_dir = str(potential_job_dir)
                            manifest_path = str(potential_manifest)
                            break
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    # Verify this job is expecting remote worker outputs
    if manifest.get('executionMode') != 'remote':
        return jsonify({'error': 'Job is not in remote execution mode'}), 400
    
    worker_info = manifest.get('worker', {})
    if worker_info.get('workerJobId') != worker_job_id:
        return jsonify({'error': 'Worker job ID mismatch'}), 400
    
    # Save the output file
    outputs_dir = os.path.join(job_dir, 'outputs')
    os.makedirs(outputs_dir, exist_ok=True)
    
    # Sanitize filename
    safe_filename = secure_filename(file.filename or 'output')
    output_path = os.path.join(outputs_dir, safe_filename)
    
    # Prevent overwriting
    if os.path.exists(output_path):
        base, ext = os.path.splitext(safe_filename)
        safe_filename = f"{base}_{session_store.new_id(4)}{ext}"
        output_path = os.path.join(outputs_dir, safe_filename)
    
    file.save(output_path)
    
    # If controller had timed out, note that we're still receiving outputs
    # (Don't clear the flag here - let /worker/complete do that)
    if manifest.get('controllerTimedOut'):
        log_event('info', 'worker_output_after_timeout',
                  jobId=job_id, workerJobId=worker_job_id,
                  filename=safe_filename, outputType=output_type,
                  timedOutAt=manifest.get('controllerTimedOutAt'))
    
    # Register output in manifest
    # Find the original filename for this input
    input_filename = None
    for inp in manifest.get('inputs', []):
        if inp.get('uploadId') == input_id:
            input_filename = inp.get('originalFilename')
            break
    
    output_entry = {
        'id': session_store.new_id(8),
        'filename': safe_filename,
        'path': output_path,
        'type': output_type,
        'forUploadId': input_id,  # Use forUploadId for consistency with local pipeline
        'inputFilename': input_filename,
        'sizeBytes': os.path.getsize(output_path),
        'fromWorker': True,
        'workerJobId': worker_job_id
    }
    
    manifest.setdefault('outputs', []).append(output_entry)
    manifest['updatedAt'] = datetime.now(timezone.utc).isoformat()
    session_store.atomic_write_json(manifest_path, manifest)
    
    log_event('info', 'worker_output_received',
              jobId=job_id, workerJobId=worker_job_id,
              filename=safe_filename, outputType=output_type)
    
    return jsonify({'status': 'ok', 'outputId': output_entry['id']})


@app.route('/api/jobs/<job_id>/worker/complete', methods=['POST'])
@require_worker_auth
def api_worker_complete(job_id):
    """
    Receive job completion notification from remote worker.
    
    Body:
    {
        "workerJobId": "wk_xxx",
        "controllerJobId": "xxx",
        "status": "complete" | "failed" | "canceled",
        "outputs": [...],
        "error": {"code": "...", "message": "..."} | null
    }
    """
    data = request.json or {}
    worker_job_id = data.get('workerJobId', '')
    status = data.get('status', 'complete')
    error = data.get('error')
    
    # Find the job
    job_dir = None
    manifest_path = None
    manifest = None
    
    if session_store.is_server_mode():
        sessions_root = Path(session_store.data_root()) / 'sessions'
        if sessions_root.exists():
            for session_dir in sessions_root.iterdir():
                if session_dir.is_dir():
                    potential_job_dir = session_dir / 'jobs' / job_id
                    potential_manifest = potential_job_dir / 'job.json'
                    if potential_manifest.exists():
                        manifest = session_store.read_json(str(potential_manifest))
                        if manifest:
                            job_dir = str(potential_job_dir)
                            manifest_path = str(potential_manifest)
                            break
    
    if not manifest:
        return jsonify({'error': 'Job not found'}), 404
    
    # Verify worker job ID
    worker_info = manifest.get('worker', {})
    if worker_info.get('workerJobId') != worker_job_id:
        return jsonify({'error': 'Worker job ID mismatch'}), 400
    
    # Update manifest with final status
    now = datetime.now(timezone.utc).isoformat()
    
    # Check if controller had timed out - clear the flag since worker is reporting back
    was_timed_out = manifest.get('controllerTimedOut', False)
    if was_timed_out:
        manifest['controllerTimedOut'] = False
        manifest['controllerTimedOutResolvedAt'] = now
        log_event('info', 'worker_callback_after_timeout',
                  jobId=job_id, workerJobId=worker_job_id,
                  timedOutAt=manifest.get('controllerTimedOutAt'),
                  resolvedAt=now)
    
    # Check how many outputs we actually received
    outputs_received = len(manifest.get('outputs', []))
    worker_reported_outputs = len(data.get('outputs', []))
    
    # Update remote metadata
    if manifest.get('remote') is None:
        manifest['remote'] = {}
    manifest['remote']['completedAt'] = now
    manifest['remote']['workerReportedStatus'] = status
    manifest['remote']['workerReportedOutputs'] = worker_reported_outputs
    manifest['remote']['outputsReceived'] = outputs_received
    manifest['remote']['wasTimedOut'] = was_timed_out
    
    if status == 'complete':
        # Check if we actually received outputs
        if outputs_received == 0 and worker_reported_outputs > 0:
            # Worker says it sent outputs but we didn't receive any
            manifest['status'] = 'failed'
            manifest['error'] = {
                'code': 'REMOTE_OUTPUTS_MISSING',
                'message': f'Worker reported {worker_reported_outputs} outputs but controller received 0. Check worker logs for upload errors.'
            }
            log_event('error', 'worker_outputs_missing',
                      jobId=job_id, workerJobId=worker_job_id,
                      workerReported=worker_reported_outputs, received=outputs_received)
        elif outputs_received == 0:
            # Worker didn't report any outputs either - unusual but not necessarily an error
            manifest['status'] = 'complete'
            manifest['progress'] = {
                'stage': 'complete',
                'percent': 100,
                'totalFiles': manifest.get('progress', {}).get('totalFiles', 0),
                'currentFileIndex': manifest.get('progress', {}).get('totalFiles', 0)
            }
            log_event('warning', 'worker_no_outputs',
                      jobId=job_id, workerJobId=worker_job_id)
        else:
            manifest['status'] = 'complete'
            manifest['progress'] = {
                'stage': 'complete',
                'percent': 100,
                'totalFiles': manifest.get('progress', {}).get('totalFiles', 0),
                'currentFileIndex': manifest.get('progress', {}).get('totalFiles', 0)
            }
    elif status == 'failed':
        if outputs_received > 0:
            manifest['status'] = 'complete_with_errors'
            manifest['error'] = error or {'code': 'WORKER_FAILED_PARTIAL', 'message': 'Worker job failed after producing some outputs'}
            manifest['progress'] = {
                'stage': 'complete',
                'percent': 100,
                'totalFiles': manifest.get('progress', {}).get('totalFiles', 0),
                'currentFileIndex': manifest.get('progress', {}).get('totalFiles', 0)
            }
        else:
            manifest['status'] = 'failed'
            manifest['error'] = error or {'code': 'WORKER_FAILED', 'message': 'Worker job failed'}
    elif status == 'canceled':
        if outputs_received > 0:
            manifest['status'] = 'complete_with_errors'
            manifest['error'] = {'code': 'WORKER_CANCELED_PARTIAL', 'message': 'Worker job was canceled after producing some outputs'}
            manifest['progress'] = {
                'stage': 'complete',
                'percent': 100,
                'totalFiles': manifest.get('progress', {}).get('totalFiles', 0),
                'currentFileIndex': manifest.get('progress', {}).get('totalFiles', 0)
            }
        else:
            manifest['status'] = 'canceled'
            manifest['error'] = {'code': 'WORKER_CANCELED', 'message': 'Worker job was canceled'}
    
    manifest['finishedAt'] = now
    manifest['updatedAt'] = now
    
    session_store.atomic_write_json(manifest_path, manifest)
    
    # Clean up active jobs tracking
    session_id = manifest.get('sessionId')
    if session_id:
        with _active_jobs_lock:
            if (session_id, job_id) in _active_jobs:
                _active_jobs[(session_id, job_id)]['running'] = False
    
    log_event('info', 'worker_job_complete',
              jobId=job_id, workerJobId=worker_job_id,
              status=status, outputsReceived=outputs_received,
              workerReported=worker_reported_outputs, error=error)
    
    return jsonify({'status': 'ok', 'outputsReceived': outputs_received})


@app.route('/api/remote-worker/status', methods=['GET'])
def api_remote_worker_status():
    """Get status of remote worker configuration and availability."""
    try:
        from remote_worker import get_remote_worker_status
        status = get_remote_worker_status()
        return jsonify(status)
    except ImportError:
        return jsonify({
            'configured': False,
            'available': False,
            'error': 'Remote worker module not available'
        })


@app.route('/api/session', methods=['GET'])
def api_get_session():
    """Get session info including CSRF token (server mode only)."""
    if not session_store.is_server_mode():
        return jsonify({'error': 'Not available in local mode'}), 404
    
    session_id = g.session_id
    csrf_token = session_store.get_csrf_token(session_id)
    
    response = jsonify({
        'csrfToken': csrf_token
    })
    response.headers['Cache-Control'] = 'no-store'
    return response


@app.route('/api/jobs', methods=['GET'])
def api_list_jobs():
    """List jobs for current session."""
    session_id = g.session_id
    jobs = session_store.list_jobs(session_id, limit=20)
    return jsonify({'jobs': jobs})


@app.route('/api/mode', methods=['GET'])
def api_get_mode():
    """Get current application mode."""
    return jsonify({
        'mode': session_store.get_app_mode(),
        'isServer': session_store.is_server_mode(),
        'isLocal': session_store.is_local_mode()
    })


@app.route('/api/runtime', methods=['GET'])
def api_get_runtime():
    """Get runtime environment and supported compute backends."""
    env = compute_backend.get_cached_environment()
    
    # Add diarization availability info
    hf_token = os.environ.get('HF_TOKEN') or os.environ.get('HUGGINGFACE_TOKEN')
    
    # Check diarization availability
    pyannote_available = False
    diarization_reason = None
    
    # Check if HF token is present
    hf_token_present = bool(hf_token)
    env['hfTokenPresent'] = hf_token_present
    
    # Check if pyannote can be imported
    try:
        import pyannote.audio
        pyannote_available = True
        env['diarizationImportOk'] = True
    except (ImportError, AttributeError, Exception) as e:
        pyannote_available = False
        env['diarizationImportOk'] = False
        diarization_reason = f'pyannote.audio not available: {type(e).__name__}'
    
    # Check HF access if token is present and pyannote is available
    hf_access_ok = False
    hf_access_missing = []
    hf_access_message = None
    
    if hf_token_present and pyannote_available:
        try:
            from diarization import check_diarization_access
            access_info = check_diarization_access(hf_token)
            hf_access_ok = access_info['ok']
            hf_access_missing = access_info['missing']
            hf_access_message = access_info['message']
        except Exception as e:
            hf_access_ok = False
            hf_access_message = f"Failed to check access: {e}"
    
    env['hfAccessOk'] = hf_access_ok
    env['hfAccessMissingRepos'] = hf_access_missing
    env['hfAccessMessage'] = hf_access_message
    
    # Determine overall diarization availability
    diarization_available = hf_token_present and pyannote_available and hf_access_ok
    env['diarizationAvailable'] = diarization_available
    
    # Set reason if not available
    if not diarization_available:
        reasons = []
        if not hf_token_present:
            reasons.append('HF_TOKEN environment variable not set')
        if not pyannote_available:
            reasons.append('pyannote.audio not available')
        if hf_token_present and pyannote_available and not hf_access_ok:
            if hf_access_missing:
                reasons.append(f"Missing access to: {', '.join(hf_access_missing)}")
            if hf_access_message:
                reasons.append(hf_access_message)
        diarization_reason = '; '.join(reasons)
    
    if diarization_reason:
        env['diarizationReason'] = diarization_reason
    
    # Add version info for debugging
    try:
        import torch
        env['torchVersion'] = torch.__version__
    except:
        env['torchVersion'] = None
    
    try:
        import torchaudio
        env['torchaudioVersion'] = torchaudio.__version__
    except:
        env['torchaudioVersion'] = None
    
    if pyannote_available:
        try:
            import pyannote.audio
            env['pyannoteVersion'] = pyannote.audio.__version__
        except:
            env['pyannoteVersion'] = 'installed'
    else:
        env['pyannoteVersion'] = None
    
    # Add diarization policy configuration (server caps)
    from diarization_policy import get_server_policy_config
    env['diarizationPolicy'] = get_server_policy_config()
    
    # Add build version info
    env['buildCommit'] = os.environ.get('BUILD_COMMIT', 'unknown')
    env['buildTime'] = os.environ.get('BUILD_TIME', 'unknown')
    
    # Add remote worker status
    try:
        from remote_worker import get_remote_worker_status
        env['remoteWorker'] = get_remote_worker_status()
    except ImportError:
        env['remoteWorker'] = {
            'configured': False,
            'available': False,
            'mode': 'off',
            'error': None
        }
    
    return jsonify(env)


@app.route('/api/diarization/policy', methods=['POST'])
def api_compute_diarization_policy():
    """
    Compute effective diarization policy from user preferences.
    
    This endpoint ensures the UI and backend use the same policy computation.
    The UI calls this when diarization options change to preview effective values.
    """
    data = request.get_json() or {}
    
    diarization_enabled = data.get('diarizationEnabled', False)
    diarization_auto_split = data.get('diarizationAutoSplit', False)
    diarization_fast_switching = data.get('diarizationFastSwitching', False)
    requested_max_duration = data.get('requestedMaxDurationSeconds')
    requested_chunk = data.get('requestedChunkSeconds')
    requested_overlap = data.get('requestedOverlapSeconds')
    
    from diarization_policy import compute_diarization_policy, get_server_policy_config
    
    server_config = get_server_policy_config()
    
    policy = compute_diarization_policy(
        diarization_enabled=diarization_enabled,
        diarization_auto_split=diarization_auto_split,
        diarization_fast_switching=diarization_fast_switching,
        requested_max_duration_seconds=requested_max_duration,
        requested_chunk_seconds=requested_chunk,
        requested_overlap_seconds=requested_overlap,
        server_max_duration_seconds=server_config['serverMaxDurationSeconds'],
        default_max_duration_seconds=server_config['defaultMaxDurationSeconds'],
        min_chunk_seconds=server_config['minChunkSeconds'],
        max_chunk_seconds=server_config['maxChunkSeconds'],
        overlap_ratio=server_config['overlapRatio'],
        min_overlap_seconds=server_config['minOverlapSeconds'],
        max_overlap_seconds=server_config['maxOverlapSeconds'],
    )
    
    # Generate user-friendly warnings for any clamped values
    from diarization_policy import get_clamping_warnings, estimate_chunk_count
    warnings = get_clamping_warnings(policy) if policy else []
    
    # Calculate estimated chunks if file duration is provided
    file_duration = data.get('fileDurationSeconds')
    estimated_chunks = None
    if policy and file_duration and file_duration > 0:
        estimated_chunks = estimate_chunk_count(
            file_duration_seconds=file_duration,
            chunk_seconds=policy['chunkSeconds'],
            overlap_seconds=policy['overlapSeconds']
        )
    
    return jsonify({
        'policy': policy,
        'serverConfig': server_config,
        'warnings': warnings,
        'estimatedChunks': estimated_chunks
    })


@app.route('/api/admin/stats', methods=['GET'])
def api_admin_stats():
    """Admin diagnostics endpoint. Requires X-Admin-Token header."""
    admin_token = os.environ.get('ADMIN_TOKEN', '')
    if not admin_token:
        return jsonify({'error': 'Admin endpoint not configured'}), 404
    
    provided_token = request.headers.get('X-Admin-Token', '')
    if not provided_token or not secrets.compare_digest(provided_token, admin_token):
        return jsonify({'error': 'Unauthorized'}), 401
    
    import hashlib
    
    # Count sessions and jobs
    sessions_root = os.path.join(session_store.data_root(), 'sessions')
    session_count = 0
    jobs_by_status = {}
    total_disk_mb = 0.0
    
    if os.path.exists(sessions_root):
        for session_id in os.listdir(sessions_root):
            session_path = os.path.join(sessions_root, session_id)
            if not os.path.isdir(session_path):
                continue
            session_count += 1
            
            # Count disk usage
            for dirpath, dirnames, filenames in os.walk(session_path):
                for filename in filenames:
                    try:
                        total_disk_mb += os.path.getsize(os.path.join(dirpath, filename)) / (1024 * 1024)
                    except OSError:
                        pass
            
            # Count jobs by status
            jobs_path = os.path.join(session_path, 'jobs')
            if os.path.exists(jobs_path):
                for job_id in os.listdir(jobs_path):
                    manifest = session_store.read_json(os.path.join(jobs_path, job_id, 'job.json'))
                    if manifest:
                        status = manifest.get('status', 'unknown')
                        jobs_by_status[status] = jobs_by_status.get(status, 0) + 1
    
    response = jsonify({
        'sessionCount': session_count,
        'jobsByStatus': jobs_by_status,
        'totalDiskUsageMB': round(total_disk_mb, 2),
        'limits': {
            'maxUploadMB': session_store.get_max_upload_mb(),
            'maxSessionMB': session_store.get_max_session_mb(),
            'maxFilesPerJob': session_store.get_max_files_per_job(),
            'sessionTTLHours': session_store.get_session_ttl_hours(),
            'jobStaleMins': session_store.get_job_stale_minutes()
        }
    })
    response.headers['Cache-Control'] = 'no-store'
    return response


def _require_admin_token():
    """
    Check for valid admin token. Returns error response tuple if invalid, None if valid.
    """
    admin_token = os.environ.get('ADMIN_TOKEN', '')
    if not admin_token:
        return jsonify({'error': 'Admin endpoint not configured (ADMIN_TOKEN not set)'}), 404
    
    provided_token = request.headers.get('X-Admin-Token', '')
    if not provided_token or not secrets.compare_digest(provided_token, admin_token):
        return jsonify({'error': 'Unauthorized'}), 401
    
    return None


@app.route('/api/admin/remote-worker', methods=['GET'])
def api_admin_remote_worker_get():
    """
    Get remote worker configuration (without token).
    Requires X-Admin-Token header.
    """
    auth_error = _require_admin_token()
    if auth_error:
        return auth_error
    
    try:
        from config_store import get_remote_worker_config
        config = get_remote_worker_config()
        
        # Also include current effective config source
        from remote_worker import get_worker_config
        effective = get_worker_config()
        config['configSource'] = effective.get('configSource', 'unknown')
        
        response = jsonify(config)
        response.headers['Cache-Control'] = 'no-store'
        return response
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/admin/remote-worker', methods=['POST'])
def api_admin_remote_worker_post():
    """
    Update remote worker configuration.
    Requires X-Admin-Token header.
    
    Body: { url?: string, token?: string, mode?: 'off'|'optional'|'required' }
    - If token is omitted, existing token is preserved
    - If token is empty string, token is cleared
    """
    auth_error = _require_admin_token()
    if auth_error:
        return auth_error
    
    try:
        data = request.get_json() or {}
        
        from config_store import save_remote_worker_config
        
        # Extract fields (None means "don't change")
        url = data.get('url')
        token = data.get('token')  # None = keep, '' = clear
        mode = data.get('mode')
        
        # Validate mode if provided
        if mode is not None and mode not in ('off', 'optional', 'required'):
            return jsonify({'error': f"Invalid mode: {mode}. Must be 'off', 'optional', or 'required'"}), 400
        
        # Save config
        saved = save_remote_worker_config(url=url, token=token, mode=mode)
        
        # Clear the worker status cache so next request uses new config
        from remote_worker import _worker_status_cache, _worker_status_cache_time
        _worker_status_cache.clear()
        _worker_status_cache_time.clear()
        
        return jsonify(saved)
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/admin/remote-worker/test', methods=['POST'])
def api_admin_remote_worker_test():
    """
    Test connection to remote worker.
    Requires X-Admin-Token header.
    
    Body: { url?: string, token?: string }
    - If url/token omitted, uses saved config
    """
    auth_error = _require_admin_token()
    if auth_error:
        return auth_error
    
    try:
        import time
        import requests as req
        from urllib.parse import urljoin
        
        data = request.get_json() or {}
        
        # Get URL and token (from request or saved config)
        url = data.get('url')
        token = data.get('token')
        
        if not url or not token:
            # Fall back to saved config
            from config_store import get_remote_worker_url, get_remote_worker_token
            if not url:
                url = get_remote_worker_url()
            if not token:
                token = get_remote_worker_token()
        
        if not url:
            return jsonify({'ok': False, 'error': 'No URL provided or saved'}), 400
        if not token:
            return jsonify({'ok': False, 'error': 'No token provided or saved'}), 400
        
        # Test connection
        start_time = time.time()
        try:
            resp = req.get(
                urljoin(url, '/v1/ping'),
                headers={'Authorization': f'Bearer {token}'},
                timeout=10
            )
            latency_ms = int((time.time() - start_time) * 1000)
            
            if resp.status_code == 200:
                ping_data = resp.json()
                return jsonify({
                    'ok': True,
                    'latencyMs': latency_ms,
                    'identity': ping_data.get('identity'),
                    'version': ping_data.get('version'),
                    'gpu': ping_data.get('gpu'),
                    'gpuName': ping_data.get('gpuName'),
                    'diarization': ping_data.get('diarization'),
                    'warnings': ping_data.get('warnings')
                })
            elif resp.status_code in (401, 403):
                return jsonify({
                    'ok': False,
                    'error': 'Authentication failed (invalid token)',
                    'latencyMs': latency_ms
                })
            else:
                return jsonify({
                    'ok': False,
                    'error': f'Worker returned status {resp.status_code}',
                    'latencyMs': latency_ms
                })
        except req.exceptions.Timeout:
            return jsonify({'ok': False, 'error': 'Connection timed out'})
        except req.exceptions.ConnectionError as e:
            return jsonify({'ok': False, 'error': f'Connection failed: {str(e)[:100]}'})
        except Exception as e:
            return jsonify({'ok': False, 'error': str(e)[:100]})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


def run_with_webview():
    """Run the app with PyWebView for a native window experience."""
    import socket
    import threading
    
    def find_free_port():
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            s.bind(('127.0.0.1', 0))
            s.listen(1)
            port = s.getsockname()[1]
        return port
    
    port = find_free_port()
    
    # Start Flask in a background thread
    def run_flask():
        app.run(debug=False, host='127.0.0.1', port=port, use_reloader=False)
    
    flask_thread = threading.Thread(target=run_flask, daemon=True)
    flask_thread.start()
    
    # Give Flask a moment to start
    import time
    time.sleep(0.5)
    
    # Create native window with PyWebView
    import webview
    
    # Get the app icon path
    import os
    script_dir = os.path.dirname(os.path.abspath(__file__))
    icon_path = os.path.join(script_dir, 'AppIcon.icns')
    if not os.path.exists(icon_path):
        icon_path = None
    
    window = webview.create_window(
        'Bulk Transcribe',
        f'http://localhost:{port}',
        width=1100,
        height=850,
        min_size=(800, 600)
    )
    webview.start(icon=icon_path)


def run_with_browser():
    """Run the app with browser (fallback mode)."""
    import socket
    import webbrowser
    import threading
    
    def find_free_port():
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            s.bind(('127.0.0.1', 0))
            s.listen(1)
            port = s.getsockname()[1]
        return port
    
    port = find_free_port()
    print("Starting Bulk Transcribe...")
    print(f"Open http://localhost:{port} in your browser")
    
    def open_browser():
        import time
        time.sleep(1.5)
        webbrowser.open(f'http://localhost:{port}')
    
    threading.Thread(target=open_browser, daemon=True).start()
    app.run(debug=False, host='127.0.0.1', port=port)


# Required for PyInstaller on macOS - must be at module level
from multiprocessing import freeze_support

if __name__ == '__main__':
    freeze_support()
    
    import sys
    
    # Check for --browser flag to force browser mode
    if '--browser' in sys.argv:
        run_with_browser()
    else:
        try:
            run_with_webview()
        except ImportError:
            print("PyWebView not available, falling back to browser mode")
            run_with_browser()
